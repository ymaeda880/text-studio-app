# -*- coding: utf-8 -*-
# lib/image_embed/word_builder.py
# ============================================================
# Word画像埋込 Word写真一覧表作成
# ============================================================
from __future__ import annotations

from io import BytesIO
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

def _set_cell_shading(cell, fill: str) -> None:
    """
    Word表セルの背景色を設定する。
    fill は RRGGBB。
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))

    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:fill"), fill)


def _clear_cell(cell) -> None:
    cell.text = ""


def _write_center_text(
    *,
    cell,
    title: str,
    description: str,
) -> None:
    _clear_cell(cell)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)

    if description:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(description)
        run2.font.size = Pt(8)

# ============================================================
# image convert helper
# ============================================================
def _convert_image_for_docx(
    *,
    image_path: Path,
) -> BytesIO:
    """
    python-docx が読める形式に画像を変換する。

    InBoxのサムネは Streamlit では表示できても，
    python-docx が直接読めない場合があるため，
    PillowでPNGへ変換して BytesIO として渡す。
    """
    bio = BytesIO()

    with Image.open(image_path) as img:
        if img.mode not in ("RGB", "RGBA"):
            img = img.convert("RGB")

        img.save(bio, format="PNG")

    bio.seek(0)
    return bio

def _add_center_image(
    *,
    cell,
    image_path: Path,
    width_inches: float,
) -> None:
    _clear_cell(cell)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    image_stream = _convert_image_for_docx(
        image_path=image_path,
    )

    run = p.add_run()
    run.add_picture(image_stream, width=Inches(width_inches))


def build_photo_table_docx_bytes(
    *,
    inbox_root: Path,
    sub: str,
    selected_df: pd.DataFrame,
    document_title: str,
    columns: int,
    image_width_inches: float,
    cell_background_hex: str,
) -> bytes:
    """
    選択画像からWord写真一覧表を作成し，docx bytesで返す。

    Word表は，
    - タイトル行
    - 画像行
    を1セットとして繰り返す。
    """
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    title_text = str(document_title or "").strip()
    if title_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title_text)
        r.bold = True
        r.font.size = Pt(16)

    columns = max(1, int(columns))
    n = len(selected_df)
    block_rows = (n + columns - 1) // columns
    table_rows = block_rows * 2

    table = doc.add_table(rows=table_rows, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    idx = 0

    for block_i in range(block_rows):
        title_row_idx = block_i * 2
        image_row_idx = title_row_idx + 1

        for col_i in range(columns):
            title_cell = table.cell(title_row_idx, col_i)
            image_cell = table.cell(image_row_idx, col_i)

            title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            image_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            _set_cell_shading(title_cell, cell_background_hex)
            _set_cell_shading(image_cell, cell_background_hex)

            if idx >= n:
                title_cell.text = ""
                image_cell.text = ""
                continue

            row = selected_df.iloc[idx]

            title = str(row.get("タイトル") or "").strip()
            desc = str(row.get("説明") or "").strip()

            image_path = Path(str(row.get("path") or "").strip())

            _write_center_text(
                cell=title_cell,
                title=title,
                description=desc,
            )

            if not str(image_path).strip() or str(image_path) == ".":
                image_cell.text = "元画像ファイルのパスが取得できません。"
            elif not image_path.exists():
                image_cell.text = f"元画像ファイルが見つかりません: {image_path}"
            elif not image_path.is_file():
                image_cell.text = f"元画像ファイルではありません: {image_path}"
            else:
                _add_center_image(
                    cell=image_cell,
                    image_path=image_path,
                    width_inches=image_width_inches,
                )

            idx += 1

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()