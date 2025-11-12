# lib/text_loaders.py
# ------------------------------------------------------------
# 📥 テキスト読込ユーティリティ（.txt / .docx / .pdf / paste）
#  - UI（Streamlit等）に依存しない純粋関数群
#  - 例外は RuntimeError を中心に投げる（UI側で捕捉して表示）
# ------------------------------------------------------------
from __future__ import annotations
from io import BytesIO
from typing import List, Dict, Union

# ========== 基本ローダ（.txt / .docx / .pdf） ==========
def read_txt(file_or_bytes: Union[bytes, "UploadedFile"]) -> str:
    """
    .txt を文字列化（エンコード自動判定）。
    file_or_bytes: bytes か、.read() できるオブジェクト（Streamlit UploadedFile 等）
    """
    data = file_or_bytes if isinstance(file_or_bytes, (bytes, bytearray)) else file_or_bytes.read()
    for enc in ("utf-8", "utf-16", "shift_jis", "cp932"):
        try:
            return bytes(data).decode(enc)
        except Exception:
            continue
    return bytes(data).decode("utf-8", errors="ignore")


def read_docx(file_or_bytes: Union[bytes, "UploadedFile"]) -> str:
    """
    .docx を段落＋表から抽出して文字列化。
    """
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError("python-docx が必要です。`pip install python-docx` を実行してください。") from e

    data = file_or_bytes if isinstance(file_or_bytes, (bytes, bytearray)) else file_or_bytes.read()
    doc = Document(BytesIO(data))
    texts: List[str] = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts).strip()


def extract_pdf_text(data: bytes) -> Dict[str, Union[str, int]]:
    """
    PDFのテキスト層を抽出し、統計情報を返す。
    returns: {"text": str, "visible": int, "pages": int}
      - text: 改行付きの抽出テキスト
      - visible: 空白以外の印字可能文字数（≈テキスト層の有無の目安）
      - pages: PDFページ数
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise RuntimeError("pymupdf が必要です。`pip install pymupdf` を実行してください。") from e

    try:
        doc = fitz.open(stream=data, filetype="pdf")
    except Exception as e:
        raise RuntimeError(f"PDFを開けませんでした: {e}") from e

    if doc.is_encrypted:
        doc.close()
        raise RuntimeError("このPDFは暗号化（パスワード保護）されています。")

    texts: List[str] = []
    try:
        for page in doc:
            t = page.get_text("text") or ""
            texts.append(t.strip())
        pages = len(doc)
    finally:
        doc.close()

    full = "\n".join(texts).strip()
    visible = sum(ch.isprintable() and not ch.isspace() for ch in full)
    return {"text": full, "visible": int(visible), "pages": int(pages)}


def load_text_generic(uploaded_file) -> str:
    """
    拡張子で簡易分岐して文字列化（.txt / .docx / .pdf）。
    - .pdf はテキスト層がない場合でも空文字を返し得る（UI側でOCRを案内する想定）
    """
    name = (uploaded_file.name or "").lower()
    if name.endswith(".txt"):
        return read_txt(uploaded_file)
    elif name.endswith(".docx"):
        return read_docx(uploaded_file)
    elif name.endswith(".pdf"):
        data = uploaded_file.read()
        stats = extract_pdf_text(data)
        return stats["text"].strip()
    else:
        raise RuntimeError("対応形式は .txt / .docx / .pdf に限られます。")

# ========== 貼り付け（paste）関連ユーティリティ ==========
def normalize_newlines(text: str) -> str:
    """改行コードを LF に正規化。"""
    return text.replace("\r\n", "\n").replace("\r", "\n")

def strip_bom(text: str) -> str:
    """先頭のBOMを除去。"""
    return text.lstrip("\ufeff")

def collapse_blank_lines(text: str, keep: int = 1) -> str:
    """
    連続する空行を keep 行までに圧縮。keep>=1 を推奨。
    """
    lines = normalize_newlines(text).split("\n")
    out: List[str] = []
    blank_run = 0
    for ln in lines:
        if ln.strip() == "":
            blank_run += 1
            if blank_run <= keep:
                out.append("")
        else:
            blank_run = 0
            out.append(ln)
    return "\n".join(out)

def trim_trailing_spaces(text: str) -> str:
    """各行末の余分な空白を削除。"""
    return "\n".join([ln.rstrip() for ln in normalize_newlines(text).split("\n")])

def load_text_from_paste(
    text: str,
    *,
    normalize: bool = True,
    collapse_blanks: bool = False,
    keep_blank_lines: int = 1,
    trim_trailing: bool = True,
) -> str:
    """
    貼り付けテキストを整形して返す。
    - normalize=True: 改行をLFに統一しBOM除去
    - collapse_blanks=True: 連続空行を keep_blank_lines 行まで圧縮
    - trim_trailing=True: 各行末の余分な空白を削る
    """
    if text is None:
        return ""
    s = text
    if normalize:
        s = strip_bom(normalize_newlines(s))
    if trim_trailing:
        s = trim_trailing_spaces(s)
    if collapse_blanks:
        s = collapse_blank_lines(s, keep=keep_blank_lines)
    return s.strip()
