# -*- coding: utf-8 -*-
# lib/table/helpers.py
#
# Word表作成用のユーティリティ群
# - テキスト→2次元配列パース
# - 列幅計算
# - マーカー(<同上>, <同左>)による結合情報計算
# - docx への反映
# - HTML プレビュー生成

from __future__ import annotations
import io, csv, re, unicodedata
from typing import List, Tuple, Optional

import pandas as pd  # 将来拡張用に残しておく（未使用でもOK）

from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# セル内改行マーカー（<改行> / ＜改行＞ の両方に対応）
_BREAK_RE = re.compile(r"[<＜]\s*改行\s*[>＞]")

# =========================================================
# パース関連
# =========================================================
def _detect_delimiter(text: str) -> str:
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return "\t"
    head = lines[0]
    if "\t" in head:
        return "\t"
    if head.count(",") >= head.count(";") and head.count(",") > 0:
        return ","
    if head.count(";") > 0:
        return ";"
    if re.search(r"\s{2,}", head):
        return r"\s+"
    return "\t"


def _parse_table(text: str) -> List[List[str]]:
    """
    Excel からの貼り付けテキストなどを 2次元配列（rows[row][col]）に変換。
    列数の足りない行は右側を "" で埋めて揃える。
    """
    text = text.strip("\n\r ")
    if not text:
        return []
    delim = _detect_delimiter(text)
    if delim == r"\s+":
        rows = [re.split(r"\s{2,}", ln.strip()) for ln in text.splitlines() if ln.strip()]
    else:
        reader = csv.reader(io.StringIO(text), delimiter=("\t" if delim == "\t" else delim))
        rows = [list(r) for r in reader]

    # 完全な空行は除外
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return []

    # 列数を最大列数に合わせて右側パディング
    maxc = max(len(r) for r in rows)
    rows = [r + [""] * (maxc - len(r)) for r in rows]
    return rows


# =========================================================
# 色・フォント関連
# =========================================================
def _hex_to_rgb(h: str) -> Tuple[int, int, int]:
    s = h.strip().lstrip("#")
    if len(s) == 3:
        s = "".join(c * 2 for c in s)
    return int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16)


def _apply_font_run(run, font_name: str, size_pt: int, color_hex: str, bold: bool = False):
    """
    段落 run にフォント名・サイズ・色・太字を適用。
    日本語フォントを eastAsia にも設定する。
    """
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    r, g, b = _hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    run.bold = bold


def _set_cell_shading(cell, hex_color: Optional[str]):
    """
    セル背景色を塗る（None/空なら何もしない）。
    """
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)


def _set_cell_border(cell, **kwargs):
    """
    cell に対して罫線を設定するヘルパ。
    edge="left/right/top/bottom/insideH/insideV" ごとに
    (val, sz, color) を指定する。
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "right", "top", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            val, sz, color = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), val)
            tag.set(qn("w:sz"), str(sz))
            tag.set(qn("w:color"), color)
            tcBorders.append(tag)


# =========================================================
# 列幅関連
# =========================================================
def _visual_len(s: str) -> int:
    """
    全角=2, 半角=1 とみなした「見かけ上の文字幅」を返す。
    """
    if s is None:
        return 0
    t = 0
    for ch in str(s):
        t += 2 if unicodedata.east_asian_width(ch) in ("W", "F") else 1
    return t


def _compute_col_widths_cm(
    rows: List[List[str]],
    total_cm: float = 16.0,
    min_cm: float = 2.2,
    max_cm: float = 8.0,
) -> List[float]:
    """
    各列の「見かけ上の文字幅」に応じて列幅(cm)を計算する。
    total_cm を全幅として、min_cm〜max_cm の範囲でクランプし、
    最後に合計がちょうど total_cm になるよう微調整。
    """
    if not rows:
        return []
    n = len(rows[0])
    scores: List[int] = []
    for c in range(n):
        mx = 0
        for r in range(len(rows)):
            mx = max(mx, _visual_len(rows[r][c]))
        scores.append(int(mx * 1.1) or 1)
    ssum = sum(scores) or 1
    raw = [total_cm * (sc / ssum) for sc in scores]
    clamped = [max(min_cm, min(max_cm, x)) for x in raw]
    # 合計補正
    diff = total_cm - sum(clamped)
    if abs(diff) > 1e-6:
        room = [(max_cm - w) if diff > 0 else (w - min_cm) for w in clamped]
        room_sum = sum(x for x in room if x > 0) or 1
        adj = [(diff * (r / room_sum) if r > 0 else 0) for r in room]
        clamped = [w + a for w, a in zip(clamped, adj)]
    return clamped


def _widths_to_pct(widths_cm: List[float]) -> List[float]:
    """
    列幅(cm)のリストを合計100%に正規化したパーセント表現に変換。
    """
    tot = sum(widths_cm) or 1
    return [round(w * 100.0 / tot, 2) for w in widths_cm]


def _apply_docx_col_widths(table, widths_cm: List[float]):
    """
    docx.Table に対して固定レイアウト + 表幅 + グリッド + 各セル幅を設定。
    """
    # 固定レイアウト + 表幅 + グリッド + セル幅
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    # 表幅
    tot_cm = sum(widths_cm) or 1.0
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    # 1cm ≒ 567 dxa
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(tot_cm * 567)))

    # グリッド作り直し
    for child in list(table._tbl.iterchildren()):
        if child.tag == qn("w:tblGrid"):
            table._tbl.remove(child)
    grid = OxmlElement("w:tblGrid")
    for wcm in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(wcm * 567)))
        grid.append(gc)
    table._tbl.insert(1, grid)

    # 各セル幅
    for c, wcm in enumerate(widths_cm):
        for r in range(len(table.rows)):
            table.cell(r, c).width = Cm(wcm)


# =========================================================
# マーカー方式の結合
# =========================================================
_MARK_UP_RE   = re.compile(r'^[<＜]\s*同上\s*[>＞]$')
_MARK_LEFT_RE = re.compile(r'^[<＜]\s*同左\s*[>＞]$')


def _is_mark_up(s: Optional[str]) -> bool:
    if s is None:
        return False
    s = str(s).strip()
    return bool(_MARK_UP_RE.match(s))


def _is_mark_left(s: Optional[str]) -> bool:
    if s is None:
        return False
    s = str(s).strip()
    return bool(_MARK_LEFT_RE.match(s))


def _compute_spans_markers(
    rows: List[List[str]],
    *,
    use_up: bool = True,      # 「<同上＞」を有効にする
    use_left: bool = True,    # 「<同左＞」を有効にする
    header_rows: int = 1,     # 今は未使用だが将来拡張用
    strict_rect: bool = True  # 横結合は矩形制約（rowspan一致）を要求
):
    """
    マーカー方式の結合規則：
      ・セルが「<同上＞」なら “上の非マーカーセル” に縦結合（連続なら上端まで伸長）
      ・セルが「<同左＞」なら “左の非マーカーセル” に横結合（連続なら左端まで伸長）
    ※ 同値自動結合は行わない
    ※ ヘッダー行も含めて全行でマーカーを解釈する
    """
    if not rows:
        return []

    R, C = len(rows), len(rows[0])
    spans = [[{"rowspan": 1, "colspan": 1, "skip": False} for _ in range(C)] for __ in range(R)]

    # 以前の ignore_header ロジックは廃止。全行を対象とする。
    start_r = 0

    # --- 1) 縦方向（<同上＞）---
    if use_up:
        for r in range(start_r, R):
            for c in range(C):
                if not _is_mark_up(rows[r][c]):
                    continue
                a = r - 1
                while a >= start_r and _is_mark_up(rows[a][c]):
                    a -= 1
                if a >= 0 and not _is_mark_up(rows[a][c]):
                    spans[a][c]["rowspan"] += 1
                    spans[r][c]["skip"] = True

    # --- 2) 横方向（<同左＞）---
    if use_left:
        for r in range(start_r, R):
            for c in range(C):
                if not _is_mark_left(rows[r][c]):
                    continue
                b = c - 1
                while b >= 0 and _is_mark_left(rows[r][b]):
                    b -= 1
                if b >= 0 and (not _is_mark_left(rows[r][b])) and (not spans[r][b]["skip"]):
                    if strict_rect:
                        base_rs = spans[r][b]["rowspan"]
                        ok = True
                        for cc in range(b, c):
                            if spans[r][cc]["rowspan"] != base_rs:
                                ok = False
                                break
                        if not ok:
                            continue
                    spans[r][b]["colspan"] += 1
                    spans[r][c]["skip"] = True

    return spans


# =========================================================
# docx への反映（ヘッダー行 + ヘッダー列）
# =========================================================
def _merge_docx_by_spans(
    table,
    rows: List[List[str]],
    spans,
    font_name: str,
    base_size: int,
    header_size: int,
    header_fg: str,
    body_fg: str,
    header_bg: Optional[str],
    body_bg: Optional[str],
    zebra: bool,
    header_same: bool,
    body_bg_on: bool,
    header_rows: int = 1,      # 先頭から何行をヘッダー行とするか（0 も許可）
    header_cols: int = 0,      # 左から何列を「ヘッダー列」とするか
):
    """
    spans 情報に基づいて docx.Table にテキスト・背景色・フォントを適用。
    - 先頭 header_rows 行は「ヘッダー行」として header_bg / header_fg / header_size / 太字 を適用
    - それ以外の行でも、左から header_cols 列は「ヘッダー列」として同様のスタイルを適用
    - それ以外は本文として body_bg / body_fg / base_size を適用
    """
    R, C = len(rows), len(rows[0])

    hb = header_bg or "#EEEEEE"
    zebra_alt = "#F7F9FC"
    body_fill_default = (body_bg if body_bg_on else None)

    # ガード
    header_rows = max(0, min(header_rows, R))
    header_cols = max(0, min(header_cols, C))

    # --- 1) 全セル空に ---
    for r in range(R):
        for c in range(C):
            table.cell(r, c).text = ""

    # --- 2) テキストを anchor セルだけへ ---
    for r in range(R):
        for c in range(C):
            if spans[r][c]["skip"]:
                continue

            cell = table.cell(r, c)
            raw = str(rows[r][c])

            raw = raw.replace("\r\n", "\n").replace("\r", "\n")
            parts_raw = _BREAK_RE.split(raw)

            parts: List[str] = []
            for p in parts_raw:
                p = p.strip("\n")
                if p.strip() == "":
                    continue
                parts.append(p)

            if not parts:
                parts = [""]

            p0 = cell.paragraphs[0]
            p0.text = parts[0]
            for extra in parts[1:]:
                cell.add_paragraph(extra)

    # --- 3) 背景 + フォント ---
    for r in range(R):
        for c in range(C):
            if spans[r][c]["skip"]:
                continue
            cell = table.cell(r, c)

            # ヘッダー行 or ヘッダー列 ならヘッダー扱い
            is_header_row = (r < header_rows)
            is_header_col = (c < header_cols)
            is_header = is_header_row or is_header_col

            # 背景色
            if is_header:
                _set_cell_shading(cell, hb)
            else:
                zebra_idx = r - header_rows
                fill = zebra_alt if (zebra and (zebra_idx % 2 == 1)) else body_fill_default
                if fill:
                    _set_cell_shading(cell, fill)

            # フォント
            for p in cell.paragraphs:
                for run in p.runs:
                    if is_header:
                        _apply_font_run(run, font_name, header_size, header_fg, bold=True)
                    else:
                        _apply_font_run(run, font_name, base_size, body_fg, bold=False)


# =========================================================
# 罫線（物理結合なし版）
# =========================================================
def _apply_table_borders_robust(
    table,
    spans,
    *,
    inner_h: bool = True,
    inner_v: bool = True,
    outer: bool = True,
    outer_mode: str = "box",   # ← 追加： "box" or "top_bottom"
    sz_inner: int = 6,
    sz_outer: int = 12,
    color: str = "000000",
):
    """
    物理結合なし版の罫線設定：
      - spans から anchor[r][c] を作成
      - 同じ anchor 同士の境目には内側罫線を引かない
      - anchor が変わる境目だけに線を引く

    outer_mode:
      - "box"       : 外枠を四辺すべてに引く（従来どおり）
      - "top_bottom": 外枠の「上」と「下」だけ太線を引く（左右は引かない）
    """
    R = len(spans)
    if R == 0:
        return
    C = len(spans[0])

    # --- anchor マップ作成 ---
    anchor: List[List[Optional[Tuple[int, int]]]] = [[None for _ in range(C)] for __ in range(R)]
    for r in range(R):
        for c in range(C):
            info = spans[r][c]
            if info["skip"]:
                continue
            rs = info["rowspan"]
            cs = info["colspan"]
            for rr in range(r, r + rs):
                for cc in range(c, c + cs):
                    anchor[rr][cc] = (r, c)

    # --- 内側 横 ---
    if inner_h:
        for r in range(R - 1):
            for c in range(C):
                if anchor[r][c] != anchor[r + 1][c]:
                    up = table.cell(r, c)
                    down = table.cell(r + 1, c)
                    _set_cell_border(up, bottom=("single", sz_inner, color))
                    _set_cell_border(down, top=("single", sz_inner, color))

    # --- 内側 縦 ---
    if inner_v:
        for r in range(R):
            for c in range(C - 1):
                if anchor[r][c] != anchor[r][c + 1]:
                    left = table.cell(r, c)
                    right = table.cell(r, c + 1)
                    _set_cell_border(left, right=("single", sz_inner, color))
                    _set_cell_border(right, left=("single", sz_inner, color))

    # --- 外周 ---
    if outer and anchor:
        # 上下だけにしたい場合と、全周描きたい場合を分岐
        if outer_mode == "top_bottom":
            # 上端
            for c in range(C):
                for r in range(R):
                    if anchor[r][c] is None:
                        continue
                    if r == 0 or anchor[r - 1][c] != anchor[r][c]:
                        cell = table.cell(r, c)
                        _set_cell_border(cell, top=("single", sz_outer, color))
                        break

            # 下端
            for c in range(C):
                for r in range(R - 1, -1, -1):
                    if anchor[r][c] is None:
                        continue
                    if r == R - 1 or anchor[r + 1][c] != anchor[r][c]:
                        cell = table.cell(r, c)
                        _set_cell_border(cell, bottom=("single", sz_outer, color))
                        break

        else:
            # outer_mode="box"（従来どおり四辺すべて）
            # 上端
            for c in range(C):
                for r in range(R):
                    if anchor[r][c] is None:
                        continue
                    if r == 0 or anchor[r - 1][c] != anchor[r][c]:
                        cell = table.cell(r, c)
                        _set_cell_border(cell, top=("single", sz_outer, color))
                        break

            # 下端
            for c in range(C):
                for r in range(R - 1, -1, -1):
                    if anchor[r][c] is None:
                        continue
                    if r == R - 1 or anchor[r + 1][c] != anchor[r][c]:
                        cell = table.cell(r, c)
                        _set_cell_border(cell, bottom=("single", sz_outer, color))
                        break

            # 左端
            for r in range(R):
                for c in range(C):
                    if anchor[r][c] is None:
                        continue
                    if c == 0 or anchor[r][c - 1] != anchor[r][c]:
                        cell = table.cell(r, c)
                        _set_cell_border(cell, left=("single", sz_outer, color))
                        break

            # 右端
            for r in range(R):
                for c in range(C - 1, -1, -1):
                    if anchor[r][c] is None:
                        continue
                    if c == C - 1 or anchor[r][c + 1] != anchor[r][c]:
                        cell = table.cell(r, c)
                        _set_cell_border(cell, right=("single", sz_outer, color))
                        break

# =========================================================
# HTML プレビュー生成（ヘッダー行 + ヘッダー列）
# =========================================================
def _build_html_table_with_spans(
    rows: List[List[str]],
    spans,
    *,
    header_rows: int = 1,
    header_bg: Optional[str],
    header_fg: str,
    header_size: int,
    header_bold: bool,
    body_bg: Optional[str],
    body_fg: str,
    body_size: int,
    font_name: str,
    zebra: bool,
    inner_h: bool,
    inner_v: bool,
    outer: bool,
    outer_mode: str = "box",   # ★ 追加："box" / "top_bottom"
    note_text: Optional[str] = None,
    col_width_pct: Optional[List[float]] = None,
    row_header_cols: int = 0,   # 左から何列を「ヘッダー列」として扱うか（HTML側）
) -> str:

    if not rows:
        return ""
    R, C = len(rows), len(rows[0])

    header_rows = max(0, min(header_rows, R))
    row_header_cols = max(0, min(row_header_cols, C))

    # ------------------------------------------------------------
    # ★ outer_mode で table 外枠を決める
    # ------------------------------------------------------------
    border_css = "border-collapse:collapse;"

    if not outer:
        # 外枠なし
        outer_css = "border:0;"
    else:
        if outer_mode == "top_bottom":
            # 上下だけ線：左右は 0
            outer_css = "border-style:solid; border-color:#000; border-width:1px 0;"
        else:
            # "box"（デフォルト）：全周のボックス
            outer_css = "border:1px solid #000;"

    # ------------------------------------------------------------
    # inner_h / inner_v によるセル内罫線
    # ------------------------------------------------------------
    if inner_h and inner_v:
        td_border_css = "border:1px solid #000;"
    elif inner_h and not inner_v:
        td_border_css = "border-style:solid; border-color:#000; border-width:1px 0;"
    elif inner_v and not inner_h:
        td_border_css = "border-style:solid; border-color:#000; border-width:0 1px;"
    else:
        td_border_css = "border:0;"

    td_base = "padding:6px 10px; vertical-align:top;"
    th_weight = "font-weight:700;" if header_bold else "font-weight:400;"
    zebra_alt = "#F7F9FC"

    # ------------------------------------------------------------
    # HTML table 開始
    # ------------------------------------------------------------
    html: List[str] = [f'<table style="{border_css}{outer_css} width:100%;">']

    # colgroup（列幅%指定）
    if col_width_pct:
        html.append("<colgroup>")
        for p in col_width_pct:
            html.append(f'<col style="width:{p}%">')
        html.append("</colgroup>")

    hb = header_bg or "#EEEEEE"

    # ------------------------------------------------------------
    # thead：ヘッダー行
    # ------------------------------------------------------------
    if header_rows > 0:
        html.append("<thead>")
        for r in range(header_rows):
            html.append("<tr>")
            for c in range(C):
                info = spans[r][c]
                if info["skip"]:
                    continue

                attrs: List[str] = []
                if info["rowspan"] > 1:
                    attrs.append(f'rowspan="{info["rowspan"]}"')
                if info["colspan"] > 1:
                    attrs.append(f'colspan="{info["colspan"]}"')

                raw_header = str(rows[r][c])
                header_text = _BREAK_RE.sub("<br/>", raw_header)

                html.append(
                    f'<th {" ".join(attrs)} style="{td_base}{td_border_css}'
                    f'background:{hb}; color:{header_fg}; '
                    f'font-family:{font_name}; font-size:{header_size}pt; {th_weight} text-align:left;">'
                    f'{header_text}'
                    f"</th>"
                )
            html.append("</tr>")
        html.append("</thead>")

    # ------------------------------------------------------------
    # tbody
    # ------------------------------------------------------------
    html.append("<tbody>")
    for r in range(header_rows, R):
        zebra_idx = r - header_rows
        row_bg = (zebra_alt if (zebra and zebra_idx % 2 == 1) else (body_bg or "transparent"))
        html.append(f'<tr style="background:{row_bg};">')

        for c in range(C):
            info = spans[r][c]
            if info["skip"]:
                continue

            attrs: List[str] = []
            if info["rowspan"] > 1:
                attrs.append(f'rowspan="{info["rowspan"]}"')
            if info["colspan"] > 1:
                attrs.append(f'colspan="{info["colspan"]}"')

            raw_body = str(rows[r][c])
            body_text = _BREAK_RE.sub("<br/>", raw_body)

            is_header_col = (c < row_header_cols)

            if is_header_col:
                cell_style = (
                    f'{td_base}{td_border_css}'
                    f'background:{hb}; color:{header_fg}; '
                    f'font-family:{font_name}; font-size:{header_size}pt; {th_weight} text-align:left;'
                )
            else:
                cell_style = (
                    f'{td_base}{td_border_css}'
                    f'font-family:{font_name}; font-size:{body_size}pt; color:{body_fg}; text-align:left;'
                )

            html.append(
                f'<td {" ".join(attrs)} style="{cell_style}">{body_text}</td>'
            )

        html.append("</tr>")
    html.append("</tbody></table>")

    # ------------------------------------------------------------
    # 注
    # ------------------------------------------------------------
    if note_text and note_text.strip():
        esc = (
            note_text.strip()
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        ).replace("\n", "<br/>")
        html.append(
            f"<div style='margin-top:8px; font-family:{font_name}; font-size:{body_size}pt; color:#444;'>{esc}</div>"
        )

    return "\n".join(html)
