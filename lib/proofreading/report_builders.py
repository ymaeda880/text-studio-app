# -*- coding: utf-8 -*-
from __future__ import annotations

import io
import datetime as _dt
from typing import Dict, Tuple

from .md_table_parser import parse_plan_md_tables


def build_policy_docx_bytes(
    *,
    original_numbered_preview: str,
    plan_md: str,
    model: str,
    mode: str,
    extra_prompt: str,
    src_name: str,
    format_cost_lines,
    usage_summary: Dict[str, int] | None = None,
    usd_jpy: float | None = None,
) -> Tuple[bytes, str]:
    """Word（.docx）でレポート出力：原文=等幅風/自動折返し、校正方針=Markdown文字列そのまま"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        add_lines = []
        if usage_summary and usd_jpy is not None:
            in_t = int(usage_summary.get("input_tokens", 0))
            out_t = int(usage_summary.get("output_tokens", 0))
            add_lines += [""] + format_cost_lines(model=model, in_t=in_t, out_t=out_t, usd_jpy=usd_jpy)

        data = "\n".join([
            "=== 校正方針レポート（TXTフォールバック） ===",
            f"生成日時: {_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
            f"モデル: {model}",
            f"モード: {mode}",
            f"追加プロンプト: {extra_prompt.strip()}" if extra_prompt.strip() else "",
            f"入力ファイル: {src_name or '-'}",
            *add_lines,
            "\n--- 原文（行番号つき） ---\n",
            original_numbered_preview or "(空)",
            "\n--- 校正方針（Markdown） ---\n",
            plan_md or "(なし)",
        ]).encode("utf-8")
        return data, ".txt"

    doc = Document()
    h = doc.add_heading("校正方針レポート", 0)
    try:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    p = doc.add_paragraph()
    p.add_run("生成日時: ").bold = True
    p.add_run(_dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p.add_run("\nモデル: ").bold = True
    p.add_run(model)
    p.add_run("\nモード: ").bold = True
    p.add_run(mode)
    if extra_prompt.strip():
        p.add_run("\n追加プロンプト: ").bold = True
        p.add_run(extra_prompt.strip())
    p.add_run("\n入力ファイル: ").bold = True
    p.add_run(src_name or "-")

    if usage_summary and (usd_jpy is not None):
        in_t = int(usage_summary.get("input_tokens", 0))
        out_t = int(usage_summary.get("output_tokens", 0))
        for ln in format_cost_lines(model=model, in_t=in_t, out_t=out_t, usd_jpy=float(usd_jpy)):
            p.add_run("\n" + ln)

    doc.add_heading("原文（行番号つき）", level=1)
    para = doc.add_paragraph()
    r = para.add_run(original_numbered_preview if original_numbered_preview else "(空)")
    try:
        r.font.name = "Courier New"
        r.font.size = Pt(10)
    except Exception:
        pass

    doc.add_heading("校正方針（Markdown）", level=1)
    para2 = doc.add_paragraph()
    r2 = para2.add_run(plan_md if plan_md else "(なし)")
    try:
        r2.font.name = "Courier New"
        r2.font.size = Pt(10)
    except Exception:
        pass

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue(), ".docx"


def build_policy_pdf_bytes(
    *,
    original_numbered_preview: str,
    plan_md: str,
    model: str,
    mode: str,
    extra_prompt: str,
    src_name: str,
    format_cost_lines,
    usage_summary: Dict[str, int] | None = None,
    usd_jpy: float | None = None,
) -> bytes:
    """整形PDF：原文は1行=1行の行テーブル、校正方針はMarkdown表をパースして表描画"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.lib.units import mm
        import re
    except Exception:
        return b""

    font_name = None
    for fname in ("HeiseiMin-W3", "HeiseiKakuGo-W5"):
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(fname))
            font_name = fname
            break
        except Exception:
            continue
    if not font_name:
        return b""

    buf = io.BytesIO()
    pagesize = A4
    margin = 18 * mm
    doc = SimpleDocTemplate(
        buf, pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=20 * mm, bottomMargin=18 * mm,
        title="校正方針レポート",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleJP", fontName=font_name, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Meta", fontName=font_name, fontSize=10, leading=13, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle(name="H1", fontName=font_name, fontSize=13, leading=16, spaceBefore=6, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body", fontName=font_name, fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="MonoCJK", fontName=font_name, fontSize=9.5, leading=13, wordWrap="CJK"))

    text_width = pagesize[0] - doc.leftMargin - doc.rightMargin

    def _page_number(canvas, doc_):
        canvas.setFont(font_name, 9)
        canvas.drawRightString(pagesize[0] - doc_.rightMargin, 12 * mm, f"{doc_.page}")

    ZWSP = "&#8203;"
    LONG_TOKEN = re.compile(r"([A-Za-z0-9_/#%~\-\.\?\=&]{30,})")

    def _soften(s: str) -> str:
        return LONG_TOKEN.sub(
            lambda m: ZWSP.join(m.group(1)[i:i+30] for i in range(0, len(m.group(1)), 30)),
            s
        )

    def _make_original_table(text: str) -> Table:
        rows = []
        MAX_CHARS_PER_ROW = 80

        for raw in (text or "").splitlines() or ["(空)"]:
            s = raw.replace("\t", "    ")
            if len(s) > MAX_CHARS_PER_ROW:
                segments = [s[i:i+MAX_CHARS_PER_ROW] for i in range(0, len(s), MAX_CHARS_PER_ROW)]
            else:
                segments = [s]

            for seg in segments:
                seg = _soften(seg).replace("  ", "&nbsp;&nbsp;")
                p = Paragraph(seg if seg else "&nbsp;", styles["MonoCJK"])
                rows.append([p])

        col_w = max(10, text_width - 2)
        t = Table(rows, colWidths=[col_w], splitByRow=1)
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    items = parse_plan_md_tables(plan_md)
    headers = ["行", "重要度", "原文", "修正案", "理由"]
    table_data = [headers] + [[
        it.get("行", ""), it.get("重要度", ""),
        it.get("原文", ""), it.get("修正案", ""), it.get("理由", "")
    ] for it in items]

    col_w = [14 * mm, 18 * mm]
    remain = text_width - sum(col_w)
    col_w += [remain * 0.30, remain * 0.30, remain * 0.40]

    def _p(s: str) -> Paragraph:
        s = (s or "")
        s = s.replace("<br>", "<br/>").replace("<br />", "<br/>")
        s = s.replace("\n", "<br/>")
        return Paragraph(s, styles["Body"])

    table_para = [table_data[0]] + [[_p(x) for x in r] for r in table_data[1:]]
    tbl = Table(table_para, colWidths=col_w, repeatRows=1)

    zebra = colors.Color(0.95, 0.95, 0.98)
    header_bg = colors.Color(0.88, 0.90, 0.95)
    grid = colors.Color(0.75, 0.78, 0.85)

    st_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        ("ALIGN", (0, 0), (1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("GRID", (0, 0), (-1, -1), 0.3, grid),
    ]
    for i in range(1, len(table_para)):
        if i % 2 == 1:
            st_cmds.append(("BACKGROUND", (0, i), (-1, i), zebra))
    tbl.setStyle(TableStyle(st_cmds))

    story = []
    story.append(Paragraph("校正方針レポート", styles["TitleJP"]))

    meta_lines = [
        f"生成日時：{_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        f"モデル：{model}",
        f"モード：{mode}",
        f"入力ファイル：{src_name or '-'}",
    ]
    if (extra_prompt or "").strip():
        meta_lines.insert(3, f"追加プロンプト：{extra_prompt.strip()}")

    if usage_summary and (usd_jpy is not None):
        in_t = int(usage_summary.get("input_tokens", 0))
        out_t = int(usage_summary.get("output_tokens", 0))
        for ln in format_cost_lines(model=model, in_t=in_t, out_t=out_t, usd_jpy=float(usd_jpy)):
            meta_lines.append(ln)

    for ln in meta_lines:
        story.append(Paragraph(ln, styles["Meta"]))

    story.append(Paragraph("原文（行番号つき）", styles["H1"]))
    story.append(_make_original_table(original_numbered_preview or "(空)"))
    story.append(Spacer(1, 6 * mm))

    story.append(Paragraph("校正方針（テーブル）", styles["H1"]))
    story.append(tbl)

    doc.build(story, onFirstPage=_page_number, onLaterPages=_page_number)
    pdf_bytes = buf.getvalue()
    buf.close()
    return pdf_bytes
