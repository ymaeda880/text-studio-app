# -*- coding: utf-8 -*-
# lib/translate/report_builders.py
# ============================================================
# 翻訳レポート出力
# ============================================================

from __future__ import annotations

# ============================================================
# imports
# ============================================================
import datetime as _dt
import html
import io
from typing import Dict, Tuple


# ============================================================
# text 出力
# ============================================================
def build_translation_txt_bytes(
    *,
    original_numbered_preview: str,
    translated_text_numbered: str,
    translated_text: str,
    summary_text: str,
    model: str,
    direction: str,
    extra_prompt: str,
    src_name: str,
    usage_summary: Dict[str, int] | None = None,
) -> bytes:

    # ------------------------------------------------------------
    # usage
    # ------------------------------------------------------------
    usage_lines = []
    if usage_summary:
        usage_lines = [
            f"input_tokens: {usage_summary.get('input_tokens', '')}",
            f"output_tokens: {usage_summary.get('output_tokens', '')}",
            f"total_tokens: {usage_summary.get('total_tokens', '')}",
        ]


    # ------------------------------------------------------------
    # レポート本文
    # ------------------------------------------------------------
    sections = [
        "=== 翻訳結果レポート ===",
        f"生成日時: {_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        f"モデル: {model}",
        f"翻訳方向: {direction}",
        f"追加プロンプト: {extra_prompt.strip()}" if extra_prompt.strip() else "",
        f"入力ファイル: {src_name or '-'}",
        *usage_lines,
        "",
        "--- 原文（番号つき） ---",
        original_numbered_preview or "(空)",
        "",
        "--- 翻訳結果（番号付き） ---",
        translated_text_numbered or "(なし)",
        "",
        "--- 翻訳結果 ---",
        translated_text or "(なし)",
    ]

    # ------------------------------------------------------------
    # 要約
    # ------------------------------------------------------------
    if str(summary_text or "").strip():

        sections.extend([
            "",
            "--- 要約 ---",
            summary_text,
        ])

    data = "\n".join(sections)


    return data.encode("utf-8")


# ============================================================
# Word 出力
# ============================================================
def build_translation_docx_bytes(
    *,
    original_numbered_preview: str,
    translated_text_numbered: str,
    translated_text: str,
    summary_text: str,
    model: str,
    direction: str,
    extra_prompt: str,
    src_name: str,
    usage_summary: Dict[str, int] | None = None,
) -> Tuple[bytes, str]:

    # ------------------------------------------------------------
    # import
    # ------------------------------------------------------------
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return (
            build_translation_txt_bytes(
                original_numbered_preview=original_numbered_preview,
                translated_text=translated_text,
                model=model,
                direction=direction,
                extra_prompt=extra_prompt,
                src_name=src_name,
                usage_summary=usage_summary,
            ),
            ".txt",
        )

    # ------------------------------------------------------------
    # document
    # ------------------------------------------------------------
    doc = Document()

    h = doc.add_heading("翻訳結果レポート", 0)
    try:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    # ------------------------------------------------------------
    # metadata
    # ------------------------------------------------------------
    p = doc.add_paragraph()
    p.add_run("生成日時: ").bold = True
    p.add_run(_dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p.add_run("\nモデル: ").bold = True
    p.add_run(model)
    p.add_run("\n翻訳方向: ").bold = True
    p.add_run(direction)
    if extra_prompt.strip():
        p.add_run("\n追加プロンプト: ").bold = True
        p.add_run(extra_prompt.strip())
    p.add_run("\n入力ファイル: ").bold = True
    p.add_run(src_name or "-")

    if usage_summary:
        p.add_run("\ninput_tokens: ").bold = True
        p.add_run(str(usage_summary.get("input_tokens", "")))
        p.add_run("\noutput_tokens: ").bold = True
        p.add_run(str(usage_summary.get("output_tokens", "")))
        p.add_run("\ntotal_tokens: ").bold = True
        p.add_run(str(usage_summary.get("total_tokens", "")))

    # ------------------------------------------------------------
    # original
    # ------------------------------------------------------------
    doc.add_heading("原文（番号つき）", level=1)
    para = doc.add_paragraph()
    r = para.add_run(original_numbered_preview if original_numbered_preview else "(空)")
    try:
        r.font.name = "Courier New"
        r.font.size = Pt(10)
    except Exception:
        pass

    # ------------------------------------------------------------
    # 翻訳結果（番号付き）
    # ------------------------------------------------------------
    doc.add_heading("翻訳結果（番号付き）", level=1)

    para2 = doc.add_paragraph()

    r2 = para2.add_run(
        translated_text_numbered
        if translated_text_numbered
        else "(なし)"
    )

    try:
        r2.font.name = "Courier New"
        r2.font.size = Pt(10)
    except Exception:
        pass

    # ------------------------------------------------------------
    # 翻訳結果
    # ------------------------------------------------------------
    doc.add_heading("翻訳結果", level=1)

    para3 = doc.add_paragraph()

    r3 = para3.add_run(
        translated_text
        if translated_text
        else "(なし)"
    )

    try:
        r3.font.name = "Courier New"
        r3.font.size = Pt(10)
    except Exception:
        pass

    # ------------------------------------------------------------
    # 要約
    # ------------------------------------------------------------
    if str(summary_text or "").strip():

        doc.add_heading("要約", level=1)

        para4 = doc.add_paragraph()

        r4 = para4.add_run(summary_text)

        try:
            r4.font.name = "Courier New"
            r4.font.size = Pt(10)
        except Exception:
            pass


    # ------------------------------------------------------------
    # bytes
    # ------------------------------------------------------------
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue(), ".docx"


# ============================================================
# PDF 出力
# ============================================================
def build_translation_pdf_bytes(
    *,
    original_numbered_preview: str,
    translated_text_numbered: str,
    translated_text: str,
    summary_text: str,
    model: str,
    direction: str,
    extra_prompt: str,
    src_name: str,
    usage_summary: Dict[str, int] | None = None,
) -> bytes:

    # ------------------------------------------------------------
    # import
    # ------------------------------------------------------------
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.lib.units import mm
    except Exception:
        return b""

    # ------------------------------------------------------------
    # font
    # ------------------------------------------------------------
    font_name = None
    for fname in ("HeiseiMin-W3", "HeiseiKakuGo-W5"):
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(fname))
            font_name = fname
            break
        except Exception:
            continue

    if not font_name:
        return b""

    # ------------------------------------------------------------
    # document
    # ------------------------------------------------------------
    buf = io.BytesIO()
    pagesize = A4
    margin = 18 * mm

    doc = SimpleDocTemplate(
        buf,
        pagesize=pagesize,
        leftMargin=margin,
        rightMargin=margin,
        topMargin=20 * mm,
        bottomMargin=18 * mm,
        title="翻訳結果レポート",
    )

    # ------------------------------------------------------------
    # styles
    # ------------------------------------------------------------
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleJP", fontName=font_name, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Meta", fontName=font_name, fontSize=10, leading=13, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle(name="H1", fontName=font_name, fontSize=13, leading=16, spaceBefore=6, spaceAfter=6))
    styles.add(ParagraphStyle(name="BodyJP", fontName=font_name, fontSize=9.5, leading=13, wordWrap="CJK"))

    text_width = pagesize[0] - doc.leftMargin - doc.rightMargin

    # ------------------------------------------------------------
    # page number
    # ------------------------------------------------------------
    def _page_number(canvas, doc_) -> None:
        canvas.setFont(font_name, 9)
        canvas.drawRightString(pagesize[0] - doc_.rightMargin, 12 * mm, f"{doc_.page}")

    # ------------------------------------------------------------
    # table helper
    # ------------------------------------------------------------
    def _make_text_table(text: str) -> Table:
        rows = []

        for raw in (text or "").splitlines() or ["(空)"]:
            s = html.escape(raw.replace("\t", "    "))
            s = s.replace("  ", "&nbsp;&nbsp;")
            p = Paragraph(s if s else "&nbsp;", styles["BodyJP"])
            rows.append([p])

        t = Table(rows, colWidths=[text_width - 2], splitByRow=1)
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    # ------------------------------------------------------------
    # story
    # ------------------------------------------------------------
    story = []
    story.append(Paragraph("翻訳結果レポート", styles["TitleJP"]))

    meta_lines = [
        f"生成日時：{_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        f"モデル：{model}",
        f"翻訳方向：{direction}",
        f"入力ファイル：{src_name or '-'}",
    ]

    if (extra_prompt or "").strip():
        meta_lines.insert(3, f"追加プロンプト：{extra_prompt.strip()}")

    if usage_summary:
        meta_lines.append(f"input_tokens：{usage_summary.get('input_tokens', '')}")
        meta_lines.append(f"output_tokens：{usage_summary.get('output_tokens', '')}")
        meta_lines.append(f"total_tokens：{usage_summary.get('total_tokens', '')}")

    for ln in meta_lines:
        story.append(Paragraph(html.escape(ln), styles["Meta"]))

    story.append(Paragraph("原文（番号つき）", styles["H1"]))
    story.append(_make_text_table(original_numbered_preview or "(空)"))
    story.append(Spacer(1, 6 * mm))

    # story.append(Paragraph("翻訳結果", styles["H1"]))
    # story.append(_make_text_table(translated_text or "(なし)"))

    # ------------------------------------------------------------
    # 翻訳結果（番号付き）
    # ------------------------------------------------------------
    story.append(
        Paragraph(
            "翻訳結果（番号付き）",
            styles["H1"],
        )
    )

    story.append(
        _make_text_table(
            translated_text_numbered or "(なし)"
        )
    )

    story.append(
        Spacer(1, 6 * mm)
    )

    # ------------------------------------------------------------
    # 翻訳結果
    # ------------------------------------------------------------
    story.append(
        Paragraph(
            "翻訳結果",
            styles["H1"],
        )
    )

    story.append(
        _make_text_table(
            translated_text or "(なし)"
        )
    )

    # ------------------------------------------------------------
    # 要約
    # ------------------------------------------------------------
    if str(summary_text or "").strip():

        story.append(
            Spacer(1, 6 * mm)
        )

        story.append(
            Paragraph(
                "要約",
                styles["H1"],
            )
        )

        story.append(
            _make_text_table(summary_text)
        )

    doc.build(story, onFirstPage=_page_number, onLaterPages=_page_number)

    pdf_bytes = buf.getvalue()
    buf.close()
    return pdf_bytes