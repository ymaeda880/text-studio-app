# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/figuretable/builder.py
# ============================================================
# figureTable Word builder
#
# - FigureTableBlock を Word の画像表として出力する
# - 図番号・キャプションもここで出力する
# - width / col_ratio により表幅・列幅を制御する
# ============================================================

from __future__ import annotations

# ============================================================
# imports
# ============================================================
from pathlib import Path

from docx.document import Document as DocumentObject
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from lib.wordtex.blocks import FigureTableBlock, FigureTableItem
from lib.wordtex.settings import WordTexSettings
from lib.wordtex.figuretable.constants import get_figure_table_bg_hex

from lib.image_embed.image_utils import calc_docx_image_width_inches
from lib.image_embed.inbox_image_resolver import resolve_inbox_image_path_by_filename


# ============================================================
# キャプション文字列作成
# ============================================================
def build_figure_caption_text(
    *,
    settings: WordTexSettings,
    block: FigureTableBlock,
) -> tuple[str, str]:
    """
    FigureTableBlock 用の図題文字列を作る。
    """
    fig = settings.figure

    figure_number_text = fig.format_number()
    title_text = str(block.caption or block.title or "").strip()
    caption_text = fig.format_caption(title_text)

    return caption_text, figure_number_text


# ============================================================
# キャプション出力
# ============================================================
def add_figure_caption(
    *,
    doc: DocumentObject,
    caption_text: str,
) -> None:
    text = str(caption_text or "").strip()

    if not text:
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run(text)
    run.bold = True


# ============================================================
# itemを行単位に分割
# ============================================================
def chunk_items(
    *,
    items: list[FigureTableItem],
    columns: int,
) -> list[list[FigureTableItem]]:
    if columns <= 0:
        columns = 1

    return [
        items[i : i + columns]
        for i in range(0, len(items), columns)
    ]


# ============================================================
# セルテキスト設定
# ============================================================
def set_title_cell_text(
    *,
    cell,
    item: FigureTableItem,
) -> None:
    title = str(item.title or item.file or "").strip()
    desp = str(item.desp or "").strip()

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if title:
        r = p.add_run(title)
        r.bold = True

    if desp:
        if title:
            p.add_run().add_break()
        p.add_run(desp)


# ============================================================
# 画像セル設定
# ============================================================
def set_image_cell_content(
    *,
    cell,
    inbox_root: Path,
    sub: str,
    settings: WordTexSettings,
    item: FigureTableItem,
    image_width_cm: float,
) -> None:
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fig_table_path = str(
        getattr(settings, "fig_table_path", "") or ""
    ).strip()

    if not fig_table_path or fig_table_path.lower() == "inbox":
        image_path = resolve_inbox_image_path_by_filename(
            inbox_root=inbox_root,
            sub=sub,
            file_name=item.file,
        )
    else:
        image_path = Path(fig_table_path) / str(item.file)

    if image_path is None or not image_path.exists():
        p.add_run(f"画像が見つかりません: {item.file}")
        return

    try:
        # ------------------------------------------------------------
        # 画像がセル上罫線に接触すると，
        # タイトル行と画像行の境界線が画像の背後に隠れて見える。
        # 画像の直前に小さな改行を入れて，罫線と画像の間に余白を作る。
        # ------------------------------------------------------------
        p.add_run().add_break()

        run = p.add_run()
        run.add_picture(
            str(image_path),
            width=Cm(float(image_width_cm)),
        )

        # ------------------------------------------------------------
        # 画像と下罫線が接触しないよう，
        # 画像の直後に改行を追加して下余白を確保する。
        # ------------------------------------------------------------
        p.add_run().add_break()

    except Exception as e:
        p.add_run(f"画像を貼り込めません: {item.file} / {e}")

# ============================================================
# セル背景色設定
# ============================================================
def set_cell_shading(
    cell,
    fill_hex: str,
) -> None:
    """
    Word表セルの背景色を設定する。
    fill_hex は FFFFFF のような #なし16進カラー。
    """
    fill = str(fill_hex or "").strip().replace("#", "")

    if not fill:
        return

    tc_pr = cell._tc.get_or_add_tcPr()

    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)

    shd.set(qn("w:fill"), fill)

# ============================================================
# セル罫線設定
# ============================================================
def set_cell_border(
    cell,
    *,
    size: int = 12,
    color: str = "808080",
) -> None:
    """
    セル四辺の罫線を設定する。

    size:
        Word の罫線サイズ（8=0.5pt, 12=0.75pt, 16=1pt程度）
    """

    tc_pr = cell._tc.get_or_add_tcPr()

    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)

    for edge in ("top", "left", "bottom", "right"):

        border = tc_borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            tc_borders.append(border)

        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:color"), color)

# ============================================================
# 行セルの背景色設定
# ============================================================
def shade_row_cells(
    *,
    row,
    fill_hex: str,
) -> None:
    for cell in row.cells:
        set_cell_shading(
            cell,
            fill_hex,
        )

        set_cell_border(
            cell,
            size=16,      # 1pt程度
            color="808080",
        )


# ============================================================
# 空セル処理
# ============================================================
def clear_unused_cell(
    *,
    cell,
    fill_hex: str,
) -> None:
    set_cell_shading(
        cell,
        fill_hex,
    )
    cell.text = ""


# ============================================================
# col_ratio 解析
# ============================================================
def parse_col_ratios(
    *,
    col_ratio: str,
    columns: int,
) -> list[float]:
    ratio_text = str(col_ratio or "").strip()

    if ratio_text:
        try:
            ratios = [
                float(v.strip())
                for v in ratio_text.split(",")
                if v.strip()
            ]
        except Exception:
            ratios = []
    else:
        ratios = []

    if len(ratios) != columns:
        ratios = [1.0] * columns

    if sum(ratios) <= 0:
        ratios = [1.0] * columns

    return ratios


# ============================================================
# 表幅(cm)計算
# ============================================================
def calc_table_width_cm(
    *,
    doc: DocumentObject,
    width_text: str,
) -> float | None:
    text = str(width_text or "100%").strip()

    if text.endswith("cm"):
        try:
            return float(text[:-2])
        except Exception:
            return None

    if text.endswith("%"):
        try:
            percent = float(text[:-1]) / 100.0

            section = doc.sections[-1]
            page_width_cm = section.page_width.cm
            left_margin_cm = section.left_margin.cm
            right_margin_cm = section.right_margin.cm

            printable_width_cm = (
                page_width_cm
                - left_margin_cm
                - right_margin_cm
            )

            return printable_width_cm * percent
        except Exception:
            return None

    return None


# ============================================================
# 列幅(cm)計算
# ============================================================
def calc_column_widths_cm(
    *,
    doc: DocumentObject,
    block: FigureTableBlock,
    columns: int,
    ratios: list[float],
) -> list[float]:
    ratio_sum = sum(ratios)

    total_cm = calc_table_width_cm(
        doc=doc,
        width_text=str(block.width or "100%"),
    )

    if total_cm is not None:
        return [
            total_cm * r / ratio_sum
            for r in ratios
        ]

    default_image_width_cm = (
        calc_docx_image_width_inches(
            mode="自動設定",
            manual_width_cm=7.0,
            columns=columns,
        )
        * 2.54
    )

    return [
        default_image_width_cm
        for _ in range(columns)
    ]


# ============================================================
# twips変換
# ============================================================
def cm_to_twips(value_cm: float) -> int:
    return int(Cm(float(value_cm)).twips)


# ============================================================
# table幅・列グリッド設定
# ============================================================
def apply_table_grid_widths(
    *,
    table,
    column_widths_cm: list[float],
) -> None:
    """
    Word表の幅を tblW と tblGrid で固定する。

    cell.width やセル余白はいじらず，
    Wordに列グリッドを明示する。
    """
    try:
        table.autofit = False
    except Exception:
        pass

    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    total_cm = sum(float(w) for w in column_widths_cm)

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)

    tbl_w.set(qn("w:w"), str(cm_to_twips(total_cm)))
    tbl_w.set(qn("w:type"), "dxa")

    old_tbl_grid = tbl.find(qn("w:tblGrid"))
    if old_tbl_grid is not None:
        tbl.remove(old_tbl_grid)

    tbl_grid = OxmlElement("w:tblGrid")
    for width_cm in column_widths_cm:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(cm_to_twips(width_cm)))
        tbl_grid.append(grid_col)

    tbl.insert(0, tbl_grid)


# ============================================================
# 画像幅(cm)計算
# ============================================================
def calc_image_widths_cm(
    *,
    column_widths_cm: list[float],
) -> list[float]:
    """
    画像幅は列幅より少し小さくする。

    セル余白をXMLで触るとWord側の再計算で崩れやすいため，
    ここでは画像幅だけを控えめにする。
    """
    image_widths: list[float] = []

    for width_cm in column_widths_cm:
        width = float(width_cm)

        if width >= 8.0:
            image_widths.append(width - 0.7)
        elif width >= 5.0:
            image_widths.append(width - 0.55)
        elif width >= 3.0:
            image_widths.append(width - 0.4)
        else:
            image_widths.append(max(1.0, width * 0.82))

    return image_widths


# ============================================================
# figureTable本体出力
# ============================================================
def add_figure_table_body(
    *,
    doc: DocumentObject,
    inbox_root: Path,
    sub: str,
    settings: WordTexSettings,
    block: FigureTableBlock,
) -> None:
    columns = int(block.row or 2)
    fill_hex = get_figure_table_bg_hex(block.bc)

    ratios = parse_col_ratios(
        col_ratio=block.col_ratio,
        columns=columns,
    )

    column_widths_cm = calc_column_widths_cm(
        doc=doc,
        block=block,
        columns=columns,
        ratios=ratios,
    )

    image_widths_cm = calc_image_widths_cm(
        column_widths_cm=column_widths_cm,
    )

    item_rows = chunk_items(
        items=block.items,
        columns=columns,
    )

    if not item_rows:
        doc.add_paragraph("figureTable内に画像項目がありません。")
        return

    table = doc.add_table(
        rows=0,
        cols=columns,
    )

    table.style = "Table Grid"

    # ------------------------------------------------------------
    # figureTable本体を中央寄せにする。
    #
    # captionは通常段落として中央寄せされるため，
    # 表本体が左寄せのままだと caption だけ右にずれて見える。
    # 表本体も中央寄せにして，caption と位置基準を揃える。
    # ------------------------------------------------------------
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    apply_table_grid_widths(
        table=table,
        column_widths_cm=column_widths_cm,
    )

    for item_row in item_rows:
        # --------------------------------------------------------
        # タイトル・説明行
        # --------------------------------------------------------
        title_row = table.add_row()
        shade_row_cells(
            row=title_row,
            fill_hex=fill_hex,
        )

        for col_idx in range(columns):
            cell = title_row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if col_idx < len(item_row):
                set_title_cell_text(
                    cell=cell,
                    item=item_row[col_idx],
                )
            else:
                clear_unused_cell(
                    cell=cell,
                    fill_hex=fill_hex,
                )

        # --------------------------------------------------------
        # 画像行
        # --------------------------------------------------------
        image_row = table.add_row()
        shade_row_cells(
            row=image_row,
            fill_hex=fill_hex,
        )

        for col_idx in range(columns):
            cell = image_row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            if col_idx < len(item_row):
                set_image_cell_content(
                    cell=cell,
                    inbox_root=inbox_root,
                    sub=sub,
                    settings=settings,
                    item=item_row[col_idx],
                    image_width_cm=image_widths_cm[col_idx],
                )
            else:
                clear_unused_cell(
                    cell=cell,
                    fill_hex=fill_hex,
                )


# ============================================================
# figureTable出力
# ============================================================
def add_figure_table_block(
    *,
    doc: DocumentObject,
    inbox_root: Path,
    sub: str,
    settings: WordTexSettings,
    block: FigureTableBlock,
) -> None:
    """
    FigureTableBlock を Word に出力する。
    """
    caption_text = ""
    figure_number_text = ""

    if block.numbering:
        caption_text, figure_number_text = build_figure_caption_text(
            settings=settings,
            block=block,
        )

        if block.label:
            settings.register_label(
                label=block.label,
                value=settings.figure.format_number_core(),
            )

    # ------------------------------------------------------------
    # caption position = top
    # ------------------------------------------------------------
    if block.numbering and settings.figure.caption_position == "top":
        add_figure_caption(
            doc=doc,
            caption_text=caption_text,
        )

    # ------------------------------------------------------------
    # figureTable本体
    # ------------------------------------------------------------
    add_figure_table_body(
        doc=doc,
        inbox_root=inbox_root,
        sub=sub,
        settings=settings,
        block=block,
    )

    # ------------------------------------------------------------
    # caption position = bottom
    # ------------------------------------------------------------
    if block.numbering and settings.figure.caption_position == "bottom":
        add_figure_caption(
            doc=doc,
            caption_text=caption_text,
        )

    # ------------------------------------------------------------
    # 図番号を進める
    # ------------------------------------------------------------
    if block.numbering:
        settings.figure.increment()

    doc.add_paragraph()