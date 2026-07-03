# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/figure/builder.py

from __future__ import annotations

from pathlib import Path

from docx.document import Document as DocumentObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from lib.wordtex.blocks import FigureBlock
from lib.wordtex.settings import WordTexSettings
from lib.image_embed.inbox_image_resolver import (
    resolve_inbox_image_path_by_filename,
)


def calc_figure_width_cm(
    *,
    doc: DocumentObject,
    width_text: str,
) -> float:
    text = str(width_text or "100%").strip()

    if text.endswith("cm"):
        try:
            return float(text[:-2])
        except Exception:
            return 12.0

    if text.endswith("%"):
        try:
            percent = float(text[:-1]) / 100.0
            section = doc.sections[-1]

            printable_width_cm = (
                section.page_width.cm
                - section.left_margin.cm
                - section.right_margin.cm
            )

            return printable_width_cm * percent
        except Exception:
            return 12.0

    return 12.0


def calc_figure_left_indent_cm(
    *,
    doc: DocumentObject,
    width_cm: float,
) -> float:
    """
    中央配置した画像の左端位置を計算する。

    note の左端を画像の左端に揃えるために使う。
    """
    section = doc.sections[-1]

    printable_width_cm = (
        section.page_width.cm
        - section.left_margin.cm
        - section.right_margin.cm
    )

    indent_cm = (printable_width_cm - float(width_cm)) / 2.0

    return max(0.0, indent_cm)


def build_figure_caption_text(
    *,
    settings: WordTexSettings,
    block: FigureBlock,
) -> str:
    """
    図キャプション文字列を作る。

    caption / title が無い場合は，
    図番号を出さない。
    """
    title_text = str(block.caption or block.title or "").strip()

    if not title_text:
        return ""

    return settings.figure.format_caption(title_text)


def add_figure_caption(
    *,
    doc: DocumentObject,
    caption_text: str,
    position: str,
) -> None:
    """
    図キャプションを追加する。

    position:
        top    : 図の上キャプション
        bottom : 図の下キャプション
    """
    text = str(caption_text or "").strip()

    if not text:
        return

    pos = str(position or "bottom").strip().lower()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------
    # 図キャプションの前後余白
    # ------------------------------------------------------------
    if pos == "top":
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    else:
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)

    p.paragraph_format.line_spacing = 1.0

    run = p.add_run(text)
    run.bold = True


def add_figure_image(
    *,
    doc: DocumentObject,
    inbox_root: Path,
    sub: str,
    settings: WordTexSettings,
    block: FigureBlock,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    """
    図画像を追加する。
    """
    if block.item is None:
        p = doc.add_paragraph("figure に画像 item がありません。")
        p.paragraph_format.space_before = Pt(float(space_before_pt))
        p.paragraph_format.space_after = Pt(float(space_after_pt))
        return

    fig_path = str(getattr(settings, "fig_path", "") or "").strip()

    if not fig_path or fig_path.lower() == "inbox":
        image_path = resolve_inbox_image_path_by_filename(
            inbox_root=inbox_root,
            sub=sub,
            file_name=block.item.file,
        )
    else:
        image_path = Path(fig_path) / str(block.item.file)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ------------------------------------------------------------
    # 図画像段落の前後余白
    # ------------------------------------------------------------
    p.paragraph_format.space_before = Pt(float(space_before_pt))
    p.paragraph_format.space_after = Pt(float(space_after_pt))
    p.paragraph_format.line_spacing = 1.0

    if image_path is None or not image_path.exists():
        p.add_run(f"画像が見つかりません: {block.item.file}")
        return

    width_cm = calc_figure_width_cm(
        doc=doc,
        width_text=block.width,
    )

    run = p.add_run()
    run.add_picture(
        str(image_path),
        width=Cm(float(width_cm)),
    )


def add_figure_note(
    *,
    doc: DocumentObject,
    note: str,
    left_indent_cm: float,
) -> None:
    """
    図 note を追加する。
    """
    text = str(note or "").strip()

    if not text:
        return

    lines = [
        line.strip()
        for line in text.splitlines()
        if line.strip()
    ]

    if not lines:
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # ------------------------------------------------------------
    # note の左端を画像の左端に揃える。
    # 「注）」は付けず，note本文だけを小さい文字で出す。
    # ------------------------------------------------------------
    p.paragraph_format.left_indent = Cm(float(left_indent_cm))

    # ------------------------------------------------------------
    # 図 note の前後余白
    # - space_before : 図と note の間
    # - space_after  : note と次本文の間
    # ------------------------------------------------------------
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.0

    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()

        run = p.add_run(line)
        run.font.size = Pt(9)


def add_figure_block(
    *,
    doc: DocumentObject,
    inbox_root: Path,
    sub: str,
    settings: WordTexSettings,
    block: FigureBlock,
) -> None:
    caption_text = ""
    has_caption = bool(str(block.caption or block.title or "").strip())
    has_note = bool(str(block.note or "").strip())

    caption_position = str(
        getattr(settings.figure, "caption_position", "bottom") or "bottom"
    ).strip().lower()

    # ------------------------------------------------------------
    # caption がある図だけを番号付き図として扱う。
    # caption が無い場合は，
    # - 図キャプションを出さない
    # - label を登録しない
    # - 図番号を進めない
    # ------------------------------------------------------------
    if block.numbering and has_caption:
        caption_text = build_figure_caption_text(
            settings=settings,
            block=block,
        )

        if block.label:
            settings.register_label(
                label=block.label,
                value=settings.figure.format_number_core(),
            )

    if caption_text and caption_position == "top":
        add_figure_caption(
            doc=doc,
            caption_text=caption_text,
            position="top",
        )

    width_cm = calc_figure_width_cm(
        doc=doc,
        width_text=block.width,
    )

    # ------------------------------------------------------------
    # 図画像の前後余白
    # - 上キャプションが無い場合は，画像の上に余白を作る。
    # - note または下キャプションが無い場合は，画像の下に余白を作る。
    # ------------------------------------------------------------
    image_space_before = 0 if caption_text and caption_position == "top" else 6

    if has_note:
        image_space_after = 0
    elif caption_text and caption_position == "bottom":
        image_space_after = 0
    else:
        image_space_after = 6

    add_figure_image(
        doc=doc,
        inbox_root=inbox_root,
        sub=sub,
        settings=settings,
        block=block,
        space_before_pt=image_space_before,
        space_after_pt=image_space_after,
    )

    left_indent_cm = calc_figure_left_indent_cm(
        doc=doc,
        width_cm=width_cm,
    )

    add_figure_note(
        doc=doc,
        note=block.note,
        left_indent_cm=left_indent_cm,
    )

    if caption_text and caption_position == "bottom":
        add_figure_caption(
            doc=doc,
            caption_text=caption_text,
            position="bottom",
        )

    if block.numbering and has_caption:
        settings.figure.increment()