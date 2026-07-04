# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/builders/paragraph_heading_builder.py
# ============================================================
# wordTex paragraph heading builder
#
# - \paragraph{} を小見出しとして出力する
# - 本文とは別段落にする
# ============================================================

from __future__ import annotations

from docx.document import Document as DocumentObject
from docx.shared import Pt

from lib.wordtex.blocks import ParagraphHeadingBlock
from lib.wordtex.settings import WordTexSettings


def add_paragraph_heading_block(
    *,
    doc: DocumentObject,
    settings: WordTexSettings,
    block: ParagraphHeadingBlock,
) -> None:
    title = str(block.title or "").strip()

    if not title:
        return

    title = settings.resolve_refs(title)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run(title)
    run.bold = True