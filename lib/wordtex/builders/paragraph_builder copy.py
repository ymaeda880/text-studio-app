# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/builders/paragraph_builder.py
# ============================================================
# wordTex 段落 builder
# ============================================================

from __future__ import annotations

# ============================================================
# imports
# ============================================================
from docx.document import Document as DocumentObject
from docx.shared import Pt

from lib.wordtex.blocks import ParagraphBlock
from lib.wordtex.settings import WordTexSettings


# ============================================================
# 段落出力
# ============================================================
def add_paragraph_block(
    *,
    doc: DocumentObject,
    settings: WordTexSettings,
    block: ParagraphBlock,
) -> None:
    """
    ParagraphBlock を Word の通常段落として出力する。

    本文中の \\ref{...} は，登録済みlabel番号へ置換する。
    """
    text = str(block.text or "").strip()

    if not text:
        return

    # ------------------------------------------------------------
    # 本文中の \ref{...} を置換する
    # ------------------------------------------------------------
    text = settings.resolve_refs(text)

    # ------------------------------------------------------------
    # 空行区切りで複数段落に分ける
    # ------------------------------------------------------------
    paragraphs = [
        p.strip()
        for p in text.split("\n\n")
        if p.strip()
    ]

    for para_text in paragraphs:
        # --------------------------------------------------------
        # 段落内の単一改行は Word 上でも改行として残す
        # --------------------------------------------------------
        lines = para_text.split("\n")

        p = doc.add_paragraph()

        # --------------------------------------------------------
        # 通常本文
        # - 1行目を字下げする
        # - 段落前後の余白を入れない
        # --------------------------------------------------------
        p.paragraph_format.first_line_indent = Pt(10.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        for i, line in enumerate(lines):
            if i > 0:
                p.add_run().add_break()

            p.add_run(line)