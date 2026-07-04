# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/builders/paragraph_builder.py
# ============================================================
# wordTex 段落 builder
#
# - ParagraphBlock を Word の通常段落として出力する
# - 段落内インライン装飾に対応する
# ============================================================

from __future__ import annotations

# ============================================================
# imports
# ============================================================
from dataclasses import dataclass
import re

from docx.document import Document as DocumentObject
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

from lib.wordtex.blocks import ParagraphBlock
from lib.wordtex.settings import WordTexSettings
from lib.wordtex.font.presets import get_font_preset

from lib.wordtex.colors import COLOR_MAP


# ============================================================
# inline style
# ============================================================
@dataclass
class InlineStyle:
    size: str = ""
    color: str = ""
    bold: bool = False
    italic: bool = False
    underline: bool = False


@dataclass
class InlineRun:
    text: str
    style: InlineStyle


def parse_hex_color(value: str) -> RGBColor | None:
    """
    #RRGGBB を RGBColor に変換する。
    """
    text = str(value or "").strip()

    if not re.match(r"^#[0-9a-fA-F]{6}$", text):
        return None

    r = int(text[1:3], 16)
    g = int(text[3:5], 16)
    b = int(text[5:7], 16)

    return RGBColor(r, g, b)


# ============================================================
# style apply
# ============================================================
def apply_inline_style(
    *,
    run,
    style: InlineStyle,
) -> None:
    """
    InlineStyle を Word run に反映する。
    """
    if style.size == "small":
        run.font.size = Pt(9)

    elif style.size == "large":
        run.font.size = Pt(14)

    elif style.size == "Large":
        run.font.size = Pt(18)

    if style.color:
        if style.color in COLOR_MAP:
            run.font.color.rgb = COLOR_MAP[style.color]
        else:
            rgb = parse_hex_color(style.color)
            if rgb is not None:
                run.font.color.rgb = rgb

    if style.bold:
        run.bold = True

    if style.italic:
        run.italic = True

    if style.underline:
        run.underline = True


def parse_style_commands(command_text: str) -> InlineStyle:
    """
    {\\large\\bold text} の \\large\\bold 部分を解析する。
    """
    style = InlineStyle()

    commands = [
        c.strip()
        for c in str(command_text or "").split("\\")
        if c.strip()
    ]

    for command in commands:
        if command == "small":
            style.size = "small"

        elif command == "large":
            style.size = "large"

        elif command == "Large":
            style.size = "Large"

        elif command == "bold":
            style.bold = True

        elif command == "italic":
            style.italic = True

        elif command == "underline":
            style.underline = True

        elif command in COLOR_MAP:
            style.color = command

        elif command.startswith("color="):
            style.color = command.split("=", 1)[1].strip()

    return style


def parse_inline_runs(text: str) -> list[InlineRun]:
    """
    段落内の簡易インライン装飾を解析する。

    対応:
        {\\large 大きい文字}
        {\\Large すごく大きい文字}
        {\\small 小さい文字}
        {\\bold 太字}
        {\\italic 斜体}
        {\\underline 下線}
        {\\red 赤文字}
        {\\blue 青文字}
        {\\green 緑文字}
        {\\gray 灰色文字}
        {\\color=#FF0000 赤文字}
        {\\large\\bold\\red 大きい赤太字}

    注意:
        入れ子は初期対応しない。
    """
    source = str(text or "")

    runs: list[InlineRun] = []

    i = 0
    plain_buf: list[str] = []

    while i < len(source):

        # --------------------------------------------------------
        # 装飾ブロック開始
        # --------------------------------------------------------
        if source[i] == "{" and i + 1 < len(source) and source[i + 1] == "\\":

            # ----------------------------------------------------
            # 直前までの通常文字を出す
            # ----------------------------------------------------
            if plain_buf:
                runs.append(
                    InlineRun(
                        text="".join(plain_buf),
                        style=InlineStyle(),
                    )
                )
                plain_buf.clear()

            end = source.find("}", i + 1)

            if end == -1:
                plain_buf.append(source[i])
                i += 1
                continue

            inner = source[i + 1:end]

            split_pos = -1

            for idx, ch in enumerate(inner):
                if ch.isspace():
                    split_pos = idx
                    break

            if split_pos == -1:
                plain_buf.append(source[i:end + 1])
                i = end + 1
                continue

            command_text = inner[:split_pos]
            body_text = inner[split_pos + 1:]

            style = parse_style_commands(command_text)

            runs.append(
                InlineRun(
                    text=body_text,
                    style=style,
                )
            )

            i = end + 1
            continue

        plain_buf.append(source[i])
        i += 1

    if plain_buf:
        runs.append(
            InlineRun(
                text="".join(plain_buf),
                style=InlineStyle(),
            )
        )

    return runs


def add_inline_runs(
    *,
    paragraph,
    text: str,
    settings: WordTexSettings,
) -> None:
    """
    段落へインライン装飾つき run を追加する。
    """
    preset = get_font_preset(settings.font)
    body_font = preset.body

    for inline_run in parse_inline_runs(text):
        run = paragraph.add_run(inline_run.text)

        # --------------------------------------------------------
        # 本文フォント
        # - ascii / hAnsi と eastAsia の両方を指定する
        # - 日本語フォントが環境依存で変わるのを防ぐ
        # --------------------------------------------------------
        run.font.name = body_font
        run.font.size = Pt(settings.size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), body_font)

        apply_inline_style(
            run=run,
            style=inline_run.style,
        )

# ============================================================
# 段落出力
# ============================================================
def add_paragraph_block(
    *,
    doc: DocumentObject,
    settings: WordTexSettings,
    block: ParagraphBlock,
) -> None:
    """
    ParagraphBlock を Word の通常段落として出力する。

    本文中の \\ref{...} は，登録済みlabel番号へ置換する。
    """
    text = str(block.text or "").strip()

    if not text:
        return

    # ------------------------------------------------------------
    # 本文中の \ref{...} を置換する
    # ------------------------------------------------------------
    text = settings.resolve_refs(text)

    # ------------------------------------------------------------
    # 空行区切りで複数段落に分ける
    # ------------------------------------------------------------
    paragraphs = [
        p.strip()
        for p in text.split("\n\n")
        if p.strip()
    ]

    for para_text in paragraphs:
        # --------------------------------------------------------
        # 段落内の単一改行は Word 上でも改行として残す
        # --------------------------------------------------------
        lines = para_text.split("\n")

        p = None

        for line in lines:

            line = str(line)

            # ----------------------------------------------------
            # \newline だけの行
            # - 空段落を1つ作るだけにする
            # - add_break() は入れない
            # - 空段落 + 改行 になると空き過ぎるため
            # ----------------------------------------------------
            if line.strip() == r"\newline":

                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

                continue

            if p is None:
                p = doc.add_paragraph()

                # ------------------------------------------------
                # \noindent 指定段落
                # - 直前に \noindent がある段落だけ字下げしない
                # ------------------------------------------------
                if getattr(block, "noindent", False):
                    p.paragraph_format.first_line_indent = Pt(0)
                else:
                    p.paragraph_format.first_line_indent = Pt(10.5)
                    
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

            else:
                p.add_run().add_break()

            add_inline_runs(
                paragraph=p,
                text=line,
                settings=settings,
            )