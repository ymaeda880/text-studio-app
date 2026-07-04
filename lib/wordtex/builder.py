# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/builder.py
# ============================================================
# wordTex Word builder
#
# - ParsedWordTex を Word 文書へ変換する
# - このファイルはディスパッチャとして扱う
# - Word表などの細かい処理は専用 builder へ切り出す
# ============================================================

from __future__ import annotations

# ============================================================
# imports
# ============================================================
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from lib.wordtex.blocks import (
    ParsedWordTex,
    ParagraphBlock,
    HeadingBlock,
    ParagraphHeadingBlock,
    SetBlock,
    FigureTableBlock,
    FigureBlock,
    NewPageBlock,
    VSkipBlock,
    TitleBlock,
    AuthorBlock,
    DateBlock,
    TableOfContentsBlock,
    ItemizeBlock,
)
from lib.wordtex.settings import (
    WordTexSettings,
    apply_wordtex_setting,
)
from lib.wordtex.builders.paragraph_builder import add_paragraph_block
from lib.wordtex.builders.heading_builder import add_heading_block
from lib.wordtex.builders.paragraph_heading_builder import add_paragraph_heading_block
from lib.wordtex.figuretable.builder import add_figure_table_block
from lib.wordtex.figure.builder import add_figure_block
from lib.wordtex.itemize.builder import add_itemize_block
from lib.wordtex.table.builder import add_table_block
from lib.wordtex.builders.vskip_builder import add_vskip_block

# ============================================================
# setブロック反映
# ============================================================
def apply_set_block(
    *,
    settings: WordTexSettings,
    block: SetBlock,
) -> None:
    """
    SetBlock を WordTexSettings に反映する。
    """
    for key, value in block.values.items():
        apply_wordtex_setting(
            settings=settings,
            key=key,
            value=value,
        )

# ============================================================
# タイトルページ
# ============================================================
def add_title_page(
    *,
    doc,
    title_text: str,
    author_text: str,
    date_text: str,
) -> None:
    """
    タイトルページを作る。

    - title の改行を保持する
    - title の空行も保持する
    - タイトルページ後に新しいセクションを開始する
    """
    for _ in range(5):
        doc.add_paragraph("")

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_lines = str(title_text or "").split("\n")

    for idx, line in enumerate(title_lines):
        run = p_title.add_run(str(line))
        run.bold = True
        run.font.size = Pt(22)

        if idx < len(title_lines) - 1:
            run.add_break()

    doc.add_paragraph("")

    if date_text:
        p_date = doc.add_paragraph()
        p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_date = p_date.add_run(str(date_text))
        r_date.font.size = Pt(14)

    if author_text:
        p_author = doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER

        r_author = p_author.add_run(str(author_text))
        r_author.font.size = Pt(14)

    doc.add_section(WD_SECTION.NEW_PAGE)


# ============================================================
# ページ番号フィールド
# ============================================================
def add_page_number_to_footer(section) -> None:
    """
    指定セクションのフッター中央に PAGE フィールドを入れる。
    """
    footer = section.footer

    if not footer.paragraphs:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    text = OxmlElement("w:t")
    text.text = "1"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)



# ============================================================
# 目次
# ============================================================
def add_table_of_contents(
    *,
    doc,
) -> None:
    """
    Word の TOC フィールドを入れる。

    - \\section{}       → Heading 1
    - \\subsection{}    → Heading 2
    - \\subsubsection{} → Heading 3

    注意:
    Wordで開いた後に「目次の更新」が必要。
    """
    heading = doc.add_heading(
        "目次",
        level=1,
    )

    # ------------------------------------------------------------
    # 目次見出し文字色
    # - Word標準の Heading スタイルでは青色になるため，
    #   wordTexでは黒へ統一する。
    # ------------------------------------------------------------
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    p = doc.add_paragraph()

    # ------------------------------------------------------------
    # TOCフィールド開始
    # ------------------------------------------------------------
    run = p.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    # ------------------------------------------------------------
    # TOC命令
    # ------------------------------------------------------------
    run = p.add_run()

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)

    # ------------------------------------------------------------
    # TOCフィールド区切り
    # ------------------------------------------------------------
    run = p.add_run()

    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    run._r.append(fld_separate)

    # ------------------------------------------------------------
    # Word上で表示される仮文字
    # ------------------------------------------------------------
    p.add_run("目次を更新してください。")

    # ------------------------------------------------------------
    # TOCフィールド終了
    # ------------------------------------------------------------
    run = p.add_run()

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)

    doc.add_page_break()

# ============================================================
# 警告出力
# ============================================================
def add_warnings_to_doc(
    *,
    doc,
    settings: WordTexSettings,
) -> None:
    """
    wordTex処理中の警告をWord末尾に出力する。
    """
    if not settings.warnings:
        return

    doc.add_page_break()
    doc.add_heading("wordTex 警告", level=1)

    for warning in settings.warnings:
        doc.add_paragraph(str(warning))


# ============================================================
# label事前登録パス
# ============================================================
def collect_labels_first_pass(
    *,
    parsed: ParsedWordTex,
    settings: WordTexSettings,
) -> None:
    """
    前方参照に対応するため，Word出力前にlabelだけ登録する。

    注意:
    このパスではWordには出力しない。
    section番号とfigure番号だけを進めてlabel辞書を作る。
    """
    for block in parsed.blocks:

        if isinstance(block, HeadingBlock):
            number_text = settings.next_section_number(
                level=int(block.level),
            )

            label = str(block.label or "").strip()
            if label:
                settings.register_label(
                    label=label,
                    value=number_text,
                )

            continue

        if isinstance(block, SetBlock):
            apply_set_block(
                settings=settings,
                block=block,
            )
            continue

        if isinstance(block, FigureTableBlock):
            if block.numbering:
                fig_number = settings.figure.format_number_core()

                label = str(block.label or "").strip()
                if label:
                    settings.register_label(
                        label=label,
                        value=fig_number,
                    )

                settings.figure.increment()
            continue

        if isinstance(block, FigureBlock):
            if block.numbering:
                fig_number = settings.figure.format_number_core()

                label = str(block.label or "").strip()
                if label:
                    settings.register_label(
                        label=label,
                        value=fig_number,
                    )

                settings.figure.increment()
            continue



# ============================================================
# Word文書作成
# ============================================================
def build_wordtex_docx_bytes(
    *,
    parsed: ParsedWordTex,
    inbox_root: Path,
    sub: str,
) -> bytes:
    """
    解析済み wordTex から Word(.docx) bytes を作成する。

    対応ブロック:
    - ParagraphBlock
    - HeadingBlock
    - SetBlock
    - FigureTableBlock
    """
    doc = Document()

    # ------------------------------------------------------------
    # タイトル情報を先に取得する
    # - TitleBlock が存在する場合，最初のページをタイトルページにする
    # - AuthorBlock / DateBlock はタイトルページ内で使う
    # ------------------------------------------------------------
    title_text = ""
    author_text = ""
    date_text = ""

    for block in parsed.blocks:
        if isinstance(block, TitleBlock):
            title_text = block.text
            continue

        if isinstance(block, AuthorBlock):
            author_text = block.text
            continue

        if isinstance(block, DateBlock):
            date_text = block.text
            continue

    if title_text:
        add_title_page(
            doc=doc,
            title_text=title_text,
            author_text=author_text,
            date_text=date_text,
        )

        # --------------------------------------------------------
        # タイトルページの次のセクションからページ番号を入れる。
        # タイトルページにはページ番号を入れない。
        # --------------------------------------------------------
        add_page_number_to_footer(
            section=doc.sections[-1],
        )

    # ------------------------------------------------------------
    # 1パス目:
    # 前方参照に対応するため，labelだけ先に登録する
    # ------------------------------------------------------------
    label_settings = WordTexSettings()

    for warning in parsed.warnings:
        label_settings.add_warning(warning)

    collect_labels_first_pass(
        parsed=parsed,
        settings=label_settings,
    )

    # ------------------------------------------------------------
    # 2パス目:
    # 実際にWordへ出力する
    # 1パス目で作った labels を引き継ぐ
    # ------------------------------------------------------------
    settings = WordTexSettings()
    settings.labels = dict(label_settings.labels)

    for warning in parsed.warnings:
        settings.add_warning(warning)

    # ------------------------------------------------------------
    # ブロックを上から順に Word へ出力する
    # ------------------------------------------------------------
    for block in parsed.blocks:

        # --------------------------------------------------------
        # タイトルページ関連
        # - TitleBlock / AuthorBlock / DateBlock は
        #   文書冒頭のタイトルページ作成で使用済みなので，
        #   本文側には出力しない。
        # --------------------------------------------------------
        if isinstance(block, TitleBlock):
            continue

        if isinstance(block, AuthorBlock):
            continue

        if isinstance(block, DateBlock):
            continue

        # --------------------------------------------------------
        # 目次
        # - \tableofcontents の位置に Word の TOC フィールドを入れる
        # --------------------------------------------------------
        if isinstance(block, TableOfContentsBlock):
            add_table_of_contents(
                doc=doc,
            )
            continue


        # --------------------------------------------------------
        # 通常段落
        # --------------------------------------------------------
        if isinstance(block, ParagraphBlock):
            add_paragraph_block(
                doc=doc,
                settings=settings,
                block=block,
            )
            continue
        # --------------------------------------------------------
        # 見出し
        # --------------------------------------------------------
        if isinstance(block, HeadingBlock):
            add_heading_block(
                doc=doc,
                settings=settings,
                block=block,
            )
            continue

        # --------------------------------------------------------
        # paragraph小見出し
        # --------------------------------------------------------
        if isinstance(block, ParagraphHeadingBlock):
            add_paragraph_heading_block(
                doc=doc,
                settings=settings,
                block=block,
            )
            continue        

        # --------------------------------------------------------
        # set設定
        # - Wordには直接出力しない
        # - 後続ブロック用の内部状態だけ変更する
        # --------------------------------------------------------
        if isinstance(block, SetBlock):
            apply_set_block(
                settings=settings,
                block=block,
            )
            continue

        # --------------------------------------------------------
        # itemize
        # --------------------------------------------------------
        if isinstance(block, ItemizeBlock):
            add_itemize_block(
                doc=doc,
                settings=settings,
                block=block,
            )
            continue
        
        # --------------------------------------------------------
        # figureTable
        # --------------------------------------------------------
        if isinstance(block, FigureTableBlock):
            add_figure_table_block(
                doc=doc,
                inbox_root=Path(inbox_root),
                sub=str(sub),
                settings=settings,
                block=block,
            )
            continue

        # --------------------------------------------------------
        # figure
        # - 写真1枚専用
        # - 枠線なし
        # - noteあり
        # --------------------------------------------------------
        if isinstance(block, FigureBlock):
            add_figure_block(
                doc=doc,
                inbox_root=Path(inbox_root),
                sub=str(sub),
                settings=settings,
                block=block,
            )
            continue      

        # ------------------------------------------------------------
        # table
        # ------------------------------------------------------------
        if block.__class__.__name__ == "TableBlock":
            add_table_block(
                doc,
                block,
                settings,
                inbox_root=Path(inbox_root),
                sub=str(sub),
            )
            continue

        # --------------------------------------------------------
        # 縦方向スペース
        #
        # wordTex の
        #
        #     \vskip{1line}
        #     \vskip{0.5line}
        #     \vskip{12pt}
        #
        # に対応する。
        # --------------------------------------------------------
        if isinstance(block, VSkipBlock):
            add_vskip_block(
                doc=doc,
                settings=settings,
                block=block,
            )
            continue

        # --------------------------------------------------------
        # 改ページ
        #
        # wordTex の
        #
        #     /newpage
        #
        # に対応する。
        # --------------------------------------------------------
        if isinstance(block, NewPageBlock):
            doc.add_page_break()
            continue

        # --------------------------------------------------------
        # 未対応ブロック
        # --------------------------------------------------------
        settings.add_warning(
            f"未対応のブロックです: {type(block).__name__}"
        )

    # ------------------------------------------------------------
    # 警告があれば末尾に出す
    # ------------------------------------------------------------
    add_warnings_to_doc(
        doc=doc,
        settings=settings,
    )

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()