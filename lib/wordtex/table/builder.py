# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/table/builder.py
# ============================================================
# wordTex table builder
# ============================================================

from __future__ import annotations

from pathlib import Path
from openpyxl import load_workbook
import re

from docx.shared import Cm, Pt
from docx.enum.table import (
    WD_ROW_HEIGHT_RULE,
    WD_TABLE_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from lib.table.helpers import (
    _compute_col_widths_cm as compute_col_widths_cm,
    _compute_spans_markers as compute_spans_markers,
    _merge_docx_by_spans as merge_docx_by_spans,
    _apply_docx_col_widths as apply_docx_col_widths,
    _apply_table_borders_robust as apply_table_borders_robust,
    _apply_font_run as apply_font_run,
)
from lib.wordtex.table.style import resolve_table_preset
from lib.wordtex.table.inbox_excel_resolver import (
    resolve_inbox_excel_path_by_filename,
)

TOTAL_CM = 16.0

TABLE_IMAGE_EXTENSIONS = {
    ".jpg",
    ".jpeg",
    ".png",
    ".gif",
    ".bmp",
    ".tif",
    ".tiff",
    ".webp",
}

TABLE_IMAGE_TOKEN_RE = re.compile(
    r"^\s*<(?P<body>[^<>]+)>\s*$"
)


def parse_table_image_token(value: object) -> dict[str, str] | None:
    """
    table セル内の画像指定を解析する。

    対応例:
        <図1.png>
        <図1.png,width=2.5cm>
        <図1.png,width=2.5cm,align=center>

    ルール:
        - 先頭要素は画像ファイル名
        - 2番目以降は key=value
    """
    
    text = str(value or "").strip()

    if not text:
        return None

    # ------------------------------------------------------------
    # 画像セルに border 指定が付いている場合だけ，
    # 画像判定用の文字列から border 指定を外す。
    # rows 本体は変更しないので，後段の罫線上書きは効く。
    #
    # 例:
    #   <border:top=none><チゴユリ.png>
    #   → 画像判定では <チゴユリ.png> として扱う
    # ------------------------------------------------------------
    text = re.sub(
        r"\s*[<＜]\s*border\s*:[^<>＜＞]+[>＞]\s*",
        "",
        text,
    ).strip()

    m = TABLE_IMAGE_TOKEN_RE.match(text)
    if not m:
        return None

    body = str(m.group("body") or "").strip()
    if not body:
        return None

    parts = [p.strip() for p in body.split(",") if p.strip()]
    if not parts:
        return None

    file_name = parts[0]
    suffix = Path(file_name).suffix.lower()

    if suffix not in TABLE_IMAGE_EXTENSIONS:
        return None

    result: dict[str, str] = {
        "file": file_name,
    }

    for part in parts[1:]:
        if "=" not in part:
            result[part] = ""
            continue

        key, val = part.split("=", 1)
        result[key.strip().lower()] = val.strip()

    return result


def parse_cm_value(value: object) -> float | None:
    """
    2.5cm のような指定を cm 数値に変換する。
    """
    text = str(value or "").strip().lower()

    if not text:
        return None

    if text.endswith("cm"):
        text = text[:-2].strip()

    try:
        cm = float(text)
    except ValueError:
        return None

    if cm <= 0:
        return None

    return cm


def resolve_table_image_path(
    *,
    file_name: str,
    settings,
    inbox_root: Path | None = None,
    sub: str = "",
) -> Path | None:
    """
    table_fig_path から表セル用画像の実パスを解決する。

    table_fig_path="inbox" の場合は，
    inbox_root/sub 以下を再帰的に探す。
    """
    file_name = str(file_name or "").strip()
    if not file_name:
        return None

    table_fig_path = str(
        getattr(settings, "table_fig_path", "") or ""
    ).strip()

    if not table_fig_path:
        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"tableセル画像がありますが table_fig_path が未設定です: {file_name}"
            )
        return None

    if table_fig_path.lower() == "inbox":
        if inbox_root is None or not sub:
            if hasattr(settings, "add_warning"):
                settings.add_warning(
                    'table_fig_path="inbox" が指定されていますが inbox_root または sub が未設定です。'
                )
            return None

        search_root = Path(inbox_root) / str(sub)

        if not search_root.exists():
            if hasattr(settings, "add_warning"):
                settings.add_warning(
                    f"Inbox画像検索フォルダが見つかりません: {search_root}"
                )
            return None

        matches = list(search_root.rglob(file_name))

        if not matches:
            if hasattr(settings, "add_warning"):
                settings.add_warning(
                    f"Inbox内のtableセル画像が見つかりません: {file_name}"
                )
            return None

        return matches[0]

    image_path = Path(table_fig_path) / file_name

    if not image_path.exists():
        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"tableセル画像ファイルが見つかりません: {image_path}"
            )
        return None

    return image_path


def normalize_table_rows_for_cell_text(
    rows: list[list[str]],
) -> list[list[str]]:
    """
    tableセル内の簡易記法をWord出力用に正規化する。

    対応:
        <改行> → セル内改行
    """
    normalized: list[list[str]] = []

    for row in rows:
        new_row: list[str] = []

        for cell in row:

            text = str(cell or "")
            text = text.replace("<改行>", "\n")
            new_row.append(text)

        normalized.append(new_row)

    return normalized

def normalize_table_row_lengths(
    rows: list[list[str]],
) -> list[list[str]]:
    """
    table rows の列数を最大列数に揃える。

    Excel読み込みでは右端の空セルを落としているため，
    行によって列数が違う場合がある。
    compute_spans_markers() は矩形表を前提にするため，
    不足セルを "" で補う。
    """
    if not rows:
        return rows

    max_cols = max(len(row) for row in rows)

    if max_cols <= 0:
        return rows

    normalized: list[list[str]] = []

    for row in rows:
        new_row = list(row)

        while len(new_row) < max_cols:
            new_row.append("")

        normalized.append(new_row)

    return normalized

def clear_docx_cell(cell) -> None:
    """
    セル内の既存段落を空にする。
    """
    cell.text = ""


def add_table_image_to_cell(
    *,
    cell,
    image_path: Path,
    width_cm: float,
    align: str,
) -> None:
    """
    Word表セルに画像を挿入する。

    画像は指定幅で挿入し，縦横比は python-docx に任せる。
    """
    clear_docx_cell(cell)

    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]

    align_text = str(align or "center").strip().lower()

    if align_text == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align_text == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0

    run = p.add_run()
    run.add_picture(
        str(image_path),
        width=Cm(float(width_cm)),
    )


def apply_table_cell_images(
    *,
    table,
    rows: list[list[str]],
    widths_cm: list[float],
    settings,
    inbox_root: Path | None = None,
    sub: str = "",
) -> None:
    """
    rows内に画像指定セルがあれば，該当セルを画像に置き換える。

    対応:
        <図1.png>
        <図1.png,width=2.5cm,align=center>
    """
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            info = parse_table_image_token(value)
            if info is None:
                continue

            image_path = resolve_table_image_path(
                file_name=info.get("file", ""),
                settings=settings,
                inbox_root=inbox_root,
                sub=sub,
            )

            if image_path is None:
                continue

            if c_idx < len(widths_cm):
                default_width_cm = max(0.5, float(widths_cm[c_idx]) - 0.2)
            else:
                default_width_cm = 3.0

            width_cm = parse_cm_value(info.get("width"))
            if width_cm is None:
                width_cm = default_width_cm

            align = str(info.get("align") or "center").strip().lower()

            if align not in {"left", "center", "right"}:
                if hasattr(settings, "add_warning"):
                    settings.add_warning(
                        f"tableセル画像 align が不正です: {align}"
                    )
                align = "center"

            try:
                cell = table.cell(r_idx, c_idx)
            except Exception:
                continue

            add_table_image_to_cell(
                cell=cell,
                image_path=image_path,
                width_cm=width_cm,
                align=align,
            )

def load_table_rows_from_excel(
    *,
    block,
    settings,
    inbox_root: Path | None = None,
    sub: str = "",
) -> tuple[list[list[str]], str]:
    """
    table の src 指定から Excel 表を読み込む。

    A列に <note> がある場合:
    - その行より上を表本体にする
    - その行より下を note として返す
    """

    # raise RuntimeError(
    # f"load_table_rows_from_excel に入りました: "
    # f"src_file={getattr(block, 'src_file', '')}, "
    # f"src_sheet={getattr(block, 'src_sheet', '')}"
    # )

    src_file = str(getattr(block, "src_file", "") or "").strip()
    src_sheet = str(getattr(block, "src_sheet", "") or "").strip()

    if not src_file:
        return [], ""

    table_path = str(getattr(settings, "table_path", "") or "").strip()

    if not table_path:
        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"table src が指定されていますが table_path が未設定です: {src_file}"
            )
        return [], ""

    # ------------------------------------------------------------
    # Excelファイルの場所
    # - table_path="inbox" の場合は Inbox から探す
    # - それ以外は table_path / src_file として読む
    # ------------------------------------------------------------
    if table_path.lower() == "inbox":
        if inbox_root is None or not sub:
            if hasattr(settings, "add_warning"):
                settings.add_warning(
                    'table_path="inbox" が指定されていますが inbox_root または sub が未設定です。'
                )
            return [], ""

        xlsx_path = resolve_inbox_excel_path_by_filename(
            inbox_root=inbox_root,
            sub=sub,
            file_name=src_file,
        )

        if xlsx_path is None:
            if hasattr(settings, "add_warning"):
                settings.add_warning(
                    f"Inbox内のExcelファイルが見つかりません: {src_file}"
                )
            return [], ""

    else:
        xlsx_path = Path(table_path) / src_file

    if not xlsx_path.exists():
        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"table Excel ファイルが見つかりません: {xlsx_path}"
            )
        return [], ""

    try:
        wb = load_workbook(
            filename=xlsx_path,
            data_only=True,
            read_only=True,
        )

    except Exception as e:
        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"table Excel ファイルを読み込めません: {xlsx_path} / {e}"
            )
        return [], ""

    try:
        if src_sheet:
            if src_sheet not in wb.sheetnames:
                if hasattr(settings, "add_warning"):
                    settings.add_warning(
                        f"table Excel シートが見つかりません: {src_sheet}"
                    )
                return [], ""
            ws = wb[src_sheet]
        else:
            ws = wb[wb.sheetnames[0]]

        rows: list[list[str]] = []
        note_lines: list[str] = []
        in_note = False
        in_memo = False

        for excel_row in ws.iter_rows(values_only=True):
            values = [
                "" if v is None else str(v).strip()
                for v in excel_row
            ]

            # ----------------------------------------------------
            # 右端の空セルを落とす。
            # ----------------------------------------------------
            while values and values[-1] == "":
                values.pop()

            if not values:
                continue

            first_cell = str(values[0] or "").strip()

            # ----------------------------------------------------
            # A列が <note> の行以降は note として扱う。
            # <note> 行そのものは出力しない。
            # ----------------------------------------------------
            if first_cell == "<note>":
                in_note = True
                continue

            # ----------------------------------------------------
            # A列が <memo> の行以降は完全に無視する。
            # ----------------------------------------------------
            if first_cell == "<memo>":
                in_memo = True
                continue

            if in_memo:
                continue

            if in_note:
                note_text = "\t".join(values).strip()
                if note_text:
                    note_lines.append(note_text)
                continue

            rows.append(values)

        return rows, "\n".join(note_lines).strip()        

    finally:
        wb.close()

def parse_table_total_width_cm(value: object) -> float:
    """
    table の width 指定を cm に変換する。

    対応:
    - "100%" → TOTAL_CM
    - "120%" → TOTAL_CM * 1.2
    - "18cm" → 18.0
    - "18" → 18.0
    """
    text = str(value or "").strip().lower()

    if not text:
        return TOTAL_CM

    if text.endswith("%"):
        try:
            pct = float(text[:-1].strip())
        except ValueError:
            return TOTAL_CM

        if pct <= 0:
            return TOTAL_CM

        return TOTAL_CM * pct / 100.0

    if text.endswith("cm"):
        text = text[:-2].strip()

    try:
        cm = float(text)
    except ValueError:
        return TOTAL_CM

    if cm <= 0:
        return TOTAL_CM

    return cm


def normalize_widths_to_total(
    *,
    values: list[float],
    total_cm: float,
) -> list[float]:
    """
    比率指定を total_cm に正規化する。
    """
    cleaned = []

    for v in values:
        try:
            fv = float(v)
        except Exception:
            continue

        if fv > 0:
            cleaned.append(fv)

    if not cleaned:
        return []

    s = sum(cleaned)

    if s <= 0:
        return []

    return [
        total_cm * v / s
        for v in cleaned
    ]


def get_table_widths_cm(
    *,
    rows: list[list[str]],
    block,
    settings,
) -> list[float]:
    """
    table の列幅を決定する。

    優先順位:
    1. col_widths=[...]  cm指定
    2. col_ratio=[...]   比率指定
    3. 自動（文字数）
    """
    if not rows or not rows[0]:
        return []

    n_cols = len(rows[0])
    total_cm = parse_table_total_width_cm(
        getattr(block, "width", "100%")
    )

    col_widths = list(getattr(block, "col_widths", []) or [])
    if col_widths:
        if len(col_widths) == n_cols:
            return [float(v) for v in col_widths]

        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"table col_widths の数が列数と一致しません: "
                f"{len(col_widths)} != {n_cols}"
            )

    col_ratio = list(getattr(block, "col_ratio", []) or [])
    if col_ratio:
        if len(col_ratio) == n_cols:
            return normalize_widths_to_total(
                values=col_ratio,
                total_cm=total_cm,
            )

        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"table col_ratio の数が列数と一致しません: "
                f"{len(col_ratio)} != {n_cols}"
            )

    return compute_col_widths_cm(
        rows,
        total_cm=total_cm,
    )


# ============================================================
# compact table layout
# ============================================================
def set_cell_margins(
    cell,
    *,
    top: int = 40,
    start: int = 60,
    bottom: int = 40,
    end: int = 60,
) -> None:
    """
    セル内余白を twips 単位で設定する。
    1pt = 20twips。
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)

    for name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tcMar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tcMar.append(node)

        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

# ============================================================
# compact table layout
# ============================================================
def apply_compact_table_layout(
    table,
    *,
    mode: str = "none",
) -> None:
    """
    表内段落の行間を詰める。

    mode
        none      : 何もしない
        normal    : 行間 1.00
        tight     : 行間 0.90
        verytight : 行間 0.80
    """
    mode = str(mode or "none").strip().lower()

    if mode == "none":
        return

    if mode == "normal":
        line_spacing = 1.0

    elif mode == "tight":
        line_spacing = 0.9

    elif mode == "verytight":
        line_spacing = 0.8

    else:
        line_spacing = 1.0

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AUTO

        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = line_spacing

# ============================================================
# fixed table layout
# ============================================================
def set_table_fixed_layout(
    table,
    *,
    total_cm: float,
) -> None:
    """
    表ごとに Word の固定レイアウトを設定する。

    - width="12cm" などの指定を表全体へ反映する
    - 列幅だけでなく，表そのものの幅も固定する
    - 前の表が 16cm の場合でも，次の表が引きずられないようにする
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    # ------------------------------------------------------------
    # Word の表レイアウトを fixed にする
    # - autofit=False だけでは tblLayout が明示されない場合がある
    # - tblLayout=fixed を入れることで，表ごとの幅指定を安定させる
    # ------------------------------------------------------------
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = OxmlElement("w:tblLayout")
        tblPr.append(tblLayout)

    tblLayout.set(qn("w:type"), "fixed")

    # ------------------------------------------------------------
    # 表全体の幅を Word XML に明示する
    # - python-docx の列幅指定だけでは表全体の幅が固定されないため
    # - 1cm = 約567twips として dxa 指定する
    # ------------------------------------------------------------
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)

    width_twips = int(float(total_cm) * 567)

    tblW.set(qn("w:w"), str(width_twips))
    tblW.set(qn("w:type"), "dxa")


def add_table_block(
    doc,
    block,
    settings: dict,
    inbox_root: Path | None = None,
    sub: str = "",
) -> None:
    """
    TableBlock を Word 文書へ追加する。
    """

    rows = getattr(block, "rows", [])

    # ------------------------------------------------------------
    # table note 出力用テキスト
    # - \note{...} 由来の note を最初に保持する
    # - Excel の <note> 以下があれば後でここへ結合する
    # ------------------------------------------------------------
    note_text = str(getattr(block, "note", "") or "").strip()

    # ------------------------------------------------------------
    # src=["file.xlsx","sheet"] がある場合は Excel から表を読む。
    # 手入力 rows がある場合は，既存の手入力 table を優先する。
    # ------------------------------------------------------------
    if str(getattr(block, "src_file", "") or "").strip():
        excel_rows, excel_note = load_table_rows_from_excel(
            block=block,
            settings=settings,
            inbox_root=inbox_root,
            sub=sub,
        )

        if excel_rows:
            rows = excel_rows
            block.rows = excel_rows

        if excel_note:
            # ------------------------------------------------------------
            # Excelの <note> 以下を table note として反映する。
            # 既に \note{...} がある場合は，後ろに結合する。
            # ------------------------------------------------------------
            if note_text:
                note_text = note_text + "\n" + excel_note
            else:
                note_text = excel_note

    if not rows:
        return
    
    # ------------------------------------------------------------
    # tableセル内記法の正規化
    # - <改行> をセル内改行へ変換する
    # - 画像指定 <図1.png,width=2.5cm,align=center> は
    #   後段の apply_table_cell_images() で処理するため残す
    # ------------------------------------------------------------
    # ------------------------------------------------------------
    # table rows の正規化
    # - Excel読み込みでは行ごとに列数が異なる場合があるため，
    #   先に不足セルを "" で補って矩形表にする。
    # - その後，<改行> をセル内改行へ変換する。
    # ------------------------------------------------------------
    rows = normalize_table_row_lengths(rows)
    rows = normalize_table_rows_for_cell_text(rows)
    block.rows = rows

    preset = resolve_table_preset(getattr(block, "style", "simple"))

    header_rows = max(0, min(int(getattr(block, "header_rows", 1)), len(rows)))
    header_cols = max(0, min(int(getattr(block, "header_cols", 0)), len(rows[0])))

    font_name = preset.get("font_name", "Meiryo")
    header_bg = preset.get("header_bg") or "#EEEEEE"
    header_fg = preset.get("header_fg") or "#000000"
    body_bg = preset.get("body_bg")
    body_fg = preset.get("body_fg") or "#000000"

    inner_h = bool(preset.get("inner_h", True))
    inner_v = bool(preset.get("inner_v", True))
    outer = bool(preset.get("outer", True))
    zebra = bool(preset.get("zebra", False))
    outer_mode = preset.get("outer_mode", "box")

    body_size = int(getattr(settings, "table_body_size", 10))
    header_size = int(getattr(settings, "table_header_size", body_size + 1))

    # ------------------------------------------------------------
    # 表内文字サイズ
    # - 本文全体ではなく，この table ブロック内の文字だけを調整する
    # ------------------------------------------------------------
    fontsize = str(
        getattr(block, "fontsize", "normal")
    ).strip().lower()

    if fontsize == "small":
        body_size = max(6, body_size - 1)
        header_size = max(7, header_size - 1)

    elif fontsize == "verysmall":
        body_size = max(6, body_size - 2)
        header_size = max(7, header_size - 2)

    elif fontsize not in {"normal", ""}:
        if hasattr(settings, "add_warning"):
            settings.add_warning(
                f"table fontsize が不正です: {fontsize}"
            ) 

    table_settings = getattr(settings, "table", None)

    if table_settings is not None:
        caption_position = str(
            getattr(table_settings, "caption_position", "top")
        ).strip().lower()
    else:
        caption_position = str(
            getattr(settings, "table_caption_position", "top")
        ).strip().lower()

    caption_text = build_table_caption_text(
        block=block,
        settings=settings,
    )

    # ------------------------------------------------------------
    # caption が無い表でも，本文との間に少し余白を作る。
    # caption がある場合は caption 段落の space_before で調整する。
    # ------------------------------------------------------------
    if not caption_text:
        p = doc.add_paragraph()
        # p.paragraph_format.space_before = Pt(6)
        # p.paragraph_format.space_after = Pt(0)
        # p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)

    if caption_position == "top" and caption_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # ------------------------------------------------------------
        # 表キャプションの前後余白
        # ------------------------------------------------------------
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(caption_text)
        apply_font_run(
            run,
            font_name,
            int(getattr(settings, "size", 11)),
            "#000000",
            bold=True,
        )

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))

    # ------------------------------------------------------------
    # 表全体を固定レイアウトにする
    # - width="12cm" などの指定を表全体へ反映する
    # - col_ratio / col_widths で計算した列幅と，表全体の幅を一致させる
    # - 前の表が 16cm でも，次の表の width 指定が独立して効くようにする
    # ------------------------------------------------------------
    total_cm = parse_table_total_width_cm(
        getattr(block, "width", "100%")
    )

    set_table_fixed_layout(
        table,
        total_cm=total_cm,
    )

    # ------------------------------------------------------------
    # 表位置
    # width を 100% より大きくした場合も，
    # 左右均等に余白を減らして中央配置する。
    # ------------------------------------------------------------
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ------------------------------------------------------------
    # Word が自動で列幅を変更しないようにする
    # ------------------------------------------------------------
    table.autofit = False

    widths_cm = get_table_widths_cm(
        rows=rows,
        block=block,
        settings=settings,
    )
    apply_docx_col_widths(table, widths_cm)

    spans = compute_spans_markers(
        rows,
        use_up=True,
        use_left=True,
        header_rows=header_rows,
        strict_rect=True,
    )

    merge_docx_by_spans(
        table,
        rows,
        spans,
        font_name=font_name,
        base_size=body_size,
        header_size=header_size,
        header_fg=header_fg,
        body_fg=body_fg,
        header_bg=header_bg,
        body_bg=body_bg,
        zebra=zebra,
        header_same=False,
        body_bg_on=body_bg is not None,
        header_rows=header_rows,
        header_cols=header_cols,
    )

    # ------------------------------------------------------------
    # tableセル内画像
    # - セル内容が <図1.png,width=2.5cm,align=center> の場合，
    #   table_fig_path から画像を読み込み，そのセルへ挿入する。
    # - 通常文字セル，<同上>，<同左> の処理は既存の
    #   merge_docx_by_spans() に任せる。
    # ------------------------------------------------------------
    apply_table_cell_images(
        table=table,
        rows=rows,
        widths_cm=widths_cm,
        settings=settings,
        inbox_root=inbox_root,
        sub=sub,
    )

    # ------------------------------------------------------------
    # compact
    # ------------------------------------------------------------
    apply_compact_table_layout(
        table,
        mode=str(
            getattr(block, "compact", "none")
        ),
    )

    apply_table_borders_robust(
        table,
        spans,
        rows=rows,
        inner_h=inner_h,
        inner_v=inner_v,
        outer=outer,
        outer_mode=outer_mode,
        sz_inner=6,
        sz_outer=12,
        color="000000",
    )

    # ------------------------------------------------------------
    # caption / note が無い表でも，表の下に少し余白を作る。
    # Word の table には paragraph の space_after が無いため，
    # 表の直後に小さい空段落を入れる。
    # ------------------------------------------------------------
    if not note_text and not (caption_position == "bottom" and caption_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0


    if caption_position == "bottom" and caption_text:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.line_spacing = 1.0

        run = p.add_run(caption_text)
        apply_font_run(
            run,
            font_name,
            int(getattr(settings, "size", 11)),
            "#000000",
            bold=True,
        )
        
    note = note_text

    compact = str(
        getattr(block, "compact", "none")
    ).strip().lower()

    if compact == "normal":
        note_line_spacing = 1.0

    elif compact == "tight":
        note_line_spacing = 0.9

    elif compact == "verytight":
        note_line_spacing = 0.8

    else:
        note_line_spacing = 1.0

    if note:
        p = doc.add_paragraph()

        # ------------------------------------------------------------
        # 表 note の前後余白
        # - space_before : 表と note の間
        # - space_after  : note と次本文の間
        # ------------------------------------------------------------
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = note_line_spacing

        run = p.add_run(note)
        apply_font_run(
            run,
            font_name,
            max(8, body_size - 1),
            "#444444",
            bold=False,
        )

def build_table_caption_text(
    *,
    block,
    settings,
) -> str:
    """
    table_caption / table_caption_sep / table_numbers / table_format に従って
    表キャプションを作る。

    caption が無い場合は、
      - キャプションを出力しない
      - 表番号を進めない
      - label を登録しない
    """
    caption = str(
        getattr(block, "caption", "") or getattr(block, "title", "") or ""
    ).strip()

    # ------------------------------------------------------------
    # caption が無い表は番号を付けない。
    # label 登録もしない。
    # ------------------------------------------------------------
    if not caption:
        return ""

    label = str(getattr(block, "label", "") or "").strip()

    table_settings = getattr(settings, "table", None)

    if table_settings is None:
        table_caption = str(getattr(settings, "table_caption", "表"))
        table_caption_sep = str(getattr(settings, "table_caption_sep", "："))
        table_numbers = int(getattr(settings, "table_numbers", 1))
        table_format = str(getattr(settings, "table_format", "{1}"))

        if table_numbers:
            number = int(getattr(settings, "_table_counter", 0)) + 1
            setattr(settings, "_table_counter", number)

            number_text = table_format.replace("{1}", str(number))

            if label and hasattr(settings, "register_label"):
                settings.register_label(
                    label=label,
                    value=number_text,
                )

            return f"{table_caption}{number_text}{table_caption_sep}{caption}"

        return caption

    if label and hasattr(settings, "register_label"):
        settings.register_label(
            label=label,
            value=table_settings.format_number_core(),
        )

    caption_text = table_settings.format_caption(caption)
    table_settings.increment()

    return caption_text