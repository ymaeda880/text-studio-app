# -*- coding: utf-8 -*-
# text_studio_app/lib/wordtex/itemize/builder.py
# ============================================================
# wordTex itemize builder
# ============================================================

from __future__ import annotations

from docx.document import Document as DocumentObject
from docx.shared import Pt

from lib.wordtex.blocks import ItemizeBlock, ItemizeItem
from lib.wordtex.settings import WordTexSettings
from lib.wordtex.builders.paragraph_builder import add_inline_runs


def to_alpha_number(index: int, *, upper: bool = False) -> str:
    n = max(1, int(index))
    chars: list[str] = []

    while n > 0:
        n -= 1
        chars.append(chr(ord("A" if upper else "a") + (n % 26)))
        n //= 26

    return "".join(reversed(chars))


def to_roman_number(index: int, *, upper: bool = False) -> str:
    n = max(1, int(index))

    values = [
        (1000, "m"),
        (900, "cm"),
        (500, "d"),
        (400, "cd"),
        (100, "c"),
        (90, "xc"),
        (50, "l"),
        (40, "xl"),
        (10, "x"),
        (9, "ix"),
        (5, "v"),
        (4, "iv"),
        (1, "i"),
    ]

    result = ""

    for value, roman in values:
        while n >= value:
            result += roman
            n -= value

    return result.upper() if upper else result


def make_marker(
    *,
    style: str,
    index: int,
) -> str:
    """
    styleから箇条書きマーカーを作る。

    style=""
        マーカーなし

    style="・"
        ・

    style="#."
        1.

    style="(#)"
        (1)

    style="a."
        a.

    style="A."
        A.

    style="i."
        i.

    style="I."
        I.
    """
    style_text = str(style or "").strip()

    if not style_text:
        return ""

    if "#" in style_text:
        return style_text.replace("#", str(index))

    if "a" in style_text:
        return style_text.replace(
            "a",
            to_alpha_number(index, upper=False),
        )

    if "A" in style_text:
        return style_text.replace(
            "A",
            to_alpha_number(index, upper=True),
        )

    if "i" in style_text:
        return style_text.replace(
            "i",
            to_roman_number(index, upper=False),
        )

    if "I" in style_text:
        return style_text.replace(
            "I",
            to_roman_number(index, upper=True),
        )

    return style_text


def add_itemize_item_paragraph(
    *,
    doc: DocumentObject,
    settings: WordTexSettings,
    item: ItemizeItem,
    marker: str,
    level: int,
) -> None:
    """
    item 1件を Word 段落として出力する。
    """
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]

    # ------------------------------------------------------------
    # インデント
    # - Wordのリスト機能は使わず，wordTex側でマーカーを文字として出す
    # - renderer差異を避けるため
    # ------------------------------------------------------------
    left_indent = Pt(18 * level)
    hanging_indent = Pt(10.5)

    p.paragraph_format.left_indent = left_indent
    p.paragraph_format.first_line_indent = -hanging_indent
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    label = str(item.label or "").strip()
    text = settings.resolve_refs(str(item.text or "").strip())

    if label:
        p.add_run(label)
        p.add_run("　")
        add_inline_runs(
            paragraph=p,
            text=text,
        )
        return

    if marker:
        p.add_run(marker)
        p.add_run(" ")

    add_inline_runs(
        paragraph=p,
        text=text,
    )


def add_itemize_block(
    *,
    doc: DocumentObject,
    settings: WordTexSettings,
    block: ItemizeBlock,
    level: int = 1,
) -> None:
    """
    ItemizeBlock を Word に出力する。

    入れ子は再帰で処理する。
    """
    style = str(block.style or "").strip()

    for idx, item in enumerate(block.items, start=1):
        marker = make_marker(
            style=style,
            index=idx,
        )

        add_itemize_item_paragraph(
            doc=doc,
            settings=settings,
            item=item,
            marker=marker,
            level=level,
        )

        for child in item.children:
            add_itemize_block(
                doc=doc,
                settings=settings,
                block=child,
                level=level + 1,
            )