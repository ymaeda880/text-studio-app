# pages/20_校正.py — 解析（校正方針：ページ/行/理由）→ 本校正
from __future__ import annotations
from io import BytesIO
from typing import List, Tuple
import streamlit as st
from openai import OpenAI

st.set_page_config(page_title="Text Studio / 校正", page_icon="📝", layout="wide")

MODEL_OPTIONS = ["gpt-5-mini", "gpt-5-nano"]
DEFAULT_MODEL = "gpt-5-mini"

# 1ページあたりの行数（.txt/.docx を段落→行にした上で擬似ページ分割）
LINES_PER_PAGE = 40

if "chat_model" not in st.session_state:
    st.session_state["chat_model"] = DEFAULT_MODEL

with st.sidebar:
    st.header("設定")
    st.radio(
        "🧠 使用モデル",
        MODEL_OPTIONS,
        index=MODEL_OPTIONS.index(st.session_state["chat_model"]),
        key="chat_model",
    )
    keep_formatting = st.checkbox("改行・段落を保持（本校正時）", value=True)
    show_report = st.checkbox("本校正後に『修正ポイント要約』も付与", value=False)
    lpp = st.number_input("ページ行数（表示用）", min_value=20, max_value=100, value=LINES_PER_PAGE, step=5)

st.title("📝 校正 — 解析（ページ/行/理由） → 本校正")

st.write("Word（.docx）またはテキスト（.txt）をドロップしてください。最初に『解析』ボタンで**校正方針**を一覧表示します。")


# ====== ファイル読み込み ======
def read_txt(file) -> str:
    data = file.read()
    for enc in ("utf-8", "utf-16", "shift_jis", "cp932"):
        try:
            return data.decode(enc)
        except Exception:
            continue
    return data.decode("utf-8", errors="ignore")


def read_docx(file) -> str:
    from docx import Document
    bio = BytesIO(file.read())
    doc = Document(bio)
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    return "\n".join(texts).strip()


def load_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return read_txt(uploaded_file)
    elif name.endswith(".docx"):
        return read_docx(uploaded_file)
    else:
        st.error("対応形式は .txt または .docx です。")
        st.stop()


# ====== 行番号・ページ番号の付与 ======
def to_numbered_lines(raw: str) -> List[str]:
    """
    原文を行単位に分解し、空行も保持。ここではすでに \n 区切りのテキスト。
    """
    lines = raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    return lines


def page_and_line(idx: int, lines_per_page: int) -> Tuple[int, int]:
    """
    0始まりの行index -> (ページ番号1始まり, 行番号1始まり)
    """
    page = idx // lines_per_page + 1
    line_in_page = idx % lines_per_page + 1
    return page, line_in_page


def render_preview_with_numbers(lines: List[str], lines_per_page: int) -> str:
    """
    画面用プレビュー（[p:行] prefix）を付けたテキスト。
    例: [1:01] 最初の行
    """
    out = []
    for i, t in enumerate(lines):
        p, ln = page_and_line(i, lines_per_page)
        out.append(f"[{p}:{ln:02d}] {t}")
    return "\n".join(out)


# ====== OpenAI 呼び出し ======
def openai_client() -> OpenAI:
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])


def analyze_issues(model: str, lines: List[str], lines_per_page: int) -> str:
    """
    校正方針（何を/どう直すか/理由）を Markdown表で返す。
    大きい文書は「ページ単位」で分割して逐次解析。
    出力列：| Page | Line | Issue | Original | Suggestion | Reason |
    """
    client = openai_client()
    md_tables: List[str] = []

    total_pages = (len(lines) + lines_per_page - 1) // lines_per_page
    for pg in range(total_pages):
        start = pg * lines_per_page
        end = min((pg + 1) * lines_per_page, len(lines))
        # 該当ページの行のみを、[p:行]付きで渡す
        page_chunk = []
        for i in range(start, end):
            p, ln = page_and_line(i, lines_per_page)
            page_chunk.append(f"[{p}:{ln:02d}] {lines[i]}")
        page_text = "\n".join(page_chunk)

        sys_inst = (
            "あなたは厳密な日本語校正リーダーです。以下の番号付きテキストを読み、"
            "『何をどのように直すべきか』を、具体的な理由とともに一覧化してください。"
            "行頭の [page:line] を必ず参照して位置を示し、過度な意訳は避けてください。"
            "助詞（てにおは）、主述一致、冗長、重複、語順、誤字脱字、用語誤用、文体の不統一に特に注意。"
            "出力は **Markdownの表** で、列は次の順：Page | Line | Issue | Original | Suggestion | Reason。"
            "Original は該当箇所の短い抜粋（最大20字）に留めてください。"
        )
        user_msg = f"次のテキスト（このページのみ）を解析してください：\n---\n{page_text}"

        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": sys_inst},
                {"role": "user", "content": user_msg},
            ],
        )
        md_tables.append(resp.choices[0].message.content.strip())

    # ページごとに小見出しを付けて連結
    out = []
    for i, tbl in enumerate(md_tables, 1):
        out.append(f"#### Page {i}\n\n{tbl}\n")
    return "\n".join(out)


def proofread(model: str, content: str, keep_layout: bool, want_report: bool) -> str:
    """
    本校正。長文はチャンク分割（約6000字）し結合。
    """
    client = openai_client()
    CHUNK = 6000
    chunks = [content[i:i+CHUNK] for i in range(0, len(content), CHUNK)] or [content]

    fixed_parts = []
    for chunk in chunks:
        sys_inst = (
            "あなたは厳密な日本語校正者です。以下を徹底してください：\n"
            "- てにおは、助詞、主述一致、誤字脱字、表記ゆれを修正\n"
            "- 用語誤用の是正・語順の自然化（意味は変えない）\n"
            + ("- 改行・段落は可能な限り維持\n" if keep_layout else "")
            + "- 出力は校正後の本文のみ（前置き不要）"
        )
        user_msg = f"次の本文を校正してください：\n---\n{chunk}"

        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": sys_inst},
                {"role": "user", "content": user_msg},
            ],
        )
        fixed_parts.append(resp.choices[0].message.content.strip())

    fixed_text = ("\n\n" if keep_layout else "\n").join(fixed_parts).strip()

    if want_report:
        rep_inst = (
            "次の原文と校正後本文の差分観点で、主な修正ポイントを最大8項目で箇条書きに。\n"
            "例：助詞/活用、語順、冗長/重複、誤字脱字、文体統一、用語誤用など。"
        )
        rep_user = f"原文:\n{content[:4000]}\n\n校正後:\n{fixed_text[:4000]}"
        rep = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": rep_inst},
                {"role": "user", "content": rep_user},
            ],
        )
        fixed_text += f"\n\n---\n【修正ポイント（要約）】\n{rep.choices[0].message.content.strip()}"

    return fixed_text


# ====== UI ======
col_u, col_btn1, col_btn2 = st.columns([3, 1, 1])
with col_u:
    up = st.file_uploader(".docx / .txt をアップロード", type=["docx", "txt"])
with col_btn1:
    do_analyze = st.button("① 解析（校正方針を表示）", type="secondary", use_container_width=True, disabled=not up)
with col_btn2:
    do_fix = st.button("② 本校正を実行", type="primary", use_container_width=True, disabled=not up)

if up:
    src_text = load_text(up)
    if not src_text.strip():
        st.warning("ファイル内にテキストが見つかりませんでした。")
        st.stop()

    # 行番号プレビュー
    lines = to_numbered_lines(src_text)
    st.subheader("👀 行番号付きプレビュー")
    st.caption(f"表示上のページ行数: {lpp} 行/ページ（擬似割り付け）")
    st.text_area("原文（番号付きプレビュー）", value=render_preview_with_numbers(lines, lpp), height=260)

    if do_analyze:
        with st.spinner("解析中（校正方針を抽出）…"):
            plan_md = analyze_issues(st.session_state["chat_model"], lines, lpp)
        st.success("解析が完了しました。ページ/行/理由つきで方針を表示します。")
        st.subheader("📋 校正方針（まず何をどう直すか）")
        st.markdown(plan_md, unsafe_allow_html=False)

    if do_fix:
        with st.spinner("本校正の実行中…"):
            result = proofread(
                model=st.session_state["chat_model"],
                content=src_text,
                keep_layout=keep_formatting,
                want_report=show_report,
            )
        st.success("校正完了！下の結果をコピー/ダウンロードできます。")
        st.subheader("🧾 校正結果")
        st.text_area("校正結果", value=result, height=420)
        st.download_button(
            "校正結果をダウンロード (.txt)",
            data=result.encode("utf-8"),
            file_name=f"proofread_{up.name.rsplit('.', 1)[0]}.txt",
            mime="text/plain",
        )
else:
    st.info("ファイルを選んで『① 解析』→『② 本校正』の順に実行してください。")
