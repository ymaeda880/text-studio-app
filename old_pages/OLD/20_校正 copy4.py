# pages/20_校正.py — 解析（校正方針：ページ/行/理由）オンリー版
# （PDFが画像PDFなら警告して停止／OCR機能は削除）

from __future__ import annotations
from typing import List, Tuple, Dict
from pathlib import Path
import sys
import io
import base64
import datetime as _dt

import streamlit as st
from openai import OpenAI

# ===== 共有ライブラリ（common_lib）をパスに追加 =====
PROJECTS_ROOT = Path(__file__).resolve().parents[3]  # or 3 for pages
if str(PROJECTS_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECTS_ROOT))

# 料金見積りユーティリティ
from common_lib.openai.costs import (
    ChatUsage,
    estimate_chat_cost,
    render_chat_only_summary,
    DEFAULT_USDJPY,
    MODEL_PRICES_USD,
)

# === lib の利用（読込 / 貼り付け整形） ※OCRなし ===
from lib.text_loaders import (
    load_text_generic, extract_pdf_text, load_text_from_paste
)

# ------------------------------------------------------------
# モード定義（解析のみ使用）
# ------------------------------------------------------------
MODE_DEFS: Dict[str, Dict[str, str]] = {
    "厳格校正": {
        "desc": "助詞・主述一致・冗長/重複・語順・誤字脱字・用語誤用・文体不統一まで広く対象。意味は変えず最適化のための方針を抽出。",
        "analyze_inst": (
            "あなたは厳密な日本語校正リーダーです。以下の番号付きテキストを読み、"
            "『何をどのように直すべきか』を、具体的な理由とともに一覧化してください。\n"
            "行頭の [page:line] を必ず参照して位置を示し、過度な意訳は避けてください。\n"
            "助詞（てにおは）、主述一致、冗長、重複、語順、誤字脱字、用語誤用、文体の不統一に特に注意。\n"
            "「原文」 は該当箇所の短い抜粋（最大20字）に留めてください。"
        ),
    },

    "通常校正": {
        "desc": "助詞・主述一致・語順・誤字脱字・文体不統一などを対象。意味は変えず最適化のための方針を抽出。",
        "analyze_inst": (
            "あなたは通常の日本語校正リーダーです。以下の番号付きテキストを読み、"
            "『何をどのように直すべきか』を、理由とともに一覧化してください。\n"
            "行頭の [page:line] を必ず参照して位置を示し、過度な意訳は避けてください。\n"
            "語順変更やスタイル統一などの裁量的変更は提案しないでください。"
            "「原文」 は該当箇所の短い抜粋（最大20字）に留めてください。"
        ),
    },

    "簡易校正（ミス最小修正）": {
        "desc": "明白なミスのみ（てにおは・助詞・誤字脱字・明確な変換ミス）に絞った方針を抽出。",
        "analyze_inst": (
            "あなたは日本語の軽微校正リーダーです。以下の番号付きテキストから、"
            "『明白なミス（てにおは・助詞の誤り、誤字脱字、明確な変換ミス）』のみを抽出してください。"
            "語順変更やスタイル統一などの裁量的変更は提案しないでください。"
            "「原文」 は短い抜粋（最大20字）。「理由」 は“明白な誤り”である根拠を簡潔に。"
        ),
    },
}

# ------------------------------------------------------------
# 解析時に常に付与する共通プロンプト
# ------------------------------------------------------------
COMMON_PROMPT = (
    "【共通方針（厳守）】\n"
    "それぞれの校正の必要性を0から10までの数字で表して明記してください：\n"
    "- 必要性が高い構成には大きい数字をつけてください．\n"
    "- 誤字．脱字など，校正が必ず必要なものを10と評価してください．\n"
    "- 校正する必要の無いものを0と評価してください．\n"
    "- 出力は **Markdownの表** で、列は次の順：頁 | 行 | 重要度 | 原文 | 修正案 | 理由"
)

# ------------------------------------------------------------
# UI定数
# ------------------------------------------------------------
st.set_page_config(page_title="Text Studio / 解析（校正方針）", page_icon="📝", layout="wide")

MODEL_OPTIONS = ["gpt-5-mini", "gpt-5-nano"]
DEFAULT_MODEL = "gpt-5-mini"
DEFAULT_MODE = "通常校正"
LINES_PER_PAGE = 40  # 1ページあたりの表示行数

# セッション初期化
if "chat_model" not in st.session_state:
    st.session_state["chat_model"] = DEFAULT_MODEL
if "proof_mode" not in st.session_state:
    st.session_state["proof_mode"] = DEFAULT_MODE
if "tok_in_total" not in st.session_state:
    st.session_state["tok_in_total"] = 0
if "tok_out_total" not in st.session_state:
    st.session_state["tok_out_total"] = 0

# ------------------------------------------------------------
# 表示ユーティリティ
# ------------------------------------------------------------
def display_pdf_bytes(data: bytes, height: int = 600):
    """Streamlit PDF表示（streamlit[pdf] があれば st.pdf）。なければ iframe 埋め込み。"""
    try:
        st.pdf(data, height=height)  # Streamlit 1.31+ / streamlit[pdf]
    except Exception:
        b64 = base64.b64encode(data).decode()
        st.markdown(
            f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="{height}px"></iframe>',
            unsafe_allow_html=True
        )

def to_numbered_lines(raw: str) -> List[str]:
    return raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")

def render_preview_with_numbers(lines: List[str], lines_per_page: int) -> str:
    return "\n".join(f"[{(i//lines_per_page)+1}:{(i%lines_per_page)+1:02d}] {t}" for i, t in enumerate(lines))

# ------------------------------------------------------------
# OpenAIクライアントと使用量抽出
# ------------------------------------------------------------
def openai_client() -> OpenAI:
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def _add_usage_from_resp(resp) -> Tuple[int, int]:
    u = getattr(resp, "usage", None) or {}
    prompt = getattr(u, "prompt_tokens", None) if hasattr(u, "prompt_tokens") else u.get("prompt_tokens", 0)
    completion = getattr(u, "completion_tokens", None) if hasattr(u, "completion_tokens") else u.get("completion_tokens", 0)
    try:
        prompt = int(prompt or 0)
    except Exception:
        prompt = 0
    try:
        completion = int(completion or 0)
    except Exception:
        completion = 0
    return prompt, completion

# ------------------------------------------------------------
# 方針プレビュー＆追加プロンプト
# ------------------------------------------------------------
def get_analyze_instruction(mode: str) -> str:
    return MODE_DEFS.get(mode, MODE_DEFS["厳格校正"])["analyze_inst"]

def build_sys_inst(base: str, extra: str) -> str:
    """base（モード固有） → extra（UI入力） → COMMON_PROMPT（共通指示）の順で連結"""
    parts = [base.strip()]
    if (extra or "").strip():
        parts.append("【追加指示（厳守）】\n" + extra.strip())
    if (globals().get("COMMON_PROMPT") or "").strip():
        parts.append(COMMON_PROMPT.strip())
    return "\n\n".join(parts)

def render_policy_preview(*, mode: str) -> str:
    analyze_base = get_analyze_instruction(mode)

    st.subheader("🧭 解析プロンプト（System）")
    st.code(analyze_base, language="markdown")
    with st.expander("共通方針（Systemに毎回付与）", expanded=False):
        st.code(COMMON_PROMPT.strip(), language="markdown")

    st.markdown("**✍️ 追加プロンプト（任意）**")
    extra = st.text_area(
        "追加プロンプト",
        key="extra_user_prompt",
        placeholder="例）外来語はカタカナ優先。製品名や固有名詞は原文どおりに保持。",
        height=100,
    )
    return extra or ""

# ------------------------------------------------------------
# 解析結果の Word / PDF 生成
#   - Word: 既存の build_policy_docx_bytes を使用
#   - PDF : 簡易整形してテキストで書き出し（reportlab 使用）
# ------------------------------------------------------------
def _parse_plan_md_tables(md: str) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    if not md:
        return items

    to_jp = {
        "page": "頁", "line": "行", "issue": "重要度",
        "original": "原文", "suggestion": "修正案", "reason": "理由",
        "頁": "頁", "ページ": "頁", "行": "行",
        "重要度": "重要度", "問題点": "重要度", "原文": "原文",
        "修正案": "修正案", "提案": "修正案", "理由": "理由", "根拠": "理由",
    }
    expected = ["頁", "行", "重要度", "原文", "修正案", "理由"]

    def _norm_head_cell(s: str) -> str:
        key = s.strip().lower()
        return to_jp.get(key, to_jp.get(s.strip(), s.strip()))

    def _row_to_dict(ln: str, cols_jp: List[str]) -> Dict[str, str] | None:
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if len(cells) < len(cols_jp):
            return None
        row = dict(zip(cols_jp, cells[:len(cols_jp)]))
        for k in expected:
            row.setdefault(k, "")
        return row

    lines = [ln for ln in (l.strip() for l in md.splitlines()) if ln]
    header_seen = False
    cols_jp: List[str] = []

    for ln in lines:
        if not (ln.startswith("|") and "|" in ln):
            continue

        raw_cells = [c.strip() for c in ln.strip("|").split("|")]
        # 区切り行（--- / :--- / ---: / :---: 等）はスキップ
        def _is_md_separator(cell: str) -> bool:
            s = cell.strip().replace("-", "").replace(":", "")
            return s == ""  # ダッシュとコロン以外の文字が無ければ区切り

        if all(_is_md_separator(c) for c in raw_cells):
            continue

        if not header_seen:
            norm = [_norm_head_cell(c) for c in raw_cells]
            if len(set(norm) & set(expected)) >= 3:
                cols_jp = [c if c in expected else "" for c in norm]
                if len(cols_jp) < 6:
                    cols_jp += [""] * (6 - len(cols_jp))
                cols_jp = [cols_jp[i] or expected[i] for i in range(6)]
                header_seen = True
                continue

        if header_seen and cols_jp:
            row = _row_to_dict(ln, cols_jp)
            if row:
                items.append(row)

    return items

def build_policy_docx_bytes(
    *,
    original_numbered_preview: str,
    plan_md: str,
    model: str,
    mode: str,
    extra_prompt: str,
    src_name: str
) -> Tuple[bytes, str]:
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        # フォールバック: テキスト
        lines: List[str] = []
        lines.append("=== 校正方針レポート（TXTフォールバック） ===")
        lines.append(f"生成日時: {_dt.datetime.now():%Y-%m-%d %H:%M:%S}")
        lines.append(f"モデル: {model}")
        lines.append(f"モード: {mode}")
        if (extra_prompt or "").strip():
            lines.append(f"追加プロンプト: {extra_prompt.strip()}")
        lines.append(f"入力ファイル: {src_name or '-'}")
        lines.append("\n--- 原文（行番号つき） ---\n")
        lines.append(original_numbered_preview or "(空)")
        lines.append("\n--- 校正方針（Markdown） ---\n")
        lines.append(plan_md or "(なし)")
        data = "\n".join(lines).encode("utf-8")
        return data, ".txt"

    doc = Document()
    title = doc.add_heading("校正方針レポート", 0)
    try:
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    p = doc.add_paragraph()
    p.add_run("生成日時: ").bold = True
    p.add_run(_dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p.add_run("\nモデル: ").bold = True
    p.add_run(model)
    p.add_run("\nモード: ").bold = True
    p.add_run(mode)
    if (extra_prompt or "").strip():
        p.add_run("\n追加プロンプト: ").bold = True
        p.add_run(extra_prompt.strip())
    p.add_run("\n入力ファイル: ").bold = True
    p.add_run(src_name or "-")

    # 原文（行番号つき）
    doc.add_heading("原文（行番号つき）", level=1)
    para = doc.add_paragraph()
    run = para.add_run(original_numbered_preview if original_numbered_preview else "(空)")
    try:
        run.font.name = "Courier New"
        run.font.size = Pt(10)
    except Exception:
        pass

    # 校正方針（Markdownそのまま）
    doc.add_heading("校正方針（Markdown）", level=1)
    para2 = doc.add_paragraph()
    r2 = para2.add_run(plan_md if plan_md else "(なし)")
    try:
        r2.font.name = "Courier New"
        r2.font.size = Pt(10)
    except Exception:
        pass

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue(), ".docx"

def build_policy_pdf_bytes(
    *,
    original_numbered_preview: str,
    plan_md: str,
    model: str,
    mode: str,
    extra_prompt: str,
    src_name: str
) -> bytes:
    """整形PDF（原文は1行=1行で改行保持、長行は枠内で自動折返し）。"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, XPreformatted
        )
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.lib.units import mm
    except ImportError:
        return b""

    # 日本語CIDフォント登録
    font_name = None
    for fname in ("HeiseiMin-W3", "HeiseiKakuGo-W5"):
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(fname))
            font_name = fname
            break
        except Exception:
            continue
    if not font_name:
        return b""

    # 解析行（Markdown表→配列）
    items = _parse_plan_md_tables(plan_md)
    headers = ["頁", "行", "重要度", "原文", "修正案", "理由"]
    table_data = [headers] + [[
        (it.get("頁") or ""), (it.get("行") or ""), (it.get("重要度") or ""),
        (it.get("原文") or ""), (it.get("修正案") or ""), (it.get("理由") or "")
    ] for it in items]

    # ドキュメント
    buf = io.BytesIO()
    pagesize = A4
    margin = 18 * mm
    doc = SimpleDocTemplate(
        buf, pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=20 * mm, bottomMargin=18 * mm,
        title="校正方針レポート",
    )

    # スタイル
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleJP", fontName=font_name, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Meta", fontName=font_name, fontSize=10, leading=13, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle(name="H1", fontName=font_name, fontSize=13, leading=16, spaceBefore=6, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body", fontName=font_name, fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="Mono", fontName=font_name, fontSize=9.5, leading=13))  # 原文行用

    # ページ番号
    def _page_number(canvas, doc_):
        canvas.setFont(font_name, 9)
        canvas.drawRightString(pagesize[0] - doc_.rightMargin, 12 * mm, f"{doc_.page}")

    # 原文：1行=1行ブロック（ここが新規の make_pre_block_lines 相当）
    def _make_pre_block_lines(text: str, text_width):
        """
        各原文行を XPreformatted で1行ずつテーブル行にする。
        長い行はセル内で自動折り返し、空行も保持。
        """
        rows = []
        for raw in (text or "").splitlines() or ["(空)"]:
            s = (raw.replace("\t", "    ")) if raw else " "
            xp = XPreformatted(s, styles["Mono"])
            rows.append([xp])

        t = Table(rows, colWidths=[text_width], repeatRows=0)
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    # 幅計算
    text_width = pagesize[0] - doc.leftMargin - doc.rightMargin

    # 方針テーブル（列幅）
    col_w = [
        14 * mm,   # 頁
        14 * mm,   # 行
        18 * mm,   # 重要度
    ]
    remain = text_width - sum(col_w)
    col_w += [remain * 0.36, remain * 0.32, remain * 0.32]  # 原文/修正案/理由

    # セルParagraph化（自動折返し）
    def _p(s: str) -> Paragraph:
        return Paragraph((s or "").replace("\n", "<br/>"), styles["Body"])

    table_para = [table_data[0]]
    for r in table_data[1:]:
        table_para.append([_p(x) for x in r])

    tbl = Table(table_para, colWidths=col_w, repeatRows=1)
    zebra = colors.Color(0.95, 0.95, 0.98)
    header_bg = colors.Color(0.88, 0.90, 0.95)
    grid = colors.Color(0.75, 0.78, 0.85)
    st_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        ("ALIGN", (0, 0), (2, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("GRID", (0, 0), (-1, -1), 0.3, grid),
    ]
    for i in range(1, len(table_para)):
        if i % 2 == 1:
            st_cmds.append(("BACKGROUND", (0, i), (-1, i), zebra))
    tbl.setStyle(TableStyle(st_cmds))

    # 組み立て
    story = []
    story.append(Paragraph("校正方針レポート", styles["TitleJP"]))
    meta_lines = [
        f"生成日時：{_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        f"モデル：{model}",
        f"モード：{mode}",
        f"入力ファイル：{src_name or '-'}",
    ]
    if (extra_prompt or "").strip():
        meta_lines.insert(3, f"追加プロンプト：{extra_prompt.strip()}")
    for ln in meta_lines:
        story.append(Paragraph(ln, styles["Meta"]))

    story.append(Spacer(1, 4 * mm))
    story.append(Paragraph("原文（行番号つき）", styles["H1"]))
    story.append(Spacer(1, 1.5 * mm))
    story.append(_make_pre_block_lines(original_numbered_preview or "(空)", text_width))  # ← ここが新呼び出し
    story.append(Spacer(1, 5 * mm))

    story.append(Paragraph("校正方針（テーブル）", styles["H1"]))
    story.append(Spacer(1, 1.5 * mm))
    if len(table_para) > 1:
        story.append(tbl)
    else:
        story.append(Paragraph("(解析表の検出に失敗したか、データがありません)", styles["Body"]))

    doc.build(story, onFirstPage=_page_number, onLaterPages=_page_number)
    pdf_bytes = buf.getvalue()
    buf.close()
    return pdf_bytes



# ------------------------------------------------------------
# 解析（校正方針だけ）
# ------------------------------------------------------------
def analyze_issues(model: str, lines: List[str], lines_per_page: int, mode: str, extra: str) -> Tuple[str, ChatUsage]:
    client = openai_client()
    md_tables: List[str] = []
    used_in = used_out = 0

    total_pages = (len(lines) + lines_per_page - 1) // lines_per_page
    sys_inst_template = build_sys_inst(get_analyze_instruction(mode), extra)

    for pg in range(total_pages):
        start = pg * lines_per_page
        end = min((pg + 1) * lines_per_page, len(lines))
        page_chunk = [f"[{(i//lines_per_page)+1}:{(i%lines_per_page)+1:02d}] {lines[i]}" for i in range(start, end)]
        page_text = "\n".join(page_chunk)

        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": sys_inst_template},
                {"role": "user", "content": f"次のテキスト（このページのみ）を解析してください：\n---\n{page_text}"},
            ],
        )
        md_tables.append(resp.choices[0].message.content.strip())
        pin, pout = _add_usage_from_resp(resp)
        used_in += pin; used_out += pout

    out = []
    for i, tbl in enumerate(md_tables, 1):
        out.append(f"#### 頁 {i}\n\n{tbl}\n")
    return "\n".join(out), ChatUsage(input_tokens=used_in, output_tokens=used_out)

# ------------------------------------------------------------
# 画面（解析のみ）
# ------------------------------------------------------------
st.title("📝 解析（校正方針の抽出）")
st.write("本文を入力して **① 解析** を実行すると、ページ/行/理由つきの校正方針（Markdown表）を生成します。")

with st.sidebar:
    st.header("設定")
    st.radio(
        "🧠 使用モデル",
        MODEL_OPTIONS,
        index=MODEL_OPTIONS.index(st.session_state["chat_model"]),
        key="chat_model",
    )

    sel_model = st.session_state["chat_model"]
    if sel_model in MODEL_PRICES_USD:
        price_in = MODEL_PRICES_USD[sel_model]["in"] / 1000
        price_out = MODEL_PRICES_USD[sel_model]["out"] / 1000
        st.markdown(
            f"**💲 単価 (USD / 1K tokens)**  \n- 入力: `${price_in:.5f}`  \n- 出力: `${price_out:.5f}`"
        )
    else:
        st.info("単価情報が未登録のモデルです。")

    st.selectbox(
        "🛠 解析モード",
        list(MODE_DEFS.keys()),
        index=list(MODE_DEFS.keys()).index(st.session_state.get("proof_mode", DEFAULT_MODE)),
        key="proof_mode",
        help="\n\n".join([f"・{k}: {v['desc']}" for k, v in MODE_DEFS.items()]),
    )
    st.caption(f"説明：{MODE_DEFS[st.session_state['proof_mode']]['desc']}")

    lpp = st.number_input("ページ行数（表示用）", min_value=20, max_value=100, value=LINES_PER_PAGE, step=5)

    # 解析結果のダウンロード形式（内部キーで比較・表示テキストはformat_func）
    _DL_LABELS = {"word": "Word (.docx)", "pdf": "PDF (.pdf)", "both": "両方"}
    dl_choice_key = st.radio(
        "📦 ダウンロード形式（解析結果のレポート）",
        options=list(_DL_LABELS.keys()),
        format_func=lambda k: _DL_LABELS[k],
        index=1,  # デフォルトは Word
        key="dl_format_radio",
        help="解析レポート（原文プレビュー＋校正方針）を保存する形式です。",
    )

extra_prompt = render_policy_preview(mode=st.session_state["proof_mode"])
st.markdown("---")

# ===== 入力（ファイル / 貼り付け） =====
tab_file, tab_paste = st.tabs(["📁 ファイルから", "📝 貼り付けテキスト"])

src_text = ""
used_file_name = None

with tab_file:
    col_u, col_btn1 = st.columns([3, 1])
    with col_u:
        up = st.file_uploader(".docx / .txt / .pdf をアップロード", type=["docx", "txt", "pdf"])
    with col_btn1:
        do_analyze_file = st.button("① 解析（ファイル）", type="primary", use_container_width=True, disabled=not up, key="btn_analyze_file")

    if up:
        used_file_name = up.name
        name = up.name.lower()
        if name.endswith(".pdf"):
            data = up.read()
            try:
                stats = extract_pdf_text(data)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

            st.subheader("📄 PDFプレビュー")
            display_pdf_bytes(data, height=600)

            if int(stats.get("visible", 0)) < 20:
                st.warning("このPDFは画像PDF（テキスト層なし）と判定しました。OCRツールでテキスト化してから再度お試しください。")
                st.stop()
            else:
                src_text = (stats.get("text") or "").strip()
        else:
            try:
                src_text = load_text_generic(up)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

with tab_paste:
    pasted = st.text_area(
        "ここに本文を貼り付け",
        height=260,
        placeholder="ここに本文を貼り付けてください（改行は保持されます）。"
    )
    do_analyze_paste = st.button("① 解析（貼り付け）", type="secondary", use_container_width=True, disabled=not pasted.strip(), key="btn_analyze_paste")
    if pasted:
        src_text = load_text_from_paste(
            pasted,
            normalize=True,
            collapse_blanks=False,
            keep_blank_lines=1,
            trim_trailing=True,
        )
        used_file_name = "pasted_text.txt"

# ===== 共通：解析の実行 =====
if src_text:
    lines = to_numbered_lines(src_text)
    st.subheader("👀 行番号付きプレビュー")
    st.caption(f"表示上のページ行数: {LINES_PER_PAGE} 行/ページ（擬似割り付け）")
    st.text_area("原文（番号付きプレビュー）", value=render_preview_with_numbers(lines, LINES_PER_PAGE), height=260)

    run_in = run_out = 0
    want_analyze = (locals().get("do_analyze_file") or locals().get("do_analyze_paste"))

    if want_analyze:
        with st.spinner("解析中（校正方針を抽出）…"):
            plan_md, usage = analyze_issues(
                st.session_state["chat_model"], lines, LINES_PER_PAGE,
                mode=st.session_state["proof_mode"], extra=extra_prompt
            )
        st.success("解析が完了しました。ページ/行/理由つきで方針を表示します。")
        st.subheader("📋 校正方針（まず何をどう直すか）")
        st.markdown(plan_md, unsafe_allow_html=False)
        run_in += usage.input_tokens; run_out += usage.output_tokens

        # ▼▼ 解析レポートのダウンロード（Word / PDF / 両方） ▼▼
        st.markdown("### ⤵️ 解析レポートをダウンロード")
        numbered_preview = render_preview_with_numbers(lines, LINES_PER_PAGE)
        file_base = (used_file_name or "pasted_text").rsplit(".", 1)[0]
        file_stub = f"policy_{file_base}"

        # Word
        if dl_choice_key in ("word", "both"):
            data_docx, ext = build_policy_docx_bytes(
                original_numbered_preview=numbered_preview,
                plan_md=plan_md,
                model=st.session_state["chat_model"],
                mode=st.session_state["proof_mode"],
                extra_prompt=extra_prompt,
                src_name=used_file_name or "pasted_text.txt",
            )
            st.download_button(
                "Word（.docx）として保存" if ext == ".docx" else "テキスト（.txt）として保存",
                data=data_docx,
                file_name=f"{file_stub}{ext}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == ".docx" else "text/plain",
                use_container_width=True,
                key=f"dl_word_{file_stub}",
            )
            if ext != ".docx":
                st.caption("※ `python-docx` が無い環境では .txt で保存されます。")

        # PDF
        if dl_choice_key in ("pdf", "both"):
            pdf_bytes = build_policy_pdf_bytes(
                original_numbered_preview=numbered_preview,
                plan_md=plan_md,
                model=st.session_state["chat_model"],
                mode=st.session_state["proof_mode"],
                extra_prompt=extra_prompt,
                src_name=used_file_name or "pasted_text.txt",
            )
            if pdf_bytes:
                st.download_button(
                    "PDF（.pdf）として保存",
                    data=pdf_bytes,
                    file_name=f"{file_stub}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"dl_pdf_{file_stub}",
                )
            else:
                st.warning("PDF を生成できませんでした。`pip install reportlab` を実行し、CIDフォント（HeiseiMin-W3/HeiseiKakuGo-W5）が使えるか確認してください。")

        # 使用量と費用
        st.markdown("## 💰 使用量と概算費用")
        render_chat_only_summary(
            title="今回ランの概算",
            model=st.session_state["chat_model"],
            in_tokens=run_in,
            out_tokens=run_out,
        )
        usage_cost = estimate_chat_cost(
            st.session_state["chat_model"],
            ChatUsage(input_tokens=run_in, output_tokens=run_out),
        )
        st.caption(
            f"今回ラン tokens: in={run_in:,} / out={run_out:,} → "
            f"USD ${usage_cost['usd']:.6f} ≈ JPY {usage_cost['jpy']:.2f} "
            f"(rate={DEFAULT_USDJPY:.2f})"
        )

        # セッション累計
        st.session_state["tok_in_total"] += run_in
        st.session_state["tok_out_total"] += run_out
        st.divider()
        st.markdown("### セッション累計（このページでの合算）")
        render_chat_only_summary(
            title="累計の概算",
            model=st.session_state["chat_model"],
            in_tokens=st.session_state["tok_in_total"],
            out_tokens=st.session_state["tok_out_total"],
        )

else:
    st.info("入力タブ（📁/📝）から本文を指定して『① 解析』を実行してください。")
