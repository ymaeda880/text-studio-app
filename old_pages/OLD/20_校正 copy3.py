# pages/20_校正.py — 解析（校正方針：ページ/行/理由）→ 本校正
# （PDFが画像PDFなら警告して停止／OCR機能は削除）
from __future__ import annotations
from typing import List, Tuple, Dict, Iterable
from pathlib import Path
import sys
import io
import datetime as _dt

import streamlit as st
from openai import OpenAI

# ===== 共有ライブラリ（common_lib）をパスに追加 =====
PROJECTS_ROOT = Path(__file__).resolve().parents[3]  # or 3 for pages
if str(PROJECTS_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECTS_ROOT))

# 料金見積りユーティリティ
from common_lib.openai.costs import (
    ChatUsage,
    estimate_chat_cost,
    render_chat_only_summary,
    DEFAULT_USDJPY,
)

# === lib の利用（読込 / 貼り付け整形） ※OCRなし ===
from lib.text_loaders import (
    load_text_generic, extract_pdf_text, load_text_from_paste
)

# ------------------------------------------------------------
# モード定義
# ------------------------------------------------------------

MODE_DEFS: Dict[str, Dict[str, str]] = {
    "厳格校正": {
        "desc": "助詞・主述一致・冗長/重複・語順・誤字脱字・用語誤用・文体不統一まで広く対象。意味は変えず最適化。",
        "analyze_inst": (
            "あなたは厳密な日本語校正リーダーです。以下の番号付きテキストを読み、"
            "『何をどのように直すべきか』を、具体的な理由とともに一覧化してください。\n"
            "行頭の [page:line] を必ず参照して位置を示し、過度な意訳は避けてください。\n"
            "助詞（てにおは）、主述一致、冗長、重複、語順、誤字脱字、用語誤用、文体の不統一に特に注意。\n"
            "「原文」 は該当箇所の短い抜粋（最大20字）に留めてください。"
        ),
        "proofread_inst": (
            "あなたは厳密な日本語校正者です。以下を徹底してください：\n"
            "- てにおは、助詞、主述一致、誤字脱字、表記ゆれを修正\n"
            "- 用語誤用の是正・語順の自然化（意味は変えない）\n"
            "- 冗長・重複を圧縮し、文体を統一（尊敬/丁寧/常体の混在を解消）\n"
            "- 出力は校正後の本文のみ（前置き不要）"
        ),
    },
    "簡易校正（ミス最小修正）": {
        "desc": "明白なミスのみ（てにおは・助詞・誤字脱字・明確な変換ミス）。語順や言い換えは極力しない。",
        "analyze_inst": (
            "あなたは日本語の軽微校正リーダーです。以下の番号付きテキストから、"
            "『明白なミス（てにおは・助詞の誤り、誤字脱字、明確な変換ミス）』のみを抽出してください。"
            "語順変更やスタイル統一などの裁量的変更は提案しないでください。"
            "「原文」 は短い抜粋（最大20字）。「理由」 は“明白な誤り”である根拠を簡潔に。"
        ),
        "proofread_inst": (
            "あなたは日本語の軽微校正者です。以下を徹底してください：\n"
            "- 明白なミスのみ修正（てにおは・助詞・誤字脱字・明確な変換ミス）\n"
            "- 語順の大きな変更や言い換えは避ける（意味・文体は最大限維持）\n"
            "- 出力は校正後の本文のみ（前置き不要）"
        ),
    },
}

# ------------------------------------------------------------
# 両モード共通で常に付与されるプロンプト（ここに固定記述）
# ------------------------------------------------------------
COMMON_PROMPT = (
    "【共通方針（厳守）】\n"
    "それぞれの校正の必要性を0から10までの数字で表して明記してください：\n"
    "- 必要性が高い構成には大きい数字をつけてください．\n"
    "- 誤字．脱字など，校正が必ず必要なものを10と評価してください．\n"
    "- 校正する必要の無いものを0と評価してください．"
    "- 出力は **Markdownの表** で、列は次の順：頁 | 行 | 重要度  | 原文 | 修正案 | 理由  "
)

# ------------------------------------------------------------
# UI定数
# ------------------------------------------------------------
st.set_page_config(page_title="Text Studio / 校正", page_icon="📝", layout="wide")

MODEL_OPTIONS = ["gpt-5-mini", "gpt-5-nano"]
DEFAULT_MODEL = "gpt-5-mini"
DEFAULT_MODE = "厳格校正"
LINES_PER_PAGE = 40  # 1ページあたりの表示行数

# セッション初期化
if "chat_model" not in st.session_state:
    st.session_state["chat_model"] = DEFAULT_MODEL
if "proof_mode" not in st.session_state:
    st.session_state["proof_mode"] = DEFAULT_MODE
if "tok_in_total" not in st.session_state:
    st.session_state["tok_in_total"] = 0
if "tok_out_total" not in st.session_state:
    st.session_state["tok_out_total"] = 0

# ------------------------------------------------------------
# 表示ユーティリティ
# ------------------------------------------------------------
def display_pdf_bytes(data: bytes, height: int = 600):
    """Streamlit PDF表示（streamlit[pdf] があれば st.pdf）。なければ iframe 埋め込み。"""
    try:
        st.pdf(data, height=height)  # Streamlit 1.31+ / streamlit[pdf]
    except Exception:
        import base64
        b64 = base64.b64encode(data).decode()
        st.markdown(
            f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="{height}px"></iframe>',
            unsafe_allow_html=True
        )

def to_numbered_lines(raw: str) -> List[str]:
    return raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")

def page_and_line(idx: int, lines_per_page: int) -> Tuple[int, int]:
    page = idx // lines_per_page + 1
    line_in_page = idx % lines_per_page + 1
    return page, line_in_page

def render_preview_with_numbers(lines: List[str], lines_per_page: int) -> str:
    return "\n".join(f"[{(i//lines_per_page)+1}:{(i%lines_per_page)+1:02d}] {t}" for i, t in enumerate(lines))

# ------------------------------------------------------------
# OpenAIクライアントと使用量抽出
# ------------------------------------------------------------
def openai_client() -> OpenAI:
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

def _add_usage_from_resp(resp) -> Tuple[int, int]:
    u = getattr(resp, "usage", None) or {}
    prompt = getattr(u, "prompt_tokens", None) if hasattr(u, "prompt_tokens") else u.get("prompt_tokens", 0)
    completion = getattr(u, "completion_tokens", None) if hasattr(u, "completion_tokens") else u.get("completion_tokens", 0)
    try:
        prompt = int(prompt or 0)
    except Exception:
        prompt = 0
    try:
        completion = int(completion or 0)
    except Exception:
        completion = 0
    return prompt, completion

# ------------------------------------------------------------
# 追加：方針プレビュー＆追加プロンプト
# ------------------------------------------------------------
def get_analyze_instruction(mode: str) -> str:
    return MODE_DEFS.get(mode, MODE_DEFS["厳格校正"])["analyze_inst"]

def get_proofread_instruction(mode: str, keep_layout: bool) -> str:
    base = MODE_DEFS.get(mode, MODE_DEFS["厳格校正"])["proofread_inst"]
    if keep_layout:
        base = base + "\n- 改行・段落は可能な限り維持"
    return base

def build_sys_inst(base: str, extra: str) -> str:
    """
    base（モード固有） → extra（UI入力） → COMMON_PROMPT（共通指示）
    の順で結合して、1本の System プロンプトにする。
    """
    parts = [base.strip()]
    if (extra or "").strip():
        parts.append("【追加指示（厳守）】\n" + extra.strip())
    if (globals().get("COMMON_PROMPT") or "").strip():
        parts.append(COMMON_PROMPT.strip())
    return "\n\n".join(parts)

def render_policy_preview(*, mode: str, keep_layout: bool) -> str:
    analyze_base = get_analyze_instruction(mode)
    proofread_base = get_proofread_instruction(mode, keep_layout=keep_layout)

    st.subheader("🧭 校正方針の確認")
    st.caption(MODE_DEFS[mode]["desc"])
    
    # ↓↓ 解析プロンプト（System） ↓↓
    with st.expander("解析で使うプロンプト（System）", expanded=False):
        st.code(analyze_base, language="markdown")

    # ↓↓ 本校正プロンプト（System） ↓↓
    with st.expander("本校正で使うプロンプト（System）", expanded=False):
        st.code(proofread_base, language="markdown")

     # ↓↓ 共通のプロンプト（常時付与） ↓↓
    with st.expander("共通方針（Systemに毎回付与）", expanded=False):
        st.code(COMMON_PROMPT.strip(), language="markdown")

    st.markdown("**✍️ 追加プロンプト（任意）**")
    st.caption("特記事項（例：外来語はカタカナ優先／製品名は原文どおり／全角英数字は半角に など）を追記できます。")

    extra = st.text_area(
        "追加プロンプト",
        key="extra_user_prompt",
        placeholder="例）外来語の表記はカタカナ優先。製品名や固有名詞は原文どおりに保持。",
        height=100,
    )

    return extra or ""

# ------------------------------------------------------------
# Word（docx）生成ユーティリティ（原文は行番号つき／方針は整形表示）
# ------------------------------------------------------------
def _parse_plan_md_tables(md: str) -> List[Dict[str, str]]:
    """
    校正方針（Markdownの表）をパースして、各行を dict 化して返す。
    期待カラム: 頁 | 行 | 重要度 | 原文 | 修正案 | 理由
    ※ 英語/日本語ヘッダを受け付け、返却は日本語キーに統一する。
    """
    items: List[Dict[str, str]] = []
    if not md:
        return items

    # 英->日 の正規化マップ
    to_jp = {
        "page": "頁",
        "line": "行",
        "issue": "重要度",
        "original": "原文",
        "suggestion": "修正案",
        "reason": "理由",
        "頁": "頁", "ページ": "頁",
        "行": "行",
        "重要度": "重要度", "問題点": "重要度",
        "原文": "原文",
        "修正案": "修正案", "提案": "修正案",
        "理由": "理由", "根拠": "理由",
    }

    expected = ["頁", "行", "重要度", "原文", "修正案", "理由"]

    def _norm_head_cell(s: str) -> str:
        key = s.strip().lower()
        return to_jp.get(key, to_jp.get(s.strip(), s.strip()))

    def _row_to_dict(ln: str, cols_jp: List[str]) -> Dict[str, str] | None:
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if len(cells) < len(cols_jp):
            return None
        row = dict(zip(cols_jp, cells[:len(cols_jp)]))
        for k in expected:
            row.setdefault(k, "")
        return row

    lines = [ln for ln in (l.strip() for l in md.splitlines()) if ln]
    header_seen = False
    cols_jp: List[str] = []

    for ln in lines:
        if ln.startswith("|") and ("|" in ln):
            raw_cells = [c.strip() for c in ln.strip("|").split("|")]

            # 区切り行はスキップ
            if all(set(c.replace("-", "")) == set() for c in raw_cells):
                continue

            if not header_seen:
                norm = [_norm_head_cell(c) for c in raw_cells]
                if len(set(norm) & set(expected)) >= 3:
                    cols_jp = [c if c in expected else "" for c in norm]
                    if len(cols_jp) < 6:
                        cols_jp += [""] * (6 - len(cols_jp))
                    cols_jp = [cols_jp[i] or expected[i] for i in range(6)]
                    header_seen = True
                    continue

            if header_seen and cols_jp:
                row = _row_to_dict(ln, cols_jp)
                if row:
                    items.append(row)

    return items

def build_policy_docx_bytes(
    *,
    original_numbered_preview: str,
    plan_md: str,
    model: str,
    mode: str,
    extra_prompt: str,
    src_name: str
) -> Tuple[bytes, str]:
    """
    python-docx が存在すれば .docx を、無ければ .txt を返す。
    戻り値: (ファイルバイト列, 拡張子 ".docx" or ".txt")
    """
    plan_items_raw = _parse_plan_md_tables(plan_md)

    def _clean(s: str) -> str:
        s = (s or "").strip()
        if set(s) <= set("-—:・ 　"):
            return ""
        return s

    def _is_meaningful(row: Dict[str, str]) -> bool:
        keys = ["頁", "行", "原文", "修正案"]
        return any(_clean(row.get(k, "")) for k in keys)

    plan_items = []
    for it in plan_items_raw:
        row = {
            "頁": _clean(it.get("頁", "")),
            "行": _clean(it.get("行", "")),
            "重要度": _clean(it.get("重要度", "")),
            "原文": _clean(it.get("原文", "")),
            "修正案": _clean(it.get("修正案", "")),
            "理由": _clean(it.get("理由", "")),
        }
        if _is_meaningful(row):
            plan_items.append(row)

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # タイトル
        title = doc.add_heading("校正方針レポート", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # メタ情報
        p = doc.add_paragraph()
        p.add_run("生成日時: ").bold = True
        p.add_run(_dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        p.add_run("\nモデル: ").bold = True
        p.add_run(model)
        p.add_run("\nモード: ").bold = True
        p.add_run(mode)
        if (extra_prompt or "").strip():
            p.add_run("\n追加プロンプト: ").bold = True
            p.add_run(extra_prompt.strip())
        p.add_run("\n入力ファイル: ").bold = True
        p.add_run(src_name or "-")

        # 原文（行番号つき）
        doc.add_heading("原文（行番号つき）", level=1)
        para = doc.add_paragraph()
        run = para.add_run(original_numbered_preview if original_numbered_preview else "(空)")
        run.font.name = "Courier New"
        run.font.size = Pt(10)

        # 校正方針（整形版）
        doc.add_heading("校正方針（読みやすい整形版）", level=1)
        if plan_items:
            for k, it in enumerate(plan_items, 1):
                if k > 1:
                    doc.add_paragraph("────────────────────────────────")

                head = doc.add_paragraph()
                head_run = head.add_run(
                    f"{k}. 頁 {it.get('頁','-')}  行 {it.get('行','-')} — 重要度 {it.get('重要度','')}"
                )
                head_run.bold = True

                if it.get("原文"):
                    p1 = doc.add_paragraph()
                    p1.add_run("原文: ").bold = True
                    p1.add_run(it.get("原文", ""))

                if it.get("修正案"):
                    p2 = doc.add_paragraph()
                    p2.add_run("修正案: ").bold = True
                    p2.add_run(it.get("修正案", ""))

                if it.get("理由"):
                    p3 = doc.add_paragraph()
                    p3.add_run("理由: ").bold = True
                    p3.add_run(it.get("理由", ""))
        else:
            doc.add_paragraph("(解析表の検出に失敗したか、データがありません)")

        # 付録：元Markdown（任意）
        doc.add_heading("付録：元Markdown（そのまま）", level=2)
        para2 = doc.add_paragraph()
        r2 = para2.add_run(plan_md if plan_md else "(なし)")
        r2.font.name = "Courier New"
        r2.font.size = Pt(10)

        bio = io.BytesIO()
        doc.save(bio)
        return bio.getvalue(), ".docx"

    except Exception:
        # フォールバック: .txt
        lines: List[str] = []
        lines.append("=== 校正方針レポート（TXTフォールバック） ===")
        lines.append(f"生成日時: {_dt.datetime.now():%Y-%m-%d %H:%M:%S}")
        lines.append(f"モデル: {model}")
        lines.append(f"モード: {mode}")
        if (extra_prompt or "").strip():
            lines.append(f"追加プロンプト: {extra_prompt.strip()}")
        lines.append(f"入力ファイル: {src_name or '-'}")

        lines.append("\n--- 原文（行番号つき） ---\n")
        lines.append(original_numbered_preview or "(空)")

        lines.append("\n--- 校正方針（整形版） ---\n")
        if plan_md and plan_items:
            for i, it in enumerate(plan_items, 1):
                if i > 1:
                    lines.append("────────────────────────────────")
                lines.append(
                    f"{i}. 頁 {it.get('頁','-')}  行 {it.get('行','-')} — 重要度 {it.get('重要度','')}"
                )
                if it.get("原文"):
                    lines.append(f"   原文   : {it.get('原文','')}")
                if it.get("修正案"):
                    lines.append(f"   修正案 : {it.get('修正案','')}")
                if it.get("理由"):
                    lines.append(f"   理由   : {it.get('理由','')}")
        else:
            lines.append("(解析表の検出に失敗)")

        lines.append("\n--- 付録：元Markdown ---\n")
        lines.append(plan_md or "(なし)")
        data = "\n".join(lines).encode("utf-8")
        return data, ".txt"

# ------------------------------------------------------------
# 追加：校正結果の PDF/Word 生成ヘルパ
# ------------------------------------------------------------
def build_plain_docx_bytes(text: str, title: str = "校正結果") -> bytes:
    """校正結果テキストを .docx に単純出力"""
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading(title, 0)
    p = doc.add_paragraph()
    run = p.add_run(text if text else "(空)")
    run.font.size = Pt(11)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

def build_pdf_from_text(text: str, title: str = "校正結果") -> bytes:
    """
    日本語対応の簡易PDF生成。reportlab がある場合のみ。
    無い場合は空bytesを返す（UI側で案内）。
    """
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import mm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont

        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        width, height = A4

        # 日本語フォント（CIDフォント）
        pdfmetrics.registerFont(UnicodeCIDFont('HeiseiMin-W3'))

        # タイトル
        c.setFont('HeiseiMin-W3', 14)
        c.drawString(20*mm, height - 20*mm, title)

        # 本文
        textobj = c.beginText(20*mm, height - 30*mm)
        textobj.setFont('HeiseiMin-W3', 10)

        for line in (text or "").splitlines():
            textobj.textLine(line if line else " ")
            # 余白確保して改ページ
            if textobj.getY() < 20*mm:
                c.drawText(textobj)
                c.showPage()
                c.setFont('HeiseiMin-W3', 14)
                c.drawString(20*mm, height - 20*mm, title)
                textobj = c.beginText(20*mm, height - 30*mm)
                textobj.setFont('HeiseiMin-W3', 10)

        c.drawText(textobj)
        c.showPage()
        c.save()
        pdf_bytes = buf.getvalue()
        buf.close()
        return pdf_bytes
    except Exception:
        return b""

# ------------------------------------------------------------
# 解析（校正方針）
# ------------------------------------------------------------
def analyze_issues(model: str, lines: List[str], lines_per_page: int, mode: str, extra: str) -> Tuple[str, ChatUsage]:
    client = openai_client()
    md_tables: List[str] = []
    used_in = used_out = 0

    total_pages = (len(lines) + lines_per_page - 1) // lines_per_page
    sys_inst_template = build_sys_inst(get_analyze_instruction(mode), extra)

    for pg in range(total_pages):
        start = pg * lines_per_page
        end = min((pg + 1) * lines_per_page, len(lines))
        page_chunk = [f"[{(i//lines_per_page)+1}:{(i%lines_per_page)+1:02d}] {lines[i]}" for i in range(start, end)]
        page_text = "\n".join(page_chunk)

        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": sys_inst_template},
                {"role": "user", "content": f"次のテキスト（このページのみ）を解析してください：\n---\n{page_text}"},
            ],
        )
        md_tables.append(resp.choices[0].message.content.strip())
        pin, pout = _add_usage_from_resp(resp)
        used_in += pin; used_out += pout

    out = []
    for i, tbl in enumerate(md_tables, 1):
        out.append(f"#### 頁 {i}\n\n{tbl}\n")
    return "\n".join(out), ChatUsage(input_tokens=used_in, output_tokens=used_out)

# ------------------------------------------------------------
# 本校正
# ------------------------------------------------------------
def proofread(model: str, content: str, keep_layout: bool, want_report: bool, mode: str, extra: str) -> Tuple[str, ChatUsage]:
    client = openai_client()
    CHUNK = 6000
    chunks = [content[i:i+CHUNK] for i in range(0, len(content), CHUNK)] or [content]

    fixed_parts: List[str] = []
    used_in = used_out = 0
    sys_inst = build_sys_inst(get_proofread_instruction(mode, keep_layout=keep_layout), extra)

    for chunk in chunks:
        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": sys_inst},
                {"role": "user", "content": f"次の本文を校正してください：\n---\n{chunk}"},
            ],
        )
        fixed_parts.append(resp.choices[0].message.content.strip())
        pin, pout = _add_usage_from_resp(resp)
        used_in += pin; used_out += pout

    fixed_text = ("\n\n" if keep_layout else "\n").join(fixed_parts).strip()

    if want_report:
        rep_inst = (
            "次の原文と校正後本文の差分観点で、主な修正ポイントを最大8項目で箇条書きに。\n"
            "例：助詞/活用、語順、冗長/重複、誤字脱字、文体統一、用語誤用など。"
        )
        rep_user = f"原文:\n{content[:4000]}\n\n校正後:\n{fixed_text[:4000]}"
        rep = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": rep_inst},
                {"role": "user", "content": rep_user},
            ],
        )
        fixed_text += f"\n\n---\n【修正ポイント（要約）】\n{rep.choices[0].message.content.strip()}"
        pin, pout = _add_usage_from_resp(rep)
        used_in += pin; used_out += pout

    return fixed_text, ChatUsage(input_tokens=used_in, output_tokens=used_out)

# ------------------------------------------------------------
# 画面
# ------------------------------------------------------------
st.title("📝 校正 — 解析（ページ/行/理由） → 本校正")
st.write("最初に**方針**を確認/追記してから、入力（ファイル or 貼り付け）を選んでください。")

with st.sidebar:
    st.header("設定")
    st.radio(
        "🧠 使用モデル",
        MODEL_OPTIONS,
        index=MODEL_OPTIONS.index(st.session_state["chat_model"]),
        key="chat_model",
    )

    from common_lib.openai.costs import MODEL_PRICES_USD
    sel_model = st.session_state["chat_model"]
    if sel_model in MODEL_PRICES_USD:
        price_in = MODEL_PRICES_USD[sel_model]["in"] / 1000
        price_out = MODEL_PRICES_USD[sel_model]["out"] / 1000
        st.markdown(
            f"""
            **💲 単価 (USD / 1K tokens)**  
            - 入力: `${price_in:.5f}`  
            - 出力: `${price_out:.5f}`
            """
        )
    else:
        st.info("単価情報が未登録のモデルです。")

    st.selectbox(
        "🛠 校正モード",
        list(MODE_DEFS.keys()),
        index=list(MODE_DEFS.keys()).index(st.session_state.get("proof_mode", DEFAULT_MODE)),
        key="proof_mode",
        help="\n\n".join([f"・{k}: {v['desc']}" for k, v in MODE_DEFS.items()]),
    )
    st.caption(f"説明：{MODE_DEFS[st.session_state['proof_mode']]['desc']}")

    keep_formatting = st.checkbox("改行・段落を保持（本校正時）", value=True)
    show_report = st.checkbox("本校正後に『修正ポイント要約』も付与", value=False)
    lpp = st.number_input("ページ行数（表示用）", min_value=20, max_value=100, value=LINES_PER_PAGE, step=5)

    # 追加：校正結果のダウンロード形式
    dl_choice = st.radio(
        "📦 ダウンロード形式（校正結果）",
        ("Word (.docx)", "PDF (.pdf)", "両方"),
        index=0,  # デフォルトは Word
        help="校正結果のダウンロード形式を選びます。方針レポートの Word 出力は従来どおり固定です。"
    )

extra_prompt = render_policy_preview(
    mode=st.session_state["proof_mode"],
    keep_layout=keep_formatting,
)
st.markdown("---")

# ===== 入力方法の選択（ファイル / 貼り付け） =====
tab_file, tab_paste = st.tabs(["📁 ファイルから", "📝 貼り付けテキスト"])

src_text = ""
used_file_name = None

# ---------- ファイル入力 ----------
with tab_file:
    col_u, col_btn1, col_btn2 = st.columns([3, 1, 1])
    with col_u:
        up = st.file_uploader(".docx / .txt / .pdf をアップロード", type=["docx", "txt", "pdf"])
    with col_btn1:
        do_analyze_file = st.button("① 解析（ファイル）", type="primary", use_container_width=True, disabled=not up, key="btn_analyze_file")
    with col_btn2:
        do_fix_file = st.button("② 本校正（ファイル）", type="primary", use_container_width=True, disabled=not up, key="btn_fix_file")

    if up:
        used_file_name = up.name
        name = up.name.lower()
        if name.endswith(".pdf"):
            data = up.read()
            try:
                stats = extract_pdf_text(data)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

            st.subheader("📄 PDFプレビュー")
            display_pdf_bytes(data, height=600)

            if int(stats.get("visible", 0)) < 20:
                st.warning("このPDFは画像PDF（テキスト層なし）と判定しました。OCRツールでテキスト化してから再度お試しください。")
                st.stop()
            else:
                src_text = (stats.get("text") or "").strip()
        else:
            try:
                src_text = load_text_generic(up)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

# ---------- 貼り付け入力 ----------
with tab_paste:
    pasted = st.text_area(
        "ここに本文を貼り付け",
        height=260,
        placeholder="ここに本文を貼り付けてください（改行は保持されます）。"
    )
    col_p1, col_p2, col_opts = st.columns([1, 1, 2])
    with col_opts:
        normalize = st.checkbox("改行とBOMを正規化", value=True, key="opt_norm")
        collapse = st.checkbox("連続空行を圧縮", value=False, key="opt_collapse")
        keep_blanks = st.number_input("空行の上限（圧縮時）", 1, 5, 1, key="opt_keep_blanks")
        trim_tail = st.checkbox("行末スペースを除去", value=True, key="opt_trim_tail")

    with col_p1:
        do_analyze_paste = st.button("① 解析（貼り付け）", type="secondary", use_container_width=True,
                                 disabled=not pasted.strip(), key="btn_analyze_paste")

    with col_p2:
        do_fix_paste = st.button("② 本校正（貼り付け）", type="primary", use_container_width=True,
                             disabled=not pasted.strip(), key="btn_fix_paste")

    if pasted:
        src_text = load_text_from_paste(
            pasted,
            normalize=normalize,
            collapse_blanks=collapse,
            keep_blank_lines=int(keep_blanks),
            trim_trailing=trim_tail,
        )
        used_file_name = "pasted_text.txt"

# ===== 共通：ここから解析/本校正の実行（ファイル or 貼り付け） =====
if src_text:
    lines = to_numbered_lines(src_text)
    st.subheader("👀 行番号付きプレビュー")
    st.caption(f"表示上のページ行数: {LINES_PER_PAGE} 行/ページ（擬似割り付け）")
    st.text_area("原文（番号付きプレビュー）", value=render_preview_with_numbers(lines, LINES_PER_PAGE), height=260)

    run_in = run_out = 0
    want_analyze = (locals().get("do_analyze_file") or locals().get("do_analyze_paste"))
    want_fix     = (locals().get("do_fix_file") or locals().get("do_fix_paste"))

    if want_analyze:
        with st.spinner("解析中（校正方針を抽出）…"):
            plan_md, usage = analyze_issues(
                st.session_state["chat_model"], lines, LINES_PER_PAGE,
                mode=st.session_state["proof_mode"], extra=extra_prompt
            )
        st.success("解析が完了しました。ページ/行/理由つきで方針を表示します。")
        st.subheader("📋 校正方針（まず何をどう直すか）")
        st.markdown(plan_md, unsafe_allow_html=False)
        run_in += usage.input_tokens; run_out += usage.output_tokens

        # ▼▼ Word（docx）ダウンロード（原文=行番号つき、方針=整形） ▼▼
        st.markdown("### ⤵️ 原文＋校正方針を Word でダウンロード")
        numbered_preview = render_preview_with_numbers(lines, LINES_PER_PAGE)
        data, ext = build_policy_docx_bytes(
            original_numbered_preview=numbered_preview,
            plan_md=plan_md,
            model=st.session_state["chat_model"],
            mode=st.session_state["proof_mode"],
            extra_prompt=extra_prompt,
            src_name=used_file_name or "pasted_text.txt",
        )
        file_base = (used_file_name or "pasted_text").rsplit(".", 1)[0]
        dl_name = f"policy_{file_base}{ext}"
        st.download_button(
            "Word（.docx）として保存" if ext == ".docx" else "テキスト（.txt）として保存",
            data=data,
            file_name=dl_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == ".docx" else "text/plain",
            use_container_width=True,
        )
        st.caption("※ 方針レポートは Word（python-docx）固定です。python-docx が無い環境では .txt で保存されます。")

    if want_fix:
        with st.spinner("本校正の実行中…"):
            result, usage = proofread(
                model=st.session_state["chat_model"],
                content=src_text,
                keep_layout=keep_formatting,
                want_report=show_report,
                mode=st.session_state["proof_mode"],
                extra=extra_prompt,
            )
        st.success("校正完了！下の結果をコピー/ダウンロードできます。")
        st.subheader("🧾 校正結果")
        st.text_area("校正結果", value=result, height=420)

        file_stub = f"proofread_{(used_file_name or 'output').rsplit('.',1)[0]}"

        # ▼▼ 校正結果のダウンロード：サイドバーの選択に従う ▼▼
        if dl_choice in ("Word (.docx)", "両方"):
            try:
                docx_bytes = build_plain_docx_bytes(result, title="校正結果")
                st.download_button(
                    "校正結果を Word (.docx) で保存",
                    data=docx_bytes,
                    file_name=f"{file_stub}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.warning(f"Word 出力に失敗しました: {e}")

        if dl_choice in ("PDF (.pdf)", "両方"):
            pdf_bytes = build_pdf_from_text(result, title="校正結果")
            if pdf_bytes:
                st.download_button(
                    "校正結果を PDF (.pdf) で保存",
                    data=pdf_bytes,
                    file_name=f"{file_stub}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            else:
                st.warning("PDF 生成に必要な reportlab が見つかりませんでした。`pip install reportlab` を実行してください。")

        # 参考として従来の .txt も残す場合は下を有効化
        # st.download_button(
        #     "校正結果をテキスト (.txt) で保存",
        #     data=result.encode("utf-8"),
        #     file_name=f"{file_stub}.txt",
        #     mime="text/plain",
        # )

        run_in += usage.input_tokens; run_out += usage.output_tokens

    # セッション累計へ反映 & 料金表示
    if want_analyze or want_fix:
        st.session_state["tok_in_total"] += run_in
        st.session_state["tok_out_total"] += run_out

        st.markdown("## 💰 使用量と概算費用")
        render_chat_only_summary(
            title="今回ランの概算",
            model=st.session_state["chat_model"],
            in_tokens=run_in,
            out_tokens=run_out,
        )
        usage_cost = estimate_chat_cost(
            st.session_state["chat_model"],
            ChatUsage(input_tokens=run_in, output_tokens=run_out),
        )
        st.caption(
            f"今回ラン tokens: in={run_in:,} / out={run_out:,} → "
            f"USD ${usage_cost['usd']:.6f} ≈ JPY {usage_cost['jpy']:.2f} "
            f"(rate={DEFAULT_USDJPY:.2f})"
        )
        st.divider()
        st.markdown("### セッション累計（このページでの合算）")
        render_chat_only_summary(
            title="累計の概算",
            model=st.session_state["chat_model"],
            in_tokens=st.session_state["tok_in_total"],
            out_tokens=st.session_state["tok_out_total"],
        )
else:
    st.info("入力タブ（📁/📝）から本文を指定して『① 解析』→『② 本校正』の順に実行してください。")
