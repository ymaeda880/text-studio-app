# -*- coding: utf-8 -*-
# pages/26_è¡¨ä½œæˆ.py
#
# Excelè²¼ã‚Šä»˜ã‘ â†’ ã‚¹ã‚¿ã‚¤ãƒ«é¸æŠ â†’ Word .docx ç”Ÿæˆ
# ãƒ»ç”»é¢ãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼è¡¨ç¤ºï¼ˆrowspan/colspan å¯¾å¿œï¼‰
# ãƒ»è¡¨ä¸‹ã®æ³¨ãƒ†ã‚­ã‚¹ãƒˆ
# ãƒ»ã‚µã‚¤ãƒ‰ãƒãƒ¼ã«åŸºæœ¬ãƒ•ã‚©ãƒ³ãƒˆã‚µã‚¤ã‚º
# ãƒ»åˆ—å¹…ãƒ¢ãƒ¼ãƒ‰ï¼šå‡ç­‰ / è‡ªå‹•ï¼ˆæ–‡å­—æ•°ï¼‰ / æ‰‹å‹•ï¼ˆcmæŒ‡å®šï¼‰
# ãƒ»çµåˆæ–¹å¼ï¼šã‚»ãƒ«ã«ã€Œ<åŒä¸Šï¼ã€ã€Œ<åŒå·¦ï¼ã€ãŒæ›¸ã‹ã‚Œã¦ã„ã‚‹å ´åˆã®ã¿çµåˆ
#   - <åŒä¸Šï¼ ãªã‚‰ä¸Šã®éãƒãƒ¼ã‚«ãƒ¼ã‚»ãƒ«ã¨ç¸¦çµåˆï¼ˆé€£ç¶šã™ã‚Œã°ä¸Šç«¯ã‚¢ãƒ³ã‚«ãƒ¼ã«ä¼¸é•·ï¼‰
#   - <åŒå·¦ï¼ ãªã‚‰å·¦ã®éãƒãƒ¼ã‚«ãƒ¼ã‚»ãƒ«ã¨æ¨ªçµåˆï¼ˆé€£ç¶šã™ã‚Œã°å·¦ç«¯ã‚¢ãƒ³ã‚«ãƒ¼ã«ä¼¸é•·ï¼‰
#   - åŒå€¤è‡ªå‹•çµåˆã¯å»ƒæ­¢
# ãƒ»Word ã§ã¯ã‚»ãƒ«ã‚’ç‰©ç†çµåˆã—ãªã„ï¼ˆä¸­èº«ã‚’ç©ºã«ã—ã¦ç½«ç·šã¨å¡—ã‚Šã§è¦‹ã‹ã‘ä¸Šã®çµåˆï¼‰
# ãƒ»HTML ãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ã¯ rowspan/colspan ã§å®Ÿéš›ã«çµåˆ
# ãƒ»è¤‡æ•°è¡Œãƒ˜ãƒƒãƒ€ãƒ¼å¯¾å¿œï¼šå…ˆé ­ header_rows è¡Œã‚’ãƒ˜ãƒƒãƒ€ãƒ¼æ‰±ã„ã«ã—ã€ãã“ã§ã‚‚ <åŒä¸Š>/<åŒå·¦> ã‚’æœ‰åŠ¹

from __future__ import annotations

from io import BytesIO

import streamlit as st
import pandas as pd

from docx import Document
from docx.shared import Pt  # ç¾çŠ¶ã¯æœªä½¿ç”¨ã ãŒã€å°†æ¥ã®æ‹¡å¼µç”¨ã«æ®‹ã—ã¦ã‚‚OK

# ãƒ˜ãƒ«ãƒ‘é–¢æ•°ç¾¤ï¼ˆlib/table/helpers.pyï¼‰
from lib.table.helpers import (
    _parse_table as parse_table,
    _compute_col_widths_cm as compute_col_widths_cm,
    _widths_to_pct as widths_to_pct,
    _compute_spans_markers as compute_spans_markers,
    _merge_docx_by_spans as merge_docx_by_spans,
    _apply_docx_col_widths as apply_docx_col_widths,
    _apply_table_borders_robust as apply_table_borders_robust,
    _build_html_table_with_spans as build_html_table_with_spans,
    _apply_font_run as apply_font_run,
)

# ãƒ—ãƒªã‚»ãƒƒãƒˆãƒ»ã‚µãƒ³ãƒ—ãƒ«ï¼ˆlib/table/presets.pyï¼‰
from lib.table.presets import PRESETS, EXAMPLE_TEXT


# ============== UI ==============
st.set_page_config(page_title="ğŸ§¾ Word è¡¨ä½œæˆ", page_icon="ğŸ§¾", layout="wide")
st.title("ğŸ§¾ Word è¡¨ä½œæˆï¼ˆExcelè²¼ã‚Šä»˜ã‘ â†’ ã‚¹ã‚¿ã‚¤ãƒ«æŒ‡å®š â†’ .docx ï¼‹ ãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ï¼‰")
st.caption("1è¡Œç›®ã‹ã‚‰=ãƒ˜ãƒƒãƒ€ãƒ¼ã€ãƒ˜ãƒƒãƒ€ãƒ¼ã®å¾Œ=ãƒ‡ãƒ¼ã‚¿ã€‚Excelã‹ã‚‰ãã®ã¾ã¾è²¼ã‚Šä»˜ã‘ã¦ãã ã•ã„ã€‚")
st.caption("ãƒ˜ãƒƒãƒ€ãƒ¼ãŒï¼‘è¡Œä»¥ä¸Šã‚ã‚‹ã¨ãã¯ã€ãƒ˜ãƒƒãƒ€ãƒ¼è¡Œæ•°ã®æŒ‡å®šã‚’è¡Œã£ã¦ãã ã•ã„ã€‚")


with st.sidebar:
    st.markdown("### 1) åŸºæœ¬è¨­å®š")
    example = st.checkbox("ã‚µãƒ³ãƒ—ãƒ«ãƒ‡ãƒ¼ã‚¿ã‚’å…¥ã‚Œã‚‹", value=False)
    base_size = st.slider("åŸºæœ¬ãƒ•ã‚©ãƒ³ãƒˆã‚µã‚¤ã‚ºï¼ˆptï¼‰", 5, 16, 10, 1)
    header_same = st.checkbox("ãƒ˜ãƒƒãƒ€ãƒ¼ã‚‚åŒã˜ã‚µã‚¤ã‚ºï¼ˆOFFãªã‚‰ +1ptï¼‰", value=False)

    st.markdown("### 2) ãƒãƒ¼ã‚«ãƒ¼ã§ã‚»ãƒ«çµåˆ")
    use_up = st.checkbox("ã€Œ<åŒä¸Šï¼ã€ã§ç¸¦çµåˆã‚’æœ‰åŠ¹åŒ–", value=True)
    use_left = st.checkbox("ã€Œ<åŒå·¦ï¼ã€ã§æ¨ªçµåˆã‚’æœ‰åŠ¹åŒ–", value=True)
    # â€» ãƒ˜ãƒƒãƒ€ãƒ¼è¡Œã‚‚ãƒãƒ¼ã‚«ãƒ¼å¯¾è±¡ã«ã™ã‚‹ã®ã§ ignore_header ã¯å»ƒæ­¢

st.markdown("**â‘  Excel ã‹ã‚‰ã‚³ãƒ”ãƒ¼ã—ãŸè¡¨ï¼ˆTSV/CSVï¼‰ã‚’è²¼ã‚Šä»˜ã‘**")
default_text = EXAMPLE_TEXT if example else ""
text = st.text_area("ã“ã“ã«è²¼ã‚Šä»˜ã‘", value=default_text, height=200)

st.markdown("**â‘¡ æ³¨ï¼ˆè¡¨ã®ä¸‹ã«è¡¨ç¤ºãƒ»ä»»æ„ï¼‰**")
note_text = st.text_area("æ³¨ã‚’ã“ã“ã«è²¼ã‚Šä»˜ã‘ï¼ˆä»»æ„ï¼‰", value="", height=100)

# â˜… ã“ã“ã§ãƒ˜ãƒƒãƒ€ãƒ¼è¡Œæ•°ã‚’æŒ‡å®šï¼ˆã‚¹ã‚¿ã‚¤ãƒ«ã®ã€Œä¸Šã€ï¼‰
st.markdown("**â‘¢ ãƒ˜ãƒƒãƒ€ãƒ¼è¡Œæ•°ã®æŒ‡å®š**")
header_rows = st.number_input(
    "ãƒ˜ãƒƒãƒ€ãƒ¼ã®è¡Œæ•°ï¼ˆå…ˆé ­ã‹ã‚‰ä½•è¡Œã‚’ãƒ˜ãƒƒãƒ€ãƒ¼æ‰±ã„ã«ã™ã‚‹ã‹ï¼‰",
    min_value=0,
    max_value=50,
    value=1,
    step=1,
    help="å…ˆé ­ã‹ã‚‰ã“ã®è¡Œæ•°åˆ†ãŒãƒ˜ãƒƒãƒ€ãƒ¼æ‰±ã„ã«ãªã‚Šã¾ã™ã€‚è¤‡æ•°è¡Œãƒ˜ãƒƒãƒ€ãƒ¼ã‚‚å¯ã€‚",
)

st.markdown("**â‘£ ã‚¹ã‚¿ã‚¤ãƒ«ï¼ˆãƒ—ãƒªã‚»ãƒƒãƒˆ â†’ è©³ç´°èª¿æ•´å¯ï¼‰**")
c1, c2 = st.columns([1.2, 1])
with c1:
    preset = st.radio("ãƒ—ãƒªã‚»ãƒƒãƒˆ", list(PRESETS.keys()), index=3)
with c2:
    st.caption("è©³ç´°è¨­å®šã§ä¸Šæ›¸ãå¯èƒ½")

# ãƒ—ãƒªã‚»ãƒƒãƒˆå€¤ã‚’åˆæœŸå€¤ã¨ã—ã¦èª­ã¿å‡ºã—
_p = PRESETS[preset]
_default_header_bg = _p["header_bg"] or "#EEEEEE"

with st.expander("ğŸ”§ è©³ç´°è¨­å®šï¼ˆå¿…è¦ãªå ´åˆã®ã¿ï¼‰", expanded=False):
    colA, colB, colC = st.columns(3)
    with colA:
        header_bg = st.color_picker("ãƒ˜ãƒƒãƒ€ãƒ¼èƒŒæ™¯è‰²", _default_header_bg)
        header_fg = st.color_picker("ãƒ˜ãƒƒãƒ€ãƒ¼æ–‡å­—è‰²", _p["header_fg"])
        header_bold = st.checkbox("ãƒ˜ãƒƒãƒ€ãƒ¼å¤ªå­—", value=_p["header_bold"])
    with colB:
        body_bg_on = st.checkbox("æœ¬æ–‡ èƒŒæ™¯ã‚’å¡—ã‚‹", value=_p["body_bg"] is not None)
        body_bg_val = st.color_picker("æœ¬æ–‡èƒŒæ™¯è‰²", _p["body_bg"] or "#FFFFFF")
        body_fg = st.color_picker("æœ¬æ–‡æ–‡å­—è‰²", _p["body_fg"])
    with colC:
        font_name = st.selectbox(
            "ãƒ•ã‚©ãƒ³ãƒˆ",
            ["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"],
            index=["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"].index(_p["font_name"])
            if _p["font_name"] in ["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"]
            else 0,
        )
        inner_h = st.checkbox("æ¨ªç½«ç·šï¼ˆå†…å´ï¼‰", value=_p["inner_h"])
        inner_v = st.checkbox("ç¸¦ç½«ç·šï¼ˆå†…å´ï¼‰", value=_p["inner_v"])
        outer = st.checkbox("å¤–æ ï¼ˆãƒœãƒƒã‚¯ã‚¹ï¼‰", value=_p["outer"])
        zebra = st.checkbox("ã‚¼ãƒ–ãƒ©è¡Œ", value=_p["zebra"])

# è©³ç´°è¨­å®šã‚’é–‹ã‹ãªã„å ´åˆã®ãƒ‡ãƒ•ã‚©ãƒ«ãƒˆï¼ˆæœªå®šç¾©ã‚¬ãƒ¼ãƒ‰ï¼‰
if "header_bg" not in locals():
    header_bg = _default_header_bg
    header_fg = _p["header_fg"]
    header_bold = _p["header_bold"]
    body_bg_on = _p["body_bg"] is not None
    body_bg_val = _p["body_bg"] or "#FFFFFF"
    body_fg = _p["body_fg"]
    font_name = _p["font_name"]
    inner_h = _p["inner_h"]
    inner_v = _p["inner_v"]
    outer = _p["outer"]
    zebra = _p["zebra"]

# ============== åˆ—å¹…ãƒ¢ãƒ¼ãƒ‰ï¼ˆæ‰‹å‹•ã‚’è¿½åŠ ï¼‰ ==============
col_width_mode = st.radio(
    "åˆ—å¹…ãƒ¢ãƒ¼ãƒ‰",
    ["å‡ç­‰", "è‡ªå‹•ï¼ˆæ–‡å­—æ•°ã§å¯å¤‰ï¼‰", "æ‰‹å‹•ï¼ˆcmæŒ‡å®šï¼‰"],
    index=0,
    horizontal=True,
)
TOTAL_CM = 16.0
body_size = int(base_size)
header_size = int(base_size if header_same else base_size + 1)

# å…¥åŠ›ãƒ†ã‚­ã‚¹ãƒˆã‹ã‚‰å…ˆã«åˆ—æ•°ã‚’æ¨å®šã—ã¦æ‰‹å‹•UIã‚’å‡ºã™
rows_preview = parse_table(text) if text.strip() else []
n_cols_preview = len(rows_preview[0]) if rows_preview else 0

# æ‰‹å‹•å¹…ã®çŠ¶æ…‹ï¼ˆã‚»ãƒƒã‚·ãƒ§ãƒ³ï¼‰
if "manual_widths" not in st.session_state:
    st.session_state.manual_widths = []

if rows_preview and (not st.session_state.manual_widths or len(st.session_state.manual_widths) != n_cols_preview):
    # 1åˆ—ã‚ãŸã‚Š TOTAL_CM/n_cols_preview ãŒ 0.5 æœªæº€ã«ãªã‚‹ã‚±ãƒ¼ã‚¹ã‚‚ã‚ã‚‹ã®ã§ max(0.5, ...) ã§ã‚¯ãƒ©ãƒ³ãƒ—
    base_width = max(0.5, round(TOTAL_CM / max(n_cols_preview, 1), 2))
    st.session_state.manual_widths = [base_width] * n_cols_preview

if col_width_mode == "æ‰‹å‹•ï¼ˆcmæŒ‡å®šï¼‰" and rows_preview:
    with st.expander("âœ‹ æ‰‹å‹•ã§åˆ—å¹…ã‚’è¨­å®šï¼ˆåˆè¨ˆã¯ç´„16.0cmæ¨å¥¨ï¼‰", expanded=True):
        ccols = st.columns(min(n_cols_preview, 6))  # 6åˆ—ãšã¤æŠ˜ã‚Šè¿”ã—
        for i in range(n_cols_preview):
            # ã‚»ãƒƒã‚·ãƒ§ãƒ³ã®å€¤ã‚’ä¸€åº¦ 0.5ã€œ20.0 ã®ç¯„å›²ã«æ•´ãˆã‚‹
            raw_val = float(st.session_state.manual_widths[i])
            safe_val = max(0.5, min(20.0, raw_val))
            st.session_state.manual_widths[i] = safe_val

            st.session_state.manual_widths[i] = ccols[i % 6].number_input(
                f"åˆ—{i+1} å¹…(cm)",
                min_value=0.5,
                max_value=20.0,
                value=safe_val,  # â† ã‚¯ãƒ©ãƒ³ãƒ—æ¸ˆã¿ã®å€¤ã‚’æ¸¡ã™
                step=0.1,
                key=f"mw_{i}",
            )
        st.write(f"åˆè¨ˆï¼š**{sum(st.session_state.manual_widths):.2f} cm** / æ¨å¥¨ {TOTAL_CM:.1f} cm")
        b1, b2, b3 = st.columns(3)
        with b1:
            if st.button("å‡ç­‰ã«ã™ã‚‹", use_container_width=True):
                st.session_state.manual_widths = [round(TOTAL_CM / n_cols_preview, 2)] * n_cols_preview
        with b2:
            if st.button("è‡ªå‹•æ¡ˆã§åŸ‹ã‚ã‚‹ï¼ˆæ–‡å­—æ•°ï¼‰", use_container_width=True) and rows_preview:
                w_list = compute_col_widths_cm(rows_preview, total_cm=TOTAL_CM)
                # ã“ã“ã§ã‚‚ 0.5ã€œ20.0 ã«ã‚¯ãƒ©ãƒ³ãƒ—
                st.session_state.manual_widths = [
                    max(0.5, min(20.0, round(x, 2))) for x in w_list
                ]
        with b3:
            if st.button("åˆè¨ˆã‚’16cmã«ã‚¹ã‚±ãƒ¼ãƒ«", use_container_width=True):
                s = sum(st.session_state.manual_widths) or 1.0
                scaled = [x * TOTAL_CM / s for x in st.session_state.manual_widths]
                st.session_state.manual_widths = [
                    max(0.5, min(20.0, round(x, 2))) for x in scaled
                ]

st.markdown("---")
make_btn = st.button("â–¶ è¡¨ä½œæˆï¼ˆWord .docx ã‚’ç”Ÿæˆï¼†ç”»é¢è¡¨ç¤ºï¼‰", type="primary", use_container_width=True)

# ============== ä½œæˆå‡¦ç† ==============
if make_btn:
    rows = parse_table(text)
    if not rows:
        st.error("è¡¨ãƒ‡ãƒ¼ã‚¿ãŒèª­ã‚ã¾ã›ã‚“ã§ã—ãŸã€‚ãƒ†ã‚­ã‚¹ãƒˆã‚’ç¢ºèªã—ã¦ãã ã•ã„ã€‚")
        st.stop()

    # å®‰å…¨ã®ãŸã‚ï¼šãƒ˜ãƒƒãƒ€ãƒ¼è¡Œæ•°ã¯å…¨ä½“è¡Œæ•°ä»¥å†…ã«æŠ‘ãˆã‚‹
    header_rows_int = max(0, min(int(header_rows), len(rows)))

    # åˆ—å¹…æ±ºå®š
    if col_width_mode == "å‡ç­‰":
        widths_cm = [TOTAL_CM / len(rows[0])] * len(rows[0])
    elif col_width_mode == "è‡ªå‹•ï¼ˆæ–‡å­—æ•°ã§å¯å¤‰ï¼‰":
        widths_cm = compute_col_widths_cm(rows, total_cm=TOTAL_CM)
    else:  # æ‰‹å‹•ï¼ˆcmæŒ‡å®šï¼‰
        if not st.session_state.manual_widths or len(st.session_state.manual_widths) != len(rows[0]):
            st.session_state.manual_widths = [TOTAL_CM / len(rows[0])] * len(rows[0])
        widths_cm = st.session_state.manual_widths[:]
    widths_pct = widths_to_pct(widths_cm)

    # ãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ç”¨DFï¼ˆã€Œå…ˆé ­è¡Œã€ã‚’ãƒ˜ãƒƒãƒ€ãƒ¼ã«ã—ã¦ã€ãƒ‡ãƒ¼ã‚¿ã¯ header_rows_int è¡Œç›®ä»¥é™ï¼‰
    if len(rows) > 1:
        header_row = [str(x) for x in rows[0]]

        # --- åˆ—åã®é‡è¤‡ã‚’é¿ã‘ã‚‹ï¼ˆåŒã˜åå‰ãŒã‚ã‚Œã° _2, _3, ... ã‚’ä»˜ã‘ã‚‹ï¼‰ ---
        seen = {}
        cols_unique = []
        for name in header_row:
            base = name
            if base in seen:
                seen[base] += 1
                cols_unique.append(f"{base}_{seen[base]}")  # ä¾‹: "<åŒå·¦>", "<åŒå·¦>_2"
            else:
                seen[base] = 0
                cols_unique.append(base)

        df = pd.DataFrame(rows[header_rows_int:], columns=cols_unique)
    else:
        df = pd.DataFrame()

    st.success(
        f"è¡¨ã‚’èª­ã¿è¾¼ã¿ã¾ã—ãŸï¼š{len(rows) - header_rows_int}è¡Œ Ã— {len(rows[0])}åˆ—"
        f"ï¼ˆãƒ˜ãƒƒãƒ€ãƒ¼è¡Œæ•°: {header_rows_int}ã€åˆ—å¹…ãƒ¢ãƒ¼ãƒ‰: {col_width_mode}ï¼‰"
    )
    st.dataframe(df, use_container_width=True)

    # ===== Word ç”Ÿæˆ =====
    doc = Document()

    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)

    # åˆ—å¹…ã¯å…ˆã«è¨­å®š
    apply_docx_col_widths(table, widths_cm)

    # spans è¨ˆç®—ï¼ˆãƒãƒ¼ã‚«ãƒ¼æ–¹å¼ï¼šãƒ˜ãƒƒãƒ€ãƒ¼è¡Œã‚‚å«ã‚ã¦ <åŒä¸Š>/<åŒå·¦> ã‚’è§£é‡ˆï¼‰
    spans = compute_spans_markers(
        rows,
        use_up=use_up,
        use_left=use_left,
        header_rows=header_rows_int,  # ãƒ˜ãƒ«ãƒ‘å´ãŒ header_rows ã‚’å—ã‘å–ã‚‹å‰æ
        strict_rect=True,
    )

    # å®Ÿã‚»ãƒ«ã¯çµåˆã—ãªã„ï¼šã‚¢ãƒ³ã‚«ãƒ¼ä»¥å¤–ã¯ç©ºæ¬„ï¼‹ç½«ç·šã§è¦‹ã‹ã‘ã ã‘çµåˆ
    merge_docx_by_spans(
        table,
        rows,
        spans,
        font_name=font_name,
        base_size=int(body_size),
        header_size=int(header_size),
        header_fg=header_fg,
        body_fg=body_fg,
        header_bg=(header_bg or None),
        body_bg=(body_bg_val if body_bg_on else None),
        zebra=zebra,
        header_same=header_same,
        body_bg_on=body_bg_on,
        header_rows=header_rows_int,
    )

    # ç½«ç·šï¼ˆspans/anchor ãƒ™ãƒ¼ã‚¹ã§ã€çµåˆãƒ–ãƒ­ãƒƒã‚¯å†…éƒ¨ã«ã¯ç·šã‚’å¼•ã‹ãªã„ï¼‰
    apply_table_borders_robust(
        table,
        spans,
        inner_h=inner_h,
        inner_v=inner_v,
        outer=outer,
        sz_inner=6,
        sz_outer=12,  # â† å¸¸ã« 12pt ã«å›ºå®š
        color="000000",
    )

    # æ³¨ï¼ˆä»»æ„ï¼‰
    if note_text.strip():
        p = doc.add_paragraph()
        run = p.add_run(note_text.strip())
        apply_font_run(run, font_name, int(body_size), "#444444")

    # ---- Word ãƒãƒƒãƒ•ã‚¡ã ã‘å…ˆã«ç”¨æ„ï¼ˆã“ã“ã§ã¯ã¾ã ãƒœã‚¿ãƒ³ã‚’å‡ºã•ãªã„ï¼‰----
    buf_docx = BytesIO()
    doc.save(buf_docx)

    # ===== HTMLãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ï¼ˆrowspan/colspan åæ˜ ï¼‰ =====
    st.subheader("ğŸ” ä½œæˆçµæœï¼ˆç”»é¢ãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ï¼‰")
    html = build_html_table_with_spans(
        rows,
        spans,
        header_rows=header_rows_int,
        header_bg=(header_bg or "#EEEEEE"),
        header_fg=header_fg,
        header_size=int(header_size),
        header_bold=True,
        body_bg=(body_bg_val if body_bg_on else None),
        body_fg=body_fg,
        body_size=int(body_size),
        font_name=font_name,
        zebra=zebra,
        inner_h=inner_h,
        inner_v=inner_v,
        outer=outer,
        note_text=note_text,
        col_width_pct=widths_pct,
    )
    st.markdown(html, unsafe_allow_html=True)

    # ===== å‡ºåŠ›ï¼ˆWord â†’ HTML ã®é †ã«ãƒœã‚¿ãƒ³ã‚’é…ç½®ï¼‰ =====
    st.markdown("### ğŸ“¥ å‡ºåŠ›")

    # --- Word ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰ãƒœã‚¿ãƒ³ï¼ˆãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ç›´ä¸‹ï¼‰ ---
    st.download_button(
        "ğŸ“¥ Wordï¼ˆ.docxï¼‰ã‚’ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰",
        data=buf_docx.getvalue(),
        file_name="table_generated.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    # --- HTML ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰ãƒœã‚¿ãƒ³ï¼ˆãã®ä¸‹ï¼‰ ---
    html_doc = f"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="UTF-8">
<title>ç”Ÿæˆã•ã‚ŒãŸè¡¨</title>
<style>
  body {{
    font-family: "{font_name}", -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    margin: 24px;
  }}
</style>
</head>
<body>
{html}
</body>
</html>
"""
    st.download_button(
        "ğŸ“„ HTML ã‚’ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰",
        data=html_doc.encode("utf-8"),
        file_name="table_generated.html",
        mime="text/html",
        use_container_width=True,
    )
