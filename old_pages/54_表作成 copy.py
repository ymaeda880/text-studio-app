# -*- coding: utf-8 -*-
# pages/54_表作成.py
#
# Excel貼り付け → スタイル選択 → Word .docx 生成
# ・画面プレビュー表示
# ・表下の注テキスト対応
# ・サイドバーの基本フォントサイズスライダー
# ・列幅モード：均等 / 自動（文字数で可変）

from __future__ import annotations
import io, csv, re, unicodedata
from typing import List, Tuple, Optional

import streamlit as st
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================
# ユーティリティ
# =========================
def _detect_delimiter(text: str) -> str:
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return "\t"
    head = lines[0]
    if "\t" in head:
        return "\t"
    comma_count = head.count(",")
    semicol_count = head.count(";")
    if comma_count >= semicol_count and comma_count > 0:
        return ","
    if semicol_count > 0:
        return ";"
    if re.search(r"\s{2,}", head):
        return r"\s+"
    return "\t"

def _parse_table(text: str) -> List[List[str]]:
    text = text.strip("\n\r ")
    if not text:
        return []
    delim = _detect_delimiter(text)
    if delim == r"\s+":
        rows = [re.split(r"\s{2,}", ln.strip()) for ln in text.splitlines() if ln.strip()]
        return rows
    reader = csv.reader(io.StringIO(text), delimiter=("\t" if delim == "\t" else delim))
    rows = [list(r) for r in reader]
    rows = [r for r in rows if any(c.strip() for c in r)]
    maxc = max((len(r) for r in rows), default=0)
    rows = [r + [""] * (maxc - len(r)) for r in rows]
    return rows

def _hex_to_rgb(hex_color: str) -> Tuple[int, int, int]:
    hc = hex_color.strip().lstrip("#")
    if len(hc) == 3:
        hc = "".join([c*2 for c in hc])
    r = int(hc[0:2], 16); g = int(hc[2:4], 16); b = int(hc[4:6], 16)
    return r, g, b

def _set_cell_shading(cell, hex_color: Optional[str]):
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip("#").upper())
    tcPr.append(shd)

def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('left','right','top','bottom','insideH','insideV'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = OxmlElement(f'w:{edge}')
            val, sz, color = edge_data
            tag.set(qn('w:val'), val)
            tag.set(qn('w:sz'), str(sz))  # 1/8 pt
            tag.set(qn('w:color'), color)
            tcBorders.append(tag)

def _apply_table_outer_border(table, val="single", sz=12, color="000000"):
    rows = table.rows; cols = table.columns
    for c in range(len(cols)):
        _set_cell_border(rows[0].cells[c], top=(val, sz, color))
        _set_cell_border(rows[-1].cells[c], bottom=(val, sz, color))
    for r in range(len(rows)):
        _set_cell_border(rows[r].cells[0], left=(val, sz, color))
        _set_cell_border(rows[r].cells[-1], right=(val, sz, color))

def _apply_table_inner_borders(table, vertical=True, horizontal=True, val="single", sz=6, color="000000"):
    rows = table.rows; cols = table.columns
    if horizontal:
        for r in range(len(rows)-1):
            for c in range(len(cols)):
                _set_cell_border(rows[r].cells[c], bottom=(val, sz, color))
    if vertical:
        for r in range(len(rows)):
            for c in range(len(cols)-1):
                _set_cell_border(rows[r].cells[c], right=(val, sz, color))

def _apply_font_run(run, font_name: str, size_pt: int, color_hex: str, bold: bool=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # CJK
    run.font.size = Pt(size_pt)
    r, g, b = _hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    run.bold = bold

# ---- 列幅（見かけ文字数）ユーティリティ ----
def _visual_len(s: str) -> int:
    """全角=2, 半角=1 の概算幅"""
    if s is None:
        return 0
    total = 0
    for ch in str(s):
        ea = unicodedata.east_asian_width(ch)
        total += 2 if ea in ("W","F") else 1
    return total

def _compute_col_widths_cm(rows: List[List[str]], total_cm: float = 16.0,
                           min_cm: float = 2.2, max_cm: float = 8.0) -> List[float]:
    """各列の最大視覚長に比例して列幅(cm)を割付"""
    if not rows:
        return []
    n_cols = len(rows[0])
    scores = []
    for c in range(n_cols):
        mx = 0
        for r in range(len(rows)):
            mx = max(mx, _visual_len(rows[r][c]))
        # 見出しは少し重み＋10%
        mx = int(mx * 1.1)
        scores.append(mx or 1)
    ssum = sum(scores) or 1
    raw = [total_cm * (sc / ssum) for sc in scores]
    clamped = [max(min_cm, min(max_cm, x)) for x in raw]

    # クランプ後の合計が total_cm とズレたら比例補正（最小値は維持）
    diff = total_cm - sum(clamped)
    if abs(diff) > 1e-6:
        # 余剰/不足を“伸ばせる列”に配る
        room = [max_cm - w for w in clamped] if diff > 0 else [w - min_cm for w in clamped]
        room_sum = sum(x for x in room if x > 0) or 1
        adj = [ (diff * (r/room_sum) if r > 0 else 0.0) for r in room ]
        clamped = [w + a for w, a in zip(clamped, adj)]
    return clamped

def _widths_to_pct(widths_cm: List[float]) -> List[float]:
    tot = sum(widths_cm) or 1
    return [round(w * 100.0 / tot, 2) for w in widths_cm]

def _apply_docx_col_widths(table, widths_cm):
    """
    Word に確実に列幅を効かせる。
    - オートフィット無効（tblLayout=fixed）
    - 表全体幅 (tblW) を設定
    - 列グリッド (tblGrid) に各列幅を設定
    - 各列セルの width も冪等的に設定
    """
    # 1) AutoFit を完全に切る（固定レイアウト）
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout')
        tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')

    # 2) 表全体の幅を twips で（1cm = 567 twips）
    tot_cm = sum(widths_cm) or 1.0
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa')
    tblW.set(qn('w:w'), str(int(tot_cm * 567)))

    # 3) 列グリッド（tblGrid）を作り直して幅指定
    #    既存があれば削除してから再作成
    for child in list(table._tbl.iterchildren()):
        if child.tag == qn('w:tblGrid'):
            table._tbl.remove(child)
    grid = OxmlElement('w:tblGrid')
    for wcm in widths_cm:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(wcm * 567)))
        grid.append(gc)
    table._tbl.insert(1, grid)  # tblPr の次あたりに入れる

    # 4) 各列のセルにも幅を明示（念押し）
    for c, wcm in enumerate(widths_cm):
        for r in range(len(table.rows)):
            cell = table.cell(r, c)
            cell.width = Cm(wcm)


def _build_html_table(rows: List[List[str]], *,
                      header_bg: str, header_fg: str, header_size: int, header_bold: bool,
                      body_bg: Optional[str], body_fg: str, body_size: int,
                      font_name: str, zebra: bool,
                      inner_h: bool, inner_v: bool, outer: bool,
                      note_text: Optional[str] = None,
                      col_width_pct: Optional[List[float]] = None) -> str:
    """画面表示用：選択スタイルを反映したHTMLテーブル＋注（任意）"""
    if not rows:
        return ""
    n_rows, n_cols = len(rows), len(rows[0])
    border_css = "border-collapse:collapse;"
    outer_css = "border:1px solid #000;" if outer else ""
    td_base = "padding:6px 10px;vertical-align:top;"
    td_border = []
    if inner_h or inner_v:
        td_border.append("border:1px solid #000;")
    elif outer:
        td_border.append("border:0;")
    else:
        td_border.append("border:0;")
    td_css = td_base + "".join(td_border)
    zebra_alt = "#F7F9FC"

    html = [f'<table style="{border_css}{outer_css}; width:100%;">']

    # 列幅（HTML）
    if col_width_pct:
        html.append("<colgroup>")
        for p in col_width_pct:
            html.append(f'<col style="width:{p}%">')
        html.append("</colgroup>")

    # header
    html.append("<thead><tr>")
    for c in range(n_cols):
        bold_css = "font-weight:700;" if header_bold else "font-weight:400;"
        html.append(
            f'<th style="{td_css}background:{header_bg};color:{header_fg};'
            f'font-family:{font_name};font-size:{header_size}pt;{bold_css}text-align:left;">'
            f'{rows[0][c]}'
            f'</th>'
        )
    html.append("</tr></thead>")

    # body
    html.append("<tbody>")
    for r in range(1, n_rows):
        row_bg = (zebra_alt if (zebra and r % 2 == 1) else (body_bg or "transparent"))
        html.append(f'<tr style="background:{row_bg};">')
        for c in range(n_cols):
            html.append(
                f'<td style="{td_css}font-family:{font_name};font-size:{body_size}pt;'
                f'color:{body_fg};text-align:left;">{rows[r][c]}</td>'
            )
        html.append("</tr>")
    html.append("</tbody></table>")

    # 注（任意）
    if note_text and note_text.strip():
        esc = (
            note_text.strip()
            .replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        ).replace("\n", "<br/>")
        html.append(
            f'<div style="margin-top:8px;font-family:{font_name};font-size:{body_size}pt;'
            f'color:#444;">{esc}</div>'
        )
    return "\n".join(html)

# =========================
# UI
# =========================
st.set_page_config(page_title="🧾 Word 表作成", page_icon="🧾", layout="wide")
st.title("🧾 Word 表作成（Excel貼り付け → スタイル指定 → .docx ＋ 画面表示＋注）")

with st.sidebar:
    st.markdown("### 1) 基本設定")
    example = st.checkbox("サンプルデータを入れる", value=False)
    base_size = st.slider("基本フォントサイズ（pt）", min_value=8, max_value=16, value=10, step=1)
    header_same = st.checkbox("ヘッダーも同じサイズを使う（OFFなら +1pt）", value=False)

st.markdown("**① Excel からコピーした表（行列）を貼り付け**（TSV/CSV どちらでもOK）")
default_text = ""
if example:
    default_text = (
        "品目\t数量\t単価(円)\t金額(円)\n"
        "りんご\t12\t120\t1440\n"
        "みかん\t8\t80\t640\n"
        "ぶどう\t2\t450\t900"
    )
text = st.text_area("ここに貼り付け", value=default_text, height=200,
                    placeholder="例）Excelを範囲コピー → ここに貼り付け")

st.markdown("**② 注（表の下に表示・任意）**")
note_text = st.text_area("注をここに貼り付け（任意）", value="", height=100,
                         placeholder="例）※ ○内はサンプル数　※ 背景色：前方分析… など")

st.markdown("**③ スタイル設定**（プリセット → 必要なら詳細調整）")

PRESETS = {
    "シンプル（横罫のみ・薄グレー見出し）": dict(
        header_bg="#EEEEEE", header_fg="#000000", header_bold=True,
        body_bg=None, body_fg="#000000",
        font_name="Meiryo",
        inner_h=True, inner_v=False, outer=True, zebra=False
    ),
    "ゼブラ（横＋縦罫・見出し濃色）": dict(
        header_bg="#2F5597", header_fg="#FFFFFF", header_bold=True,
        body_bg=None, body_fg="#222222",
        font_name="Meiryo",
        inner_h=True, inner_v=True, outer=True, zebra=True
    ),
    "横罫のみ（モノトーン）": dict(
        header_bg=None, header_fg="#000000", header_bold=True,
        body_bg=None, body_fg="#000000",
        font_name="Meiryo",
        inner_h=True, inner_v=False, outer=False, zebra=False
    ),
    "ボックス＋見出し色": dict(
        header_bg="#DDEBF7", header_fg="#000000", header_bold=True,
        body_bg="#FFFFFF", body_fg="#000000",
        font_name="Meiryo",
        inner_h=True, inner_v=True, outer=True, zebra=False
    ),
    "濃色ヘッダー（大きめ文字）": dict(
        header_bg="#1F4E79", header_fg="#FFFFFF", header_bold=True,
        body_bg="#FFFFFF", body_fg="#222222",
        font_name="Meiryo",
        inner_h=True, inner_v=True, outer=True, zebra=False
    ),
}

c1, c2 = st.columns([1.2, 1])
with c1:
    preset = st.radio("プリセットを選択", list(PRESETS.keys()), index=0)
with c2:
    st.caption("プリセットは下の詳細設定で上書き可能です。")

with st.expander("🔧 詳細設定（必要な場合のみ）", expanded=False):
    colA, colB, colC = st.columns(3)
    with colA:
        header_bg = st.color_picker("ヘッダー背景色", PRESETS[preset]["header_bg"] or "#EEEEEE")
        header_fg = st.color_picker("ヘッダー文字色", PRESETS[preset]["header_fg"])
        header_bold = st.checkbox("ヘッダー太字", value=PRESETS[preset]["header_bold"])
    with colB:
        body_bg_on = st.checkbox("本文 背景を塗る", value=PRESETS[preset]["body_bg"] is not None)
        body_bg_val = st.color_picker("本文背景色", PRESETS[preset]["body_bg"] or "#FFFFFF")
        body_fg = st.color_picker("本文文字色", PRESETS[preset]["body_fg"])
    with colC:
        font_name = st.selectbox("フォント", ["Meiryo","Yu Gothic","MS PGothic","MS Gothic","Calibri","Arial"], index=0)
        inner_h = st.checkbox("横罫線（内側）", value=PRESETS[preset]["inner_h"])
        inner_v = st.checkbox("縦罫線（内側）", value=PRESETS[preset]["inner_v"])
        outer   = st.checkbox("外枠（ボックス）", value=PRESETS[preset]["outer"])
        zebra   = st.checkbox("ゼブラ行（本文交互塗り）", value=PRESETS[preset]["zebra"])
        rounded_like = st.checkbox("外枠を太め＋余白（角丸風の近似）", value=False)

# 列幅モード（新規）
col_width_mode = st.radio("列幅モード", ["均等", "自動（文字数で可変）"], index=0, horizontal=True)

# フォントサイズ
body_size = int(base_size)
header_size = int(base_size if header_same else base_size + 1)

st.markdown("---")
make_btn = st.button("▶ 表作成（Word .docx を生成＆画面表示）", type="primary", use_container_width=True)

# =========================
# 処理
# =========================
if make_btn:
    rows = _parse_table(text)
    if not rows:
        st.error("表データが読めませんでした。テキストを確認してください。")
        st.stop()

    # 列幅算出
    if col_width_mode == "自動（文字数で可変）":
        widths_cm = _compute_col_widths_cm(rows, total_cm=16.0, min_cm=2.2, max_cm=8.0)
    else:
        n_cols = len(rows[0])
        widths_cm = [16.0 / n_cols] * n_cols
    widths_pct = _widths_to_pct(widths_cm)

    # DataFrame プレビュー
    df = pd.DataFrame(rows[1:], columns=rows[0] if rows else None)
    st.success(f"表を読み込みました：{df.shape[0]}行 × {df.shape[1]}列（基本フォント {body_size}pt / 列幅モード: {col_width_mode}）")
    st.dataframe(df, use_container_width=True)

    # Word ドキュメント作成
    doc = Document()
    if rounded_like:
        sections = doc.sections
        for sec in sections:
            sec.left_margin = Pt(36)
            sec.right_margin = Pt(36)
            sec.top_margin = Pt(36)
            sec.bottom_margin = Pt(36)

    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    _apply_docx_col_widths(table, widths_cm)  # ← 列幅反映（重要）

    # ヘッダー
    hb = header_bg or "#EEEEEE"
    for c in range(n_cols):
        cell = table.cell(0, c)
        if hb:
            _set_cell_shading(cell, hb)
        cell.text = str(rows[0][c])
        for p in cell.paragraphs:
            for run in p.runs:
                _apply_font_run(run, font_name, header_size, header_fg, bold=header_bold)

    # 本文
    zebra_alt = "#F7F9FC"
    body_bg = (body_bg_val if body_bg_on else None)
    for r in range(1, n_rows):
        for c in range(n_cols):
            cell = table.cell(r, c)
            fill = (zebra_alt if (zebra and r % 2 == 1) else body_bg)
            if fill:
                _set_cell_shading(cell, fill)
            cell.text = str(rows[r][c])
            for p in cell.paragraphs:
                for run in p.runs:
                    _apply_font_run(run, font_name, body_size, body_fg, bold=False)

    # 罫線
    _apply_table_inner_borders(table, vertical=inner_v, horizontal=inner_h, val="single", sz=6, color="000000")
    if outer:
        _apply_table_outer_border(table, val="single", sz=(16 if rounded_like else 12), color="000000")

    # 注（任意）
    if note_text and note_text.strip():
        p = doc.add_paragraph()
        run = p.add_run(note_text.strip())
        _apply_font_run(run, font_name, body_size, "#444444", bold=False)

    # .docx ダウンロード
    buf = io.BytesIO()
    doc.save(buf)
    st.download_button(
        "📥 Word（.docx）をダウンロード",
        data=buf.getvalue(),
        file_name="table_generated.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True
    )

    # 画面プレビュー（HTML）— 列幅％を colgroup に反映
    st.subheader("🔍 作成結果（画面プレビュー）")
    html = _build_html_table(
        rows,
        header_bg=hb, header_fg=header_fg, header_size=header_size, header_bold=header_bold,
        body_bg=(body_bg or None), body_fg=body_fg, body_size=body_size,
        font_name=font_name, zebra=zebra,
        inner_h=inner_h, inner_v=inner_v, outer=outer,
        note_text=note_text, col_width_pct=widths_pct
    )
    st.markdown(html, unsafe_allow_html=True)

    with st.expander("🧩 代替：コピー用HTML（Wordに貼り付け可）", expanded=False):
        st.code(html, language="html")

else:
    st.info("左のサンプルをONにして貼り付け→スタイルを選び「表作成」を押してください。")
