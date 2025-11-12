# -*- coding: utf-8 -*-
# pages/54_表作成.py
#
# Excel貼り付け → スタイル選択 → Word .docx 生成
# ・画面プレビュー表示
# ・表下の注テキスト
# ・サイドバーに基本フォントサイズ
# ・列幅モード：均等 / 自動（文字数） / 手動（cm指定）

from __future__ import annotations
import io, csv, re, unicodedata
from typing import List, Tuple, Optional

import streamlit as st
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ============== ユーティリティ ==============
def _detect_delimiter(text: str) -> str:
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines: return "\t"
    head = lines[0]
    if "\t" in head: return "\t"
    if head.count(",") >= head.count(";") and head.count(",") > 0: return ","
    if head.count(";") > 0: return ";"
    if re.search(r"\s{2,}", head): return r"\s+"
    return "\t"

def _parse_table(text: str) -> List[List[str]]:
    text = text.strip("\n\r ")
    if not text: return []
    delim = _detect_delimiter(text)
    if delim == r"\s+":
        rows = [re.split(r"\s{2,}", ln.strip()) for ln in text.splitlines() if ln.strip()]
    else:
        reader = csv.reader(io.StringIO(text), delimiter=("\t" if delim == "\t" else delim))
        rows = [list(r) for r in reader]
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows: return []
    maxc = max(len(r) for r in rows)
    rows = [r + [""] * (maxc - len(r)) for r in rows]
    return rows

def _hex_to_rgb(h: str) -> Tuple[int,int,int]:
    s = h.strip().lstrip("#")
    if len(s)==3: s = "".join(c*2 for c in s)
    return int(s[0:2],16), int(s[2:4],16), int(s[4:6],16)

def _set_cell_shading(cell, hex_color: Optional[str]):
    if not hex_color: return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), hex_color.lstrip("#").upper()); tcPr.append(shd)

def _set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
    for edge in ('left','right','top','bottom','insideH','insideV'):
        if edge in kwargs:
            val, sz, color = kwargs[edge]
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), val); tag.set(qn('w:sz'), str(sz)); tag.set(qn('w:color'), color)
            tcBorders.append(tag)

def _apply_table_outer_border(table, val="single", sz=12, color="000000"):
    rows = table.rows; cols = table.columns
    for c in range(len(cols)):
        _set_cell_border(rows[0].cells[c], top=(val, sz, color))
        _set_cell_border(rows[-1].cells[c], bottom=(val, sz, color))
    for r in range(len(rows)):
        _set_cell_border(rows[r].cells[0], left=(val, sz, color))
        _set_cell_border(rows[r].cells[-1], right=(val, sz, color))

def _apply_table_inner_borders(table, vertical=True, horizontal=True, val="single", sz=6, color="000000"):
    rows = table.rows; cols = table.columns
    if horizontal:
        for r in range(len(rows)-1):
            for c in range(len(cols)):
                _set_cell_border(rows[r].cells[c], bottom=(val, sz, color))
    if vertical:
        for r in range(len(rows)):
            for c in range(len(cols)-1):
                _set_cell_border(rows[r].cells[c], right=(val, sz, color))

def _apply_font_run(run, font_name: str, size_pt: int, color_hex: str, bold: bool=False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(size_pt)
    r,g,b = _hex_to_rgb(color_hex); run.font.color.rgb = RGBColor(r,g,b)
    run.bold = bold

# ---- 列幅用 ----
def _visual_len(s: str) -> int:
    if s is None: return 0
    t=0
    for ch in str(s):
        t += 2 if unicodedata.east_asian_width(ch) in ("W","F") else 1
    return t

def _compute_col_widths_cm(rows: List[List[str]], total_cm=16.0, min_cm=2.2, max_cm=8.0) -> List[float]:
    if not rows: return []
    n = len(rows[0]); scores=[]
    for c in range(n):
        mx = 0
        for r in range(len(rows)):
            mx = max(mx, _visual_len(rows[r][c]))
        scores.append(int(mx*1.1) or 1)
    ssum = sum(scores) or 1
    raw = [total_cm*(sc/ssum) for sc in scores]
    clamped = [max(min_cm, min(max_cm, x)) for x in raw]
    # 合計補正
    diff = total_cm - sum(clamped)
    if abs(diff) > 1e-6:
        room = [ (max_cm - w) if diff>0 else (w - 2.2) for w in clamped ]
        room_sum = sum(x for x in room if x>0) or 1
        adj = [ (diff*(r/room_sum) if r>0 else 0) for r in room ]
        clamped = [ w + a for w,a in zip(clamped, adj) ]
    return clamped

def _widths_to_pct(widths_cm: List[float]) -> List[float]:
    tot = sum(widths_cm) or 1
    return [round(w*100.0/tot, 2) for w in widths_cm]

def _apply_docx_col_widths(table, widths_cm: List[float]):
    # 固定レイアウト + 表幅 + グリッド + セル幅
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn('w:tblLayout'))
    if layout is None:
        layout = OxmlElement('w:tblLayout'); tblPr.append(layout)
    layout.set(qn('w:type'), 'fixed')
    # 表幅
    tot_cm = sum(widths_cm) or 1.0
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW'); tblPr.append(tblW)
    tblW.set(qn('w:type'), 'dxa'); tblW.set(qn('w:w'), str(int(tot_cm*567)))
    # グリッド作り直し
    for child in list(table._tbl.iterchildren()):
        if child.tag == qn('w:tblGrid'):
            table._tbl.remove(child)
    grid = OxmlElement('w:tblGrid')
    for wcm in widths_cm:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(wcm*567))); grid.append(gc)
    table._tbl.insert(1, grid)
    # 各セル幅
    for c,wcm in enumerate(widths_cm):
        for r in range(len(table.rows)):
            table.cell(r,c).width = Cm(wcm)


def _build_html_table(rows: List[List[str]], *,
                      header_bg: str, header_fg: str, header_size: int, header_bold: bool,
                      body_bg: Optional[str], body_fg: str, body_size: int,
                      font_name: str, zebra: bool,
                      inner_h: bool, inner_v: bool, outer: bool,
                      note_text: Optional[str] = None,
                      col_width_pct: Optional[List[float]] = None) -> str:
    """画面表示用：選択スタイルを反映したHTMLテーブル＋注（任意）
       ※ 横線のみ（inner_h=True, inner_v=False）の場合に縦線を描画しない
    """
    if not rows:
        return ""
    n_rows, n_cols = len(rows), len(rows[0])

    # テーブル枠（外枠）
    border_css  = "border-collapse:collapse;"
    outer_css   = "border:1px solid #000;" if outer else "border:0;"

    # 罫線スタイル（セル単位）
    # ・inner_h: 水平方向の線（top/bottom）を 1px
    # ・inner_v: 垂直方向の線（left/right）を 1px
    # ・両方ONなら border:1px
    if inner_h and inner_v:
        td_border_css = "border:1px solid #000;"
    elif inner_h and not inner_v:
        td_border_css = "border-style:solid; border-color:#000; border-width:1px 0;"
    elif inner_v and not inner_h:
        td_border_css = "border-style:solid; border-color:#000; border-width:0 1px;"
    else:
        td_border_css = "border:0;"

    # 共通
    td_base = "padding:6px 10px; vertical-align:top;"
    th_weight = "font-weight:700;" if header_bold else "font-weight:400;"
    zebra_alt = "#F7F9FC"

    html = [f'<table style="{border_css}{outer_css} width:100%;">']

    # 列幅（%）を <colgroup> に適用
    if col_width_pct:
        html.append("<colgroup>")
        for p in col_width_pct:
            html.append(f'<col style="width:{p}%">')
        html.append("</colgroup>")

    # header
    html.append("<thead><tr>")
    for c in range(n_cols):
        html.append(
            f'<th style="{td_base}{td_border_css}'
            f'background:{header_bg}; color:{header_fg}; '
            f'font-family:{font_name}; font-size:{header_size}pt; {th_weight} text-align:left;">'
            f'{rows[0][c]}'
            f'</th>'
        )
    html.append("</tr></thead>")

    # body
    html.append("<tbody>")
    for r in range(1, n_rows):
        row_bg = (zebra_alt if (zebra and r % 2 == 1) else (body_bg or "transparent"))
        html.append(f'<tr style="background:{row_bg};">')
        for c in range(n_cols):
            html.append(
                f'<td style="{td_base}{td_border_css}'
                f'font-family:{font_name}; font-size:{body_size}pt; color:{body_fg}; text-align:left;">'
                f'{rows[r][c]}'
                f'</td>'
            )
        html.append("</tr>")
    html.append("</tbody></table>")

    # 注（任意）
    if note_text and note_text.strip():
        esc = (
            note_text.strip()
            .replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        ).replace("\n", "<br/>")
        html.append(
            f'<div style="margin-top:8px; font-family:{font_name}; font-size:{body_size}pt; color:#444;">{esc}</div>'
        )

    return "\n".join(html)


# ============== UI ==============
st.set_page_config(page_title="🧾 Word 表作成", page_icon="🧾", layout="wide")
st.title("🧾 Word 表作成（Excel貼り付け → スタイル指定 → .docx ＋ プレビュー）")

with st.sidebar:
    st.markdown("### 1) 基本設定")
    example = st.checkbox("サンプルデータを入れる", value=False)
    base_size = st.slider("基本フォントサイズ（pt）", 8, 16, 10, 1)
    header_same = st.checkbox("ヘッダーも同じサイズ（OFFなら +1pt）", value=False)

st.markdown("**① Excel からコピーした表（TSV/CSV）を貼り付け**")
default_text = ""
if example:
    default_text=("品目\t数量\t単価(円)\t金額(円)\n"
                  "りんご\t12\t120\t1440\n"
                  "みかん\t8\t80\t640\n"
                  "ぶどう\t2\t450\t900")
text = st.text_area("ここに貼り付け", value=default_text, height=200)

st.markdown("**② 注（表の下に表示・任意）**")
note_text = st.text_area("注をここに貼り付け（任意）", value="", height=100)

st.markdown("**③ スタイル（プリセット → 詳細調整可）**")
PRESETS = {
    "シンプル（横罫のみ・薄グレー見出し）": dict(
        header_bg="#EEEEEE", header_fg="#000000", header_bold=True,
        body_bg=None, body_fg="#000000", font_name="Meiryo",
        inner_h=True, inner_v=False, outer=True, zebra=False
    ),
    "ゼブラ（横＋縦罫・見出し濃色）": dict(
        header_bg="#2F5597", header_fg="#FFFFFF", header_bold=True,
        body_bg=None, body_fg="#222222", font_name="Meiryo",
        inner_h=True, inner_v=True, outer=True, zebra=True
    ),
    "横罫のみ（モノトーン）": dict(
        header_bg=None, header_fg="#000000", header_bold=True,
        body_bg=None, body_fg="#000000", font_name="Meiryo",
        inner_h=True, inner_v=False, outer=False, zebra=False
    ),
    "ボックス＋見出し色": dict(
        header_bg="#DDEBF7", header_fg="#000000", header_bold=True,
        body_bg="#FFFFFF", body_fg="#000000", font_name="Meiryo",
        inner_h=True, inner_v=True, outer=True, zebra=False
    ),
        "濃色ヘッダー": dict(
        header_bg="#1F4E79", header_fg="#FFFFFF", header_bold=True,
        body_bg="#FFFFFF", body_fg="#222222",
        font_name="Meiryo",
        inner_h=True, inner_v=True, outer=True, zebra=False
    ),
}

c1, c2 = st.columns([1.2,1])
with c1: preset = st.radio("プリセット", list(PRESETS.keys()), index=0)
with c2: st.caption("詳細設定で上書き可能")

with st.expander("🔧 詳細設定（必要な場合のみ）", expanded=False):
    colA,colB,colC = st.columns(3)
    with colA:
        header_bg = st.color_picker("ヘッダー背景色", PRESETS[preset]["header_bg"] or "#EEEEEE")
        header_fg = st.color_picker("ヘッダー文字色", PRESETS[preset]["header_fg"])
        header_bold = st.checkbox("ヘッダー太字", value=PRESETS[preset]["header_bold"])
    with colB:
        body_bg_on = st.checkbox("本文 背景を塗る", value=PRESETS[preset]["body_bg"] is not None)
        body_bg_val = st.color_picker("本文背景色", PRESETS[preset]["body_bg"] or "#FFFFFF")
        body_fg = st.color_picker("本文文字色", PRESETS[preset]["body_fg"])
    with colC:
        font_name = st.selectbox("フォント", ["Meiryo","Yu Gothic","MS PGothic","MS Gothic","Calibri","Arial"], index=0)
        inner_h = st.checkbox("横罫線（内側）", value=PRESETS[preset]["inner_h"])
        inner_v = st.checkbox("縦罫線（内側）", value=PRESETS[preset]["inner_v"])
        outer   = st.checkbox("外枠（ボックス）", value=PRESETS[preset]["outer"])
        zebra   = st.checkbox("ゼブラ行", value=PRESETS[preset]["zebra"])
        rounded_like = st.checkbox("外枠を太め＋余白（角丸風）", value=False)

# ============== 列幅モード（手動を追加） ==============
col_width_mode = st.radio("列幅モード", ["均等", "自動（文字数で可変）", "手動（cm指定）"], index=0, horizontal=True)
TOTAL_CM = 16.0
body_size = int(base_size)
header_size = int(base_size if header_same else base_size+1)

# 入力テキストから先に列数を推定して手動UIを出す
rows_preview = _parse_table(text) if text.strip() else []
n_cols = len(rows_preview[0]) if rows_preview else 0

# 手動幅の状態
if "manual_widths" not in st.session_state:
    st.session_state.manual_widths = []
if rows_preview and (not st.session_state.manual_widths or len(st.session_state.manual_widths)!=n_cols):
    # 初期は均等
    st.session_state.manual_widths = [round(TOTAL_CM/max(n_cols,1),2)]*n_cols

if col_width_mode == "手動（cm指定）" and rows_preview:
    with st.expander("✋ 手動で列幅を設定（合計は約16.0cm推奨）", expanded=True):
        ccols = st.columns(min(n_cols,6))  # 6列ずつ折り返し
        for i in range(n_cols):
            st.session_state.manual_widths[i] = ccols[i%6].number_input(
                f"列{i+1} 幅(cm)", min_value=0.5, max_value=20.0,
                value=float(st.session_state.manual_widths[i]), step=0.1, key=f"mw_{i}"
            )
        st.write(f"合計：**{sum(st.session_state.manual_widths):.2f} cm** / 推奨 {TOTAL_CM:.1f} cm")
        b1,b2,b3 = st.columns(3)
        with b1:
            if st.button("均等にする", use_container_width=True):
                st.session_state.manual_widths = [round(TOTAL_CM/n_cols,2)]*n_cols
        with b2:
            if st.button("自動案で埋める（文字数）", use_container_width=True) and rows_preview:
                st.session_state.manual_widths = [round(x,2) for x in _compute_col_widths_cm(rows_preview, total_cm=TOTAL_CM)]
        with b3:
            if st.button("合計を16cmにスケール", use_container_width=True):
                s = sum(st.session_state.manual_widths) or 1.0
                st.session_state.manual_widths = [round(x*TOTAL_CM/s,2) for x in st.session_state.manual_widths]

st.markdown("---")
make_btn = st.button("▶ 表作成（Word .docx を生成＆画面表示）", type="primary", use_container_width=True)

# ============== 作成処理 ==============
if make_btn:
    rows = _parse_table(text)
    if not rows:
        st.error("表データが読めませんでした。テキストを確認してください。"); st.stop()

    # 列幅決定
    if col_width_mode == "均等":
        widths_cm = [TOTAL_CM/len(rows[0])]*len(rows[0])
    elif col_width_mode == "自動（文字数で可変）":
        widths_cm = _compute_col_widths_cm(rows, total_cm=TOTAL_CM)
    else:
        # 手動
        if not st.session_state.manual_widths or len(st.session_state.manual_widths)!=len(rows[0]):
            st.session_state.manual_widths = [TOTAL_CM/len(rows[0])]*len(rows[0])
        widths_cm = st.session_state.manual_widths[:]
    widths_pct = _widths_to_pct(widths_cm)

    # プレビュー用DF
    df = pd.DataFrame(rows[1:], columns=rows[0])
    st.success(f"表を読み込みました：{df.shape[0]}行 × {df.shape[1]}列（列幅モード: {col_width_mode}）")
    st.dataframe(df, use_container_width=True)

    # Word
    doc = Document()
    if rounded_like:
        for sec in doc.sections:
            sec.left_margin=Pt(36); sec.right_margin=Pt(36); sec.top_margin=Pt(36); sec.bottom_margin=Pt(36)

    n_rows=len(rows); n_cols=len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    _apply_docx_col_widths(table, widths_cm)  # ← 列幅反映

    hb = header_bg or "#EEEEEE"
    # header
    for c in range(n_cols):
        cell=table.cell(0,c)
        if hb: _set_cell_shading(cell, hb)
        cell.text=str(rows[0][c])
        for p in cell.paragraphs:
            for run in p.runs:
                _apply_font_run(run, font_name, header_size, header_fg, bold=True)
    # body
    zebra_alt="#F7F9FC"; body_bg=(body_bg_val if body_bg_on else None)
    for r in range(1,n_rows):
        for c in range(n_cols):
            cell=table.cell(r,c)
            fill=(zebra_alt if (zebra and r%2==1) else body_bg)
            if fill: _set_cell_shading(cell, fill)
            cell.text=str(rows[r][c])
            for p in cell.paragraphs:
                for run in p.runs:
                    _apply_font_run(run, font_name, base_size, body_fg, bold=False)
    _apply_table_inner_borders(table, vertical=inner_v, horizontal=inner_h, val="single", sz=6, color="000000")
    if outer: _apply_table_outer_border(table, val="single", sz=(16 if rounded_like else 12), color="000000")

    if note_text.strip():
        p=doc.add_paragraph(); run=p.add_run(note_text.strip()); _apply_font_run(run, font_name, base_size, "#444444")

    buf = io.BytesIO(); doc.save(buf)
    st.download_button("📥 Word（.docx）をダウンロード", data=buf.getvalue(),
                       file_name="table_generated.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                       use_container_width=True)

    # HTMLプレビュー
    st.subheader("🔍 作成結果（画面プレビュー）")
    html=_build_html_table(
        rows,
        header_bg=hb, header_fg=header_fg, header_size=(base_size if header_same else base_size+1), header_bold=True,
        body_bg=(body_bg or None), body_fg=body_fg, body_size=base_size,
        font_name=font_name, zebra=zebra,
        inner_h=inner_h, inner_v=inner_v, outer=outer,
        note_text=note_text, col_width_pct=widths_pct
    )
    st.markdown(html, unsafe_allow_html=True)

else:
    st.info("貼り付け → スタイルを選択 → 必要なら列幅モードを『手動』にして列幅を調整 → 「表作成」を押してください。")
