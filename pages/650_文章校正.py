# pages/650_文章校正.py — 解析（校正方針：ページ/行/理由）オンリー極簡版
# ・原文は「1行=1行」を厳密保持してテーブル化（CJK折返し/長語ZWSP）
# ・校正方針はMarkdown表をパースしてPDF/Word上の表に整形
# ・ダウンロード形式は PDF または Word のどちらか（デフォルト PDF）

from __future__ import annotations
from typing import List, Tuple, Dict
from pathlib import Path
import sys
import io
import base64
import datetime as _dt
import streamlit as st
from openai import OpenAI

# ===== 共有ライブラリ（common_lib）をパスに追加 =====
PROJECTS_ROOT = Path(__file__).resolve().parents[3]  # or 3 for pages
if str(PROJECTS_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECTS_ROOT))

# === lib の利用（読込 / 貼り付け整形） ※OCRなし ===
from lib.text_loaders import (
    load_text_generic, extract_pdf_text, load_text_from_paste
)

from lib.proofreading.prompts import (
    MODE_DEFS,
    COMMON_PROMPT,
    get_analyze_instruction,
    build_system_prompt,
)

from lib.proofreading.explanation import render_proof_policy_logic_expander


# ------------------------------------------------------------
# UI定数
# ------------------------------------------------------------
st.set_page_config(page_title="Text Studio / 解析（校正方針）", page_icon="📝", layout="wide")

MODEL_OPTIONS = ["gpt-5-mini", "gpt-5-nano"]
DEFAULT_MODEL = "gpt-5-mini"
DEFAULT_MODE = "通常校正"
LINES_PER_PAGE = 40  # 1ページあたりの表示行数

# セッション初期化
if "chat_model" not in st.session_state:
    st.session_state["chat_model"] = DEFAULT_MODEL
# if "proof_mode" not in st.session_state:
#     st.session_state["proof_mode"] = DEFAULT_MODE

# ------------------------------------------------------------
# 表示ユーティリティ
# ------------------------------------------------------------
def display_pdf_bytes(data: bytes, height: int = 600):
    """Streamlit PDF表示（streamlit[pdf] があれば st.pdf）。なければ iframe 埋め込み。"""
    try:
        st.pdf(data, height=height)  # Streamlit 1.31+ / streamlit[pdf]
    except Exception:
        b64 = base64.b64encode(data).decode()
        st.markdown(
            f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="{height}px"></iframe>',
            unsafe_allow_html=True
        )

def to_numbered_lines(raw: str) -> List[str]:
    return raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")

# def render_preview_with_numbers(lines: List[str], lines_per_page: int) -> str:
#     return "\n".join(f"[{(i//lines_per_page)+1}:{(i%lines_per_page)+1:02d}] {t}" for i, t in enumerate(lines))

def render_preview_with_numbers(lines: List[str], lines_per_page: int) -> str:
    return "\n".join(
        f"[{i+1:04d}] {t}"
        for i, t in enumerate(lines)
    )

# ------------------------------------------------------------
# OpenAIクライアント
# ------------------------------------------------------------
def openai_client() -> OpenAI:
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])



def render_policy_preview(*, mode: str) -> str:
    analyze_base = get_analyze_instruction(mode)

    # st.subheader("🧭 解析プロンプト設定")

    # ★ ここでタブ全体を expander に入れる
    with st.expander("🧭 解析プロンプト設定（クリックで展開）", expanded=False):

        tab3, tab1, tab2 = st.tabs([
            "✍️ 追加プロンプト",
            "🧭 System",
            "📋 共通方針",
        ])

        # --- 解析プロンプト（System） ---
        with tab1:
            st.markdown("#### 🧭 解析プロンプト（解析モード）")
            st.code(analyze_base, language="markdown")

        # --- 共通方針 ---
        with tab2:
            st.markdown("#### 📋 共通方針（毎回付与）")
            st.code(COMMON_PROMPT.strip(), language="markdown")

        # --- 追加プロンプト ---
        with tab3:
            st.markdown("#### ✍️ 追加プロンプト（任意）")
            extra = st.text_area(
                "追加プロンプトを入力",
                key="extra_user_prompt",
                placeholder="例）外来語はカタカナ優先。製品名や固有名詞は原文どおりに保持。",
                height=100,
            ) or ""

            return extra



# ------------------------------------------------------------
# Markdown表 → 配列（頁/行/重要度/原文/修正案/理由）
# ------------------------------------------------------------
def _parse_plan_md_tables(md: str) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    if not md:
        return items

    to_jp = {
        "page": "頁", "line": "行", "issue": "重要度",
        "original": "原文", "suggestion": "修正案", "reason": "理由",
        "頁": "頁", "ページ": "頁", "行": "行",
        "重要度": "重要度", "問題点": "重要度", "原文": "原文",
        "修正案": "修正案", "提案": "修正案", "理由": "理由", "根拠": "理由",
    }
    expected = ["頁", "行", "重要度", "原文", "修正案", "理由"]

    def _norm_head_cell(s: str) -> str:
        key = s.strip().lower()
        return to_jp.get(key, to_jp.get(s.strip(), s.strip()))

    def _row_to_dict(ln: str, cols_jp: List[str]) -> Dict[str, str] | None:
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if len(cells) < 3:
            return None
        row = dict(zip(cols_jp, cells[:len(cols_jp)]))
        for k in expected:
            row.setdefault(k, "")
        return row

    lines = [ln for ln in (l.strip() for l in md.splitlines()) if ln]
    header_seen = False
    cols_jp: List[str] = []

    for ln in lines:
        if not (ln.startswith("|") and "|" in ln):
            continue

        raw_cells = [c.strip() for c in ln.strip("|").split("|")]

        # 区切り行（--- / :--- / ---: / :---: 等）はスキップ
        def _is_md_separator(cell: str) -> bool:
            s = cell.strip().replace("-", "").replace(":", "")
            return s == ""

        if all(_is_md_separator(c) for c in raw_cells):
            continue

        if not header_seen:
            norm = [_norm_head_cell(c) for c in raw_cells]
            if len(set(norm) & set(expected)) >= 3:
                cols_jp = [c if c in expected else "" for c in norm]
                if len(cols_jp) < 6:
                    cols_jp += [""] * (6 - len(cols_jp))
                cols_jp = [cols_jp[i] or expected[i] for i in range(6)]
                header_seen = True
                continue

        if header_seen and cols_jp:
            row = _row_to_dict(ln, cols_jp)
            if row:
                items.append(row)

    return items

# ------------------------------------------------------------
# 解析結果の Word / PDF 生成
# ------------------------------------------------------------
def build_policy_docx_bytes(
    *, original_numbered_preview: str, plan_md: str,
    model: str, mode: str, extra_prompt: str, src_name: str
) -> Tuple[bytes, str]:
    """Word（.docx）でレポート出力：原文=等幅風/自動折返し、校正方針=Markdown文字列そのまま"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        data = "\n".join([
            "=== 校正方針レポート（TXTフォールバック） ===",
            f"生成日時: {_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
            f"モデル: {model}",
            f"モード: {mode}",
            f"追加プロンプト: {extra_prompt.strip()}" if extra_prompt.strip() else "",
            f"入力ファイル: {src_name or '-'}",
            "\n--- 原文（行番号つき） ---\n",
            original_numbered_preview or "(空)",
            "\n--- 校正方針（Markdown） ---\n",
            plan_md or "(なし)",
        ]).encode("utf-8")
        return data, ".txt"

    doc = Document()
    h = doc.add_heading("校正方針レポート", 0)
    try:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    p = doc.add_paragraph()
    p.add_run("生成日時: ").bold = True; p.add_run(_dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p.add_run("\nモデル: ").bold = True; p.add_run(model)
    p.add_run("\nモード: ").bold = True; p.add_run(mode)
    if extra_prompt.strip():
        p.add_run("\n追加プロンプト: ").bold = True; p.add_run(extra_prompt.strip())
    p.add_run("\n入力ファイル: ").bold = True; p.add_run(src_name or "-")

    doc.add_heading("原文（行番号つき）", level=1)
    para = doc.add_paragraph()
    r = para.add_run(original_numbered_preview if original_numbered_preview else "(空)")
    try:
        r.font.name = "Courier New"; r.font.size = Pt(10)
    except Exception:
        pass

    doc.add_heading("校正方針（Markdown）", level=1)
    para2 = doc.add_paragraph()
    r2 = para2.add_run(plan_md if plan_md else "(なし)")
    try:
        r2.font.name = "Courier New"; r2.font.size = Pt(10)
    except Exception:
        pass

    bio = io.BytesIO(); doc.save(bio)
    return bio.getvalue(), ".docx"

def build_policy_pdf_bytes(
    *, original_numbered_preview: str, plan_md: str,
    model: str, mode: str, extra_prompt: str, src_name: str
) -> bytes:
    """整形PDF：原文は1行=1行の行テーブル、校正方針はMarkdown表をパースして表描画"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.lib.units import mm
        import re
    except Exception:
        return b""

    # 日本語CIDフォント
    font_name = None
    for fname in ("HeiseiMin-W3", "HeiseiKakuGo-W5"):
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(fname)); font_name = fname; break
        except Exception:
            continue
    if not font_name:
        return b""

    buf = io.BytesIO()
    pagesize = A4
    margin = 18 * mm
    doc = SimpleDocTemplate(
        buf, pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=20 * mm, bottomMargin=18 * mm,
        title="校正方針レポート",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleJP", fontName=font_name, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Meta", fontName=font_name, fontSize=10, leading=13, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle(name="H1", fontName=font_name, fontSize=13, leading=16, spaceBefore=6, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body", fontName=font_name, fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="MonoCJK", fontName=font_name, fontSize=9.5, leading=13, wordWrap="CJK"))

    text_width = pagesize[0] - doc.leftMargin - doc.rightMargin

    def _page_number(canvas, doc_):
        canvas.setFont(font_name, 9)
        canvas.drawRightString(pagesize[0] - doc_.rightMargin, 12 * mm, f"{doc_.page}")

    # 原文（1行=1行のテーブル表示、CJK折返し＋英数字長語にZWSP）
    ZWSP = "&#8203;"
    LONG_TOKEN = re.compile(r"([A-Za-z0-9_/#%~\-\.\?\=&]{30,})")
    def _soften(s: str) -> str:
        return LONG_TOKEN.sub(lambda m: ZWSP.join(m.group(1)[i:i+30] for i in range(0, len(m.group(1)), 30)), s)

    def _make_original_table(text: str) -> Table:
        """
        原文（行番号付き）のテーブルを作成する。
        ただし，1行があまりにも長い場合は PDF 生成時だけ
        複数の Row に機械的に分割し，1セルの高さが1ページを超えないようにする。
        """
        rows = []

        # 1セルに入れる最大文字数の目安
        # 全角も len() では1カウントなので，かなり安全側に 80 ぐらいにしておきます
        MAX_CHARS_PER_ROW = 80

        for raw in (text or "").splitlines() or ["(空)"]:
            # タブ→スペース
            s = raw.replace("\t", "    ")

            # ★ 超長い行は複数の「物理行」に分割する
            if len(s) > MAX_CHARS_PER_ROW:
                segments = [
                    s[i : i + MAX_CHARS_PER_ROW]
                    for i in range(0, len(s), MAX_CHARS_PER_ROW)
                ]
            else:
                segments = [s]

            for seg in segments:
                # 英数字長語に ZWSP を入れて折返しを助ける
                seg = _soften(seg).replace("  ", "&nbsp;&nbsp;")
                p = Paragraph(seg if seg else "&nbsp;", styles["MonoCJK"])
                rows.append([p])

        # 列幅は従来どおり
        col_w = max(10, text_width - 2)  # 枠重なり防止に2ptマージン
        t = Table(rows, colWidths=[col_w], splitByRow=1)
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    # 校正方針テーブル
    items = _parse_plan_md_tables(plan_md)
    headers = ["行", "重要度", "原文", "修正案", "理由"]
    table_data = [headers] + [[
        it.get("行",""), it.get("重要度",""),
        it.get("原文",""), it.get("修正案",""), it.get("理由","")
    ] for it in items]

    # 列幅
    #   行: 14mm
    #   重要度: 18mm
    #   残りを「原文 30% / 修正案 30% / 理由 40%」に配分して「理由」を広めに
    col_w = [14 * mm, 18 * mm]
    remain = text_width - sum(col_w)
    col_w += [remain * 0.30, remain * 0.30, remain * 0.40]

    def _p(s: str) -> Paragraph:
        s = (s or "")
        s = s.replace("<br>", "<br/>").replace("<br />", "<br/>")
        s = s.replace("\n", "<br/>")
        return Paragraph(s, styles["Body"])

    table_para = [table_data[0]] + [[_p(x) for x in r] for r in table_data[1:]]
    tbl = Table(table_para, colWidths=col_w, repeatRows=1)
    zebra = colors.Color(0.95, 0.95, 0.98)
    header_bg = colors.Color(0.88, 0.90, 0.95)
    grid = colors.Color(0.75, 0.78, 0.85)
    st_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        # ★ 中央寄せは「行」「重要度」の2列だけ
        ("ALIGN", (0, 0), (1, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("GRID", (0, 0), (-1, -1), 0.3, grid),
    ]
    for i in range(1, len(table_para)):
        if i % 2 == 1:
            st_cmds.append(("BACKGROUND", (0, i), (-1, i), zebra))
    tbl.setStyle(TableStyle(st_cmds))

    # 組み立て
    story = []
    story.append(Paragraph("校正方針レポート", styles["TitleJP"]))
    meta_lines = [
        f"生成日時：{_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        f"モデル：{model}",
        f"モード：{mode}",
        f"入力ファイル：{src_name or '-'}",
    ]
    if (extra_prompt or "").strip():
        meta_lines.insert(3, f"追加プロンプト：{extra_prompt.strip()}")
    for ln in meta_lines:
        story.append(Paragraph(ln, styles["Meta"]))

    story.append(Paragraph("原文（行番号つき）", styles["H1"]))
    story.append(_make_original_table(original_numbered_preview or "(空)"))
    story.append(Spacer(1, 6*mm))

    story.append(Paragraph("校正方針（テーブル）", styles["H1"]))
    story.append(tbl)

    doc.build(story, onFirstPage=_page_number, onLaterPages=_page_number)
    pdf_bytes = buf.getvalue(); buf.close()
    return pdf_bytes


def analyze_issues(model: str, lines: List[str], lines_per_page: int, mode: str, extra: str) -> str:
    """
    ページ単位で GPT から Markdown 表を取得し、
    まとめて 1 個の大きな表（ヘッダー 1 回だけ）に組み直して返す。
    """
    client = openai_client()
    md_tables: List[str] = []
    total_pages = (len(lines) + lines_per_page - 1) // lines_per_page

    # System プロンプト
    sys_inst_template = build_system_prompt(mode=mode, extra=extra)

    # --- 1) GPT にページごと投げて Markdown 表を集める ---
    for pg in range(total_pages):
        start = pg * lines_per_page
        end = min((pg + 1) * lines_per_page, len(lines))
        # page_chunk = [
        #     f"[{(i // lines_per_page) + 1}:{(i % lines_per_page) + 1:02d}] {lines[i]}"
        #     for i in range(start, end)
        # ]
        page_chunk = [
            f"[{(i + 1):04d}] {lines[i]}"
            for i in range(start, end)
        ]
        page_text = "\n".join(page_chunk)

        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": sys_inst_template},
                {"role": "user", "content": f"次のテキスト（このページのみ）を解析してください：\n---\n{page_text}"},
            ],
        )
        md_tables.append(resp.choices[0].message.content.strip())

    # --- 2) まとめてパースして items にする ---
    raw_md = "\n\n".join(md_tables)
    items = _parse_plan_md_tables(raw_md)

    # ✅ ヘッダーそっくりな行（頁/行/重要度/原文/修正案/理由）を除外する
    def _is_header_like(row: Dict[str, str]) -> bool:
        def norm(s: str) -> str:
            return (s or "").strip().lower()

        return (
            norm(row.get("頁", "")) in {"頁", "page", "ページ"} and
            norm(row.get("行", "")) in {"行", "line"} and
            norm(row.get("重要度", "")) in {"重要度", "issue", "問題点"} and
            norm(row.get("原文", "")) in {"原文", "original"} and
            norm(row.get("修正案", "")) in {"修正案", "suggestion", "提案"} and
            norm(row.get("理由", "")) in {"理由", "reason", "根拠"}
        )

    items = [row for row in items if not _is_header_like(row)]

    # --- 3) items から「ヘッダー1回だけの Markdown 表」を再構築 ---
    def esc_cell(s: str) -> str:
        s = (s or "").replace("|", r"\|")
        # 修正前: s = s.replace("\n", "<br>")
        s = s.replace("\n", "<br/>")     # 修正後
        return s


    header = "| 行 | 重要度 | 原文 | 修正案 | 理由 |"
    sep    = "| --- | --- | --- | --- | --- |"

    rows = [header, sep]
    for it in items:
        rows.append(
            "| {line} | {imp} | {orig} | {sugg} | {reason} |".format(
                line=esc_cell(it.get("行", "")),
                imp=esc_cell(it.get("重要度", "")),
                orig=esc_cell(it.get("原文", "")),
                sugg=esc_cell(it.get("修正案", "")),
                reason=esc_cell(it.get("理由", "")),
            )
        )

    # この文字列が HTML プレビュー用 plan_md かつ PDF 解析用のソースになる
    return "\n".join(rows)


# ------------------------------------------------------------
# 画面（解析のみ）
# ------------------------------------------------------------
st.title("📝 文章の校正")
st.caption(
    "数段落程度の文章を貼り付けて，解析のボタンを押してください．解析には少し時間がかかります．"
    "校正にはAIを使用していますので，個人情報や機密文書の取り扱いには注意してください。"
)

st.write("""本文を入力して **① 解析** を実行すると、
         /行/理由つきの校正方針（Markdown表）を生成します。
        文章は， **1,000行以内** に区切って校正をかけることを推奨します．通常校正，厳格校正，簡易校正の３種類があります．
        解析モードを変更することで校正の程度を選択できます．""")

st.write("""
Wordの図表を貼り付けた時は，解析モードで **図表校正** を選択してください．""")

# st.write("""<p style='color:red;'>Wordファイルに校正をかけたい</p>ときは，まず **word解析** を行い， **中間テキスト** を作成してください．
#          その中間テキストをこの **校正** の **「ファイルから」タブ** より（「ファイルから」をクリック），
#          （雲のマークの上に）ドロップしてください．または，**「貼り付けテキスト」タブ** よりエリア内にペーストしてください．
#          その際に **解析モード** は <p style='color:red;'>解析文書校正</p> を選択してください．
#          文章は， **1,000行以内** になるように分割して校正をかけることを推奨します．""", unsafe_allow_html=True)

st.markdown(
    """
<div style='line-height:1.6'>
  <span style='color:red; font-size:18px; font-weight:bold;'>
    Wordファイルに校正をかけたい
  </span>
  ときは，まず <b>word解析</b> を行い， <b>中間テキスト</b> を作成してください．
  その中間テキストをこの <b>校正</b> の <b>「ファイルから」タブ</b> より，
  （雲のマークの上に）ドロップしてください．または，
  <b>「貼り付けテキスト」タブ</b> にペーストしてください．
  その際に <b>解析モード</b> は 
  <span style='color:red; font-size:18px; font-weight:bold;'>解析文書校正</span>
  を選択してください．
  文章は，<b>1,000行以内</b> になるように分割して校正することを推奨します．
  <br>
  ※ 将来的には，Wordファイルを直接校正できるように改善する予定です．
  <br><br>
</div>
    """,
    unsafe_allow_html=True
)

# ここでタブ付き expander をまとめて描画
render_proof_policy_logic_expander()


# 1) 初期化（最初の1回だけ）
if "chat_model" not in st.session_state:
    st.session_state["chat_model"] = DEFAULT_MODEL


with st.sidebar:
    st.header("設定")
    st.radio(
        "🧠 使用モデル",
        MODEL_OPTIONS,
        key="chat_model",
    )
    st.selectbox(
        "🛠 解析モード",
        list(MODE_DEFS.keys()),
        index=list(MODE_DEFS.keys()).index(st.session_state.get("proof_mode", DEFAULT_MODE)),
        key="proof_mode",
        help="\n\n".join([f"・{k}: {v['desc']}" for k, v in MODE_DEFS.items()]),
    )
    lpp = st.number_input("ページ行数（表示用）", min_value=20, max_value=100, value=LINES_PER_PAGE, step=5)

    # ダウンロード形式（PDF/Wordのみ。デフォルトPDF）
    _DL_LABELS = {"pdf": "PDF (.pdf)", "word": "Word (.docx)"}
    dl_choice_key = st.radio(
        "📦 ダウンロード形式（解析レポート）",
        options=list(_DL_LABELS.keys()),
        format_func=lambda k: _DL_LABELS[k],
        index=0,  # デフォルト PDF
        key="dl_format_radio",
    )

extra_prompt = render_policy_preview(mode=st.session_state["proof_mode"])
st.markdown("---")

# ===== 入力（ファイル / 貼り付け） =====
tab_paste, tab_file = st.tabs(["📝 貼り付けテキスト","📁 ファイルから"])

src_text = ""
used_file_name = None

# 先に初期値（必要なら一度だけ）
if "pasted_text" not in st.session_state:
    st.session_state["pasted_text"] = ""

# def _on_paste():
#     st.session_state["has_paste"] = bool(st.session_state["pasted_text"].strip())


with tab_file:
    # アップロードエリアを上に配置
    up = st.file_uploader(
        ".docx / .txt / .pdf をアップロード",
        type=["docx", "txt", "pdf"]
    )

     # ← ここで少しだけ下にスペースを入れる
    #st.markdown("<div style='height: 0.5rem;'></div>", unsafe_allow_html=True)


    # モード＋ボタンを下で横並び
    col_mode, col_btn = st.columns([1, 1])

    with col_mode:
        current_mode = st.session_state.get("proof_mode", DEFAULT_MODE)
        st.markdown(
            f"""
            <div style="
                padding:6px 10px;
                border-radius:6px;
                background-color:#ffe9c6;
                color:#8a4b0f;
                font-weight:bold;
                font-size:0.95rem;
                border:1px solid #f0b76a;
                white-space:nowrap;
                display:inline-block;
            ">
                🧭 解析モード：{current_mode}
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col_btn:
        do_analyze_file = st.button(
            "① 解析（ファイル）",
            type="primary",
            use_container_width=True,
            disabled=not up,
            key="btn_analyze_file"
        )


    if up:
        used_file_name = up.name
        name = up.name.lower()
        if name.endswith(".pdf"):
            data = up.read()
            try:
                stats = extract_pdf_text(data)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

            st.subheader("📄 PDFプレビュー")
            display_pdf_bytes(data, height=600)

            if int(stats.get("visible", 0)) < 20:
                st.warning("このPDFは画像PDF（テキスト層なし）と判定しました。OCRツールでテキスト化してから再度お試しください。")
                st.stop()
            else:
                src_text = (stats.get("text") or "").strip()
        else:
            try:
                src_text = load_text_generic(up)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

with tab_paste:
    pasted = st.text_area(
        "ここに本文を貼り付け",
        height=260,
        key="pasted_text",
        placeholder="ここに本文を貼り付けてください（改行は保持されます）。",
    )

    col_mode2, col_btn2 = st.columns([1, 1])

    with col_mode2:
        current_mode = st.session_state.get("proof_mode", DEFAULT_MODE)
        st.markdown(
            f"""
            <div style="
                padding:6px 10px;
                border-radius:6px;
                background-color:#ffe9c6;
                color:#8a4b0f;
                font-weight:bold;
                font-size:0.95rem;
                display:inline-block;
                border:1px solid #f0b76a;
                white-space:nowrap;
            ">
                🧭 解析モード：{current_mode}
            </div>
            """,
            unsafe_allow_html=True,
        )

    with col_btn2:
        do_analyze_paste = st.button(
            "① 解析（貼り付け）",
            type="primary",
            use_container_width=True,
        )



    if do_analyze_paste:
        if not pasted.strip():
            st.warning("テキストを貼り付けてください。")
            st.stop()

        src_text = load_text_from_paste(
            pasted,
            normalize=True,
            collapse_blanks=False,
            keep_blank_lines=1,
            trim_trailing=True,
        )
        used_file_name = "pasted_text.txt"


# ===== 解析の実行 =====
if src_text:
    lines = to_numbered_lines(src_text)
    st.subheader("👀 行番号付きプレビュー（テキスト表示）")
    st.caption(f"表示上のページ行数: {LINES_PER_PAGE} 行/ページ（擬似割り付け）")
    st.text_area("原文（番号付きプレビュー）", value=render_preview_with_numbers(lines, LINES_PER_PAGE), height=260)

    want_analyze = (locals().get("do_analyze_file") or locals().get("do_analyze_paste"))
    if want_analyze:
        with st.spinner("解析中（校正方針を抽出）…"):
            plan_md = analyze_issues(
                st.session_state["chat_model"], lines, LINES_PER_PAGE,
                mode=st.session_state["proof_mode"], extra=extra_prompt
            )
        st.success("解析が完了しました。ページ/行/理由つきで方針を表示します。")
        st.subheader("📋 校正方針（まず何をどう直すか）")
        # st.markdown(plan_md, unsafe_allow_html=False)

        from html import escape

        def md_table_to_html(md: str) -> str:
            lines = [ln.strip() for ln in md.splitlines() if ln.strip()]
            rows = []

            for ln in lines:
                if not ln.startswith("|"):
                    continue
                cells = [c.strip() for c in ln.strip("|").split("|")]
                rows.append(cells)

            if not rows:
                return "<p>（データなし）</p>"

            # ヘッダー行
            header = rows[0]
            body = rows[2:] if len(rows) > 2 else []

            # 列幅指定（理由列を 40%）
            col_html = """
                <colgroup>
                    <col style="width:10%">
                    <col style="width:10%">
                    <col style="width:20%">
                    <col style="width:20%">
                    <col style="width:40%">   <!-- 理由列 -->
                </colgroup>
            """

            html = "<table class='proof-table'>"
            html += col_html

            # ヘッダー
            html += "<thead><tr>"
            for h in header:
                html += f"<th>{escape(h)}</th>"
            html += "</tr></thead>"

            # 本体
            html += "<tbody>"
            for r in body:
                html += "<tr>"
                for c in r:
                    html += f"<td>{c}</td>"
                html += "</tr>"
            html += "</tbody>"

            html += "</table>"
            return html


        # ここで HTML に変換
        html_table = md_table_to_html(plan_md)

        # CSS を適用
        st.markdown("""
        <style>
        .proof-table {
            width: 100%;
            border-collapse: collapse;
        }
        .proof-table th, .proof-table td {
            border: 1px solid #ccc;
            padding: 6px;
            vertical-align: top;
        }
        .proof-table th {
            background: #e8edf7;
            text-align: center;
        }
        </style>
        """, unsafe_allow_html=True)

        st.markdown(html_table, unsafe_allow_html=True)


        # ▼▼ 解析レポートのダウンロード（PDF or Word） ▼▼
        st.markdown("### ⤵️ 解析レポートをダウンロード")
        numbered_preview = render_preview_with_numbers(lines, LINES_PER_PAGE)

        file_base = (used_file_name or "pasted_text").rsplit(".", 1)[0]

        # ★ 解析モード名をファイル名用に整形
        mode_label = st.session_state.get("proof_mode", "").replace(" ", "")
        if mode_label:
            file_stub = f"校正結果_{file_base}_[{mode_label}]"
        else:
            file_stub = f"校正結果_{file_base}"

        if dl_choice_key == "pdf":
            pdf_bytes = build_policy_pdf_bytes(
                original_numbered_preview=numbered_preview,
                plan_md=plan_md,
                model=st.session_state["chat_model"],
                mode=st.session_state["proof_mode"],
                extra_prompt=extra_prompt,
                src_name=used_file_name or "pasted_text.txt",
            )
            if pdf_bytes:
                st.download_button(
                    "PDF（.pdf）として保存",
                    data=pdf_bytes,
                    file_name=f"{file_stub}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"dl_pdf_{file_stub}",
                )
            else:
                st.warning("PDF を生成できませんでした。`pip install reportlab` を実行し、CIDフォント（HeiseiMin/HeiseiKakuGo）が使えるか確認してください。")

        else:  # word
            data_docx, ext = build_policy_docx_bytes(
                original_numbered_preview=numbered_preview,
                plan_md=plan_md,
                model=st.session_state["chat_model"],
                mode=st.session_state["proof_mode"],
                extra_prompt=extra_prompt,
                src_name=used_file_name or "pasted_text.txt",
            )
            st.download_button(
                "Word（.docx）として保存" if ext == ".docx" else "テキスト（.txt）として保存",
                data=data_docx,
                file_name=f"{file_stub}{ext}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == ".docx" else "text/plain",
                use_container_width=True,
                key=f"dl_word_{file_stub}",
            )

else:
    st.info("入力タブ（📁/📝）から本文を指定して『① 解析』を実行してください。")
