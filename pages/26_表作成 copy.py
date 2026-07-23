# -*- coding: utf-8 -*-
# text_studio_app/pages/26_表作成.py
#
# Excel貼り付け → スタイル選択 → Word .docx 生成
# ・画面プレビュー表示（rowspan/colspan 対応）
# ・表下の注テキスト
# ・サイドバーに基本フォントサイズ
# ・列幅モード：均等 / 自動（文字数） / 手動（cm指定）
# ・結合方式：セルに「<同上＞」「<同左＞」が書かれている場合のみ結合
#   - <同上＞ なら上の非マーカーセルと縦結合（連続すれば上端アンカーに伸長）
#   - <同左＞ なら左の非マーカーセルと横結合（連続すれば左端アンカーに伸長）
#   - 同値自動結合は廃止
# ・Word ではセルを物理結合しない（中身を空にして罫線と塗りで見かけ上の結合）
# ・HTML プレビューは rowspan/colspan で実際に結合
# ・複数行ヘッダー対応：先頭 header_rows 行をヘッダー扱いにし、そこでも <同上>/<同左> を有効
# ・左から header_cols 列を「行ヘッダー」としてヘッダー色・フォントを適用

from __future__ import annotations

# ============================================================
# imports（stdlib）
# ============================================================
from io import BytesIO
from pathlib import Path
import sys

# ============================================================
# imports（3rd party）
# ============================================================
import streamlit as st
import pandas as pd

# ============================================================
# ページ設定
# - st.set_page_config は最初に1回だけ実行する
# ============================================================
st.set_page_config(
    page_title="Text Studio / 表作成",
    page_icon="🧾",
    layout="wide",
)

# ============================================================
# パス設定
# - common_lib / app lib を import できるようにする
# ============================================================
_THIS = Path(__file__).resolve()
APP_DIR = _THIS.parents[1]
PROJ_DIR = _THIS.parents[2]
MONO_ROOT = _THIS.parents[3]

for p in (MONO_ROOT, PROJ_DIR, APP_DIR):
    if str(p) not in sys.path:
        sys.path.insert(0, str(p))

PROJECTS_ROOT = MONO_ROOT
APP_NAME = APP_DIR.name
PAGE_NAME = _THIS.stem


from docx import Document
from docx.shared import Pt  # 現状は未使用だが、将来の拡張用に残してもOK


# ============================================================
# common_lib（ページ共通UI）
# - header / banner / theme / login
# ============================================================
from common_lib.ui.page_header import render_standard_page_header

# ============================================================
# lib（表作成説明UI）
# ============================================================
from lib.explanations.exp_table_maker import (
    render_table_maker_page_intro,
    render_table_maker_help_expander,
)


# ヘルパ関数群（lib/table/helpers.py）
from lib.table.helpers import (
    _parse_table as parse_table,
    _compute_col_widths_cm as compute_col_widths_cm,
    _widths_to_pct as widths_to_pct,
    _compute_spans_markers as compute_spans_markers,
    _merge_docx_by_spans as merge_docx_by_spans,
    _apply_docx_col_widths as apply_docx_col_widths,
    _apply_table_borders_robust as apply_table_borders_robust,
    _build_html_table_with_spans as build_html_table_with_spans,
    _apply_font_run as apply_font_run,
)

# プリセット・サンプル（lib/table/presets.py）
from lib.table.presets import PRESETS, EXAMPLE_TEXT


# ============================================================
# 共通ヘッダー
# - settings.toml から BANNER_KEY を取得
# - banner / theme / intro CSS を描画
# - page_session_heartbeat を実行
# - title / subtitle / ログイン状態を描画
# ============================================================
sub, theme, BANNER_KEY, settings = render_standard_page_header(
    st_module=st,
    projects_root=PROJECTS_ROOT,
    app_dir=APP_DIR,
    app_name=APP_NAME,
    page_name=PAGE_NAME,
    title="🧾 表作成",
    subtitle_text="Excelから貼り付けた表をWord・HTML形式に変換",
    default_banner_key="navy_dark",
)

# ============================================================
# ページ説明
# ============================================================
render_table_maker_page_intro()

# ============================================================
# ヘルプ
# ============================================================
render_table_maker_help_expander(
    theme=theme,
)

with st.sidebar:
    st.markdown("### 1) 基本設定")
    example = st.checkbox("サンプルデータを入れる", value=False)
    base_size = st.slider("基本フォントサイズ（pt）", 5, 16, 10, 1)
    header_same = st.checkbox("ヘッダーも同じサイズ（OFFなら +1pt）", value=False)

    st.markdown("### 2) マーカーでセル結合")
    use_up = st.checkbox("「<同上＞」で縦結合を有効化", value=True)
    use_left = st.checkbox("「<同左＞」で横結合を有効化", value=True)
    # ※ ヘッダー行もマーカー対象にするので ignore_header は廃止

st.markdown("**① Excel からコピーした表（TSV/CSV）を貼り付け**")
default_text = EXAMPLE_TEXT if example else ""
text = st.text_area("ここに貼り付け", value=default_text, height=200)

st.markdown("**② 注（表の下に表示・任意）**")
note_text = st.text_area("注をここに貼り付け（任意）", value="", height=100)

st.markdown("**③ ヘッダー行数・ヘッダー列数の指定**")

# ラベル + 入力 を 4 カラムで 1 行に収める
c1, c2, c3, c4,c5 = st.columns([0.5, 0.5, 0.5, 0.5,1])  # 数字を変えると幅の比率を調整できます

with c1:
    st.text("ヘッダー行数")

with c2:
    header_rows = st.number_input(
        "header_rows",
        min_value=0,
        max_value=50,
        value=1,
        step=1,
        label_visibility="collapsed",
        help="先頭からこの行数分がヘッダー扱いになります。ヘッダーなしの場合は 0。",
        key="header_rows",
    )

with c3:
    st.text("ヘッダー列数")

with c4:
    header_cols = st.number_input(
        "header_cols",
        min_value=0,
        max_value=50,
        value=0,
        step=1,
        label_visibility="collapsed",
        help="左端からこの列数分のセルをヘッダーとして扱います。",
        key="header_cols",
    )


st.markdown("**④ スタイル（プリセット → 詳細調整可）**")
c1, c2 = st.columns([1.2, 1])
with c1:
    preset = st.radio("プリセット", list(PRESETS.keys()), index=3)
with c2:
    st.caption("詳細設定で上書き可能")

# プリセット値を初期値として読み出し
_p = PRESETS[preset]
_default_header_bg = _p["header_bg"] or "#EEEEEE"

outer_mode = _p.get("outer_mode", "box")  # 既存プリセットは "box" とみなす


with st.expander("🔧 詳細設定（必要な場合のみ）", expanded=False):
    colA, colB, colC = st.columns(3)
    with colA:
        header_bg = st.color_picker("ヘッダー背景色", _default_header_bg)
        header_fg = st.color_picker("ヘッダー文字色", _p["header_fg"])
        header_bold = st.checkbox("ヘッダー太字", value=_p["header_bold"])
    with colB:
        body_bg_on = st.checkbox("本文 背景を塗る", value=_p["body_bg"] is not None)
        body_bg_val = st.color_picker("本文背景色", _p["body_bg"] or "#FFFFFF")
        body_fg = st.color_picker("本文文字色", _p["body_fg"])
    with colC:
        font_name = st.selectbox(
            "フォント",
            ["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"],
            index=[
                "Meiryo",
                "Yu Gothic",
                "MS PGothic",
                "MS Gothic",
                "Calibri",
                "Arial",
            ].index(_p["font_name"])
            if _p["font_name"] in ["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"]
            else 0,
        )
        inner_h = st.checkbox("横罫線（内側）", value=_p["inner_h"])
        inner_v = st.checkbox("縦罫線（内側）", value=_p["inner_v"])
        outer = st.checkbox("外枠（ボックス）", value=_p["outer"])
        zebra = st.checkbox("ゼブラ行", value=_p["zebra"])

# 詳細設定を開かない場合のデフォルト（未定義ガード）
if "header_bg" not in locals():
    header_bg = _default_header_bg
    header_fg = _p["header_fg"]
    header_bold = _p["header_bold"]
    body_bg_on = _p["body_bg"] is not None
    body_bg_val = _p["body_bg"] or "#FFFFFF"
    body_fg = _p["body_fg"]
    font_name = _p["font_name"]
    inner_h = _p["inner_h"]
    inner_v = _p["inner_v"]
    outer = _p["outer"]
    zebra = _p["zebra"]

# ============== 列幅モード（手動を追加） ==============
col_width_mode = st.radio(
    "列幅モード",
    ["均等", "自動（文字数で可変）", "手動（cm指定）"],
    index=0,
    horizontal=True,
)
TOTAL_CM = 16.0
body_size = int(base_size)
header_size = int(base_size if header_same else base_size + 1)

# 入力テキストから先に列数を推定して手動UIを出す
rows_preview = parse_table(text) if text.strip() else []
n_cols_preview = len(rows_preview[0]) if rows_preview else 0

# 手動幅の状態（セッション）
if "manual_widths" not in st.session_state:
    st.session_state.manual_widths = []

if rows_preview and (
    (not st.session_state.manual_widths)
    or len(st.session_state.manual_widths) != n_cols_preview
):
    # 1列あたり TOTAL_CM/n_cols_preview が 0.5 未満になるケースもあるので max(0.5, ...) でクランプ
    base_width = max(0.5, round(TOTAL_CM / max(n_cols_preview, 1), 2))
    st.session_state.manual_widths = [base_width] * n_cols_preview

if col_width_mode == "手動（cm指定）" and rows_preview:
    with st.expander("✋ 手動で列幅を設定（合計は約16.0cm推奨）", expanded=True):
        ccols = st.columns(min(n_cols_preview, 6))  # 6列ずつ折り返し
        for i in range(n_cols_preview):
            # セッションの値を一度 0.5〜20.0 の範囲に整える
            raw_val = float(st.session_state.manual_widths[i])
            safe_val = max(0.5, min(20.0, raw_val))
            st.session_state.manual_widths[i] = safe_val

            st.session_state.manual_widths[i] = ccols[i % 6].number_input(
                f"列{i+1} 幅(cm)",
                min_value=0.5,
                max_value=20.0,
                value=safe_val,  # ← クランプ済みの値を渡す
                step=0.1,
                key=f"mw_{i}",
            )
        st.write(
            f"合計：**{sum(st.session_state.manual_widths):.2f} cm** / 推奨 {TOTAL_CM:.1f} cm"
        )
        b1, b2, b3 = st.columns(3)
        with b1:
            if st.button("均等にする", use_container_width=True):
                st.session_state.manual_widths = [
                    round(TOTAL_CM / n_cols_preview, 2)
                ] * n_cols_preview
        with b2:
            if st.button("自動案で埋める（文字数）", use_container_width=True) and rows_preview:
                w_list = compute_col_widths_cm(rows_preview, total_cm=TOTAL_CM)
                # ここでも 0.5〜20.0 にクランプ
                st.session_state.manual_widths = [
                    max(0.5, min(20.0, round(x, 2))) for x in w_list
                ]
        with b3:
            if st.button("合計を16cmにスケール", use_container_width=True):
                s = sum(st.session_state.manual_widths) or 1.0
                scaled = [x * TOTAL_CM / s for x in st.session_state.manual_widths]
                st.session_state.manual_widths = [
                    max(0.5, min(20.0, round(x, 2))) for x in scaled
                ]

st.markdown("---")
make_btn = st.button("▶ 表作成（Word .docx を生成＆画面表示）", type="primary", use_container_width=True)

# ============== 作成処理 ==============
if make_btn:
    rows = parse_table(text)
    if not rows:
        st.error("表データが読めませんでした。テキストを確認してください。")
        st.stop()

    # 安全のため：ヘッダー行数・列数は全体サイズ以内に抑える
    header_rows_int = max(0, min(int(header_rows), len(rows)))
    header_cols_int = max(0, min(int(header_cols), len(rows[0])))

    # 列幅決定
    if col_width_mode == "均等":
        widths_cm = [TOTAL_CM / len(rows[0])] * len(rows[0])
    elif col_width_mode == "自動（文字数で可変）":
        widths_cm = compute_col_widths_cm(rows, total_cm=TOTAL_CM)
    else:  # 手動（cm指定）
        if (not st.session_state.manual_widths) or (
            len(st.session_state.manual_widths) != len(rows[0])
        ):
            st.session_state.manual_widths = [TOTAL_CM / len(rows[0])] * len(rows[0])
        widths_cm = st.session_state.manual_widths[:]
    widths_pct = widths_to_pct(widths_cm)

    # プレビュー用DF（「先頭行」をヘッダーにして、データは header_rows_int 行目以降）
    if len(rows) > 1:
        header_row = [str(x) for x in rows[0]]

        # --- 列名の重複を避ける（同じ名前があれば _2, _3, ... を付ける） ---
        seen = {}
        cols_unique = []
        for name in header_row:
            base = name
            if base in seen:
                seen[base] += 1
                cols_unique.append(f"{base}_{seen[base]}")  # 例: "<同左>", "<同左>_2"
            else:
                seen[base] = 0
                cols_unique.append(base)

        df = pd.DataFrame(rows[header_rows_int:], columns=cols_unique)
    else:
        df = pd.DataFrame()

    st.success(
        f"表を読み込みました：{len(rows) - header_rows_int}行 × {len(rows[0])}列"
        f"（ヘッダー行数: {header_rows_int}、ヘッダー列数: {header_cols_int}、列幅モード: {col_width_mode}）"
    )
    st.dataframe(df, use_container_width=True)

    # ===== Word 生成 =====
    doc = Document()

    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)

    # 列幅は先に設定
    apply_docx_col_widths(table, widths_cm)

    # spans 計算（マーカー方式：ヘッダー行も含めて <同上>/<同左> を解釈）
    spans = compute_spans_markers(
        rows,
        use_up=use_up,
        use_left=use_left,
        header_rows=header_rows_int,  # 将来拡張用（現状ロジックでは未使用）
        strict_rect=True,
    )

    # 実セルは結合しない：アンカー以外は空欄＋罫線で見かけだけ結合
    merge_docx_by_spans(
        table,
        rows,
        spans,
        font_name=font_name,
        base_size=int(body_size),
        header_size=int(header_size),
        header_fg=header_fg,
        body_fg=body_fg,
        header_bg=(header_bg or None),
        body_bg=(body_bg_val if body_bg_on else None),
        zebra=zebra,
        header_same=header_same,
        body_bg_on=body_bg_on,
        header_rows=header_rows_int,
        header_cols=header_cols_int,  # ★ 行ヘッダー列を指定
    )

    # 罫線（spans/anchor ベースで、結合ブロック内部には線を引かない）
    apply_table_borders_robust(
        table,
        spans,
        inner_h=inner_h,
        inner_v=inner_v,
        outer=outer,
        outer_mode=outer_mode,   # ★ 追加
        sz_inner=6,
        sz_outer=12,  # ← 常に 12pt に固定
        color="000000",
    )

    # 注（任意）
    if note_text.strip():
        p = doc.add_paragraph()
        run = p.add_run(note_text.strip())
        apply_font_run(run, font_name, int(body_size), "#444444")

    # ---- Word バッファだけ先に用意（ここではまだボタンを出さない）----
    buf_docx = BytesIO()
    doc.save(buf_docx)

    # ===== HTMLプレビュー（rowspan/colspan 反映） =====
    st.subheader("🔍 作成結果（画面プレビュー）")
    html = build_html_table_with_spans(
        rows,
        spans,
        header_rows=header_rows_int,
        header_bg=(header_bg or "#EEEEEE"),
        header_fg=header_fg,
        header_size=int(header_size),
        header_bold=header_bold,
        body_bg=(body_bg_val if body_bg_on else None),
        body_fg=body_fg,
        body_size=int(body_size),
        font_name=font_name,
        zebra=zebra,
        inner_h=inner_h,
        inner_v=inner_v,
        outer=outer,
        outer_mode=outer_mode,   # ★ これが無いとプレビューは絶対に治りません
        note_text=note_text,
        col_width_pct=widths_pct,
        row_header_cols=header_cols_int,  # ★ HTML側にも行ヘッダー列を反映
    )
    st.markdown(html, unsafe_allow_html=True)

    # ===== 出力（Word → HTML の順にボタンを配置） =====
    st.markdown("### 📥 出力")

    # --- Word ダウンロードボタン（プレビュー直下） ---
    st.download_button(
        "📥 Word（.docx）をダウンロード",
        data=buf_docx.getvalue(),
        file_name="table_generated.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    # --- HTML ダウンロードボタン（その下） ---
    html_doc = f"""<!DOCTYPE html>
<html lang="ja">
<head>
<meta charset="UTF-8">
<title>生成された表</title>
<style>
  body {{
    font-family: "{font_name}", -apple-system, BlinkMacSystemFont, "Segoe UI", sans-serif;
    margin: 24px;
  }}
</style>
</head>
<body>
{html}
</body>
</html>
"""
    st.download_button(
        "📄 HTML をダウンロード",
        data=html_doc.encode("utf-8"),
        file_name="table_generated.html",
        mime="text/html",
        use_container_width=True,
    )
