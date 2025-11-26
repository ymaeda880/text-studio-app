# -*- coding: utf-8 -*-
# pages/20_スタイルチェック.py
#
# Word(.docx) のスタイルチェック用ページ
# - Wordファイルをアップロードして解析
# - 段落ごとのフォント / サイズ / スタイルを解析
# - 同一段落内で複数フォントが混在している箇所を検出
# - 文書全体のフォント一覧と文字数を集計
# - ベース明朝／ゴシック以外のフォントを使っている段落を抜き出して表示
# - 結果を画面表示し、レポートを Word / PDF でダウンロード可能

from __future__ import annotations

from io import BytesIO
from collections import Counter
from typing import Dict, Any, List, Optional

import streamlit as st

try:
    from docx import Document
except ImportError:
    Document = None  # type: ignore

# PDF レポート用
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    # 日本語フォントを登録（HeiseiMin-W3：ReportLab に組込みの和文 CID フォント）
    pdfmetrics.registerFont(UnicodeCIDFont("HeiseiMin-W3"))
    HAS_REPORTLAB = True
except Exception:
    HAS_REPORTLAB = False


# =========================================================
# ヘルパー：フォント・サイズの「実効値」を取得
# =========================================================
from typing import Optional

def get_effective_font_name(run, paragraph, doc) -> Optional[str]:
    """
    実際に表示されるフォントにできるだけ近づけるため，
    Word の rFonts (ascii / hAnsi / eastAsia) を優先的に参照する。

    優先順位イメージ：
      - run のテキストに CJK（日本語）が含まれる → eastAsia を最優先
      - 英数字だけ → ascii / hAnsi
      - それでも不明なら，従来どおり python-docx の font.name / style.font.name / Normal にフォールバック
    """

    # ---- 0) この run に CJK（日本語）が含まれているかざっくり判定 ----
    text = run.text or ""
    has_cjk = any(
        ("\u3040" <= ch <= "\u30ff")   # ひらがな・カタカナ
        or ("\u4e00" <= ch <= "\u9fff")  # CJK 統合漢字
        for ch in text
    )

    # ---- 1) XML の rFonts を直接見る ----
    #   <w:rPr><w:rFonts w:ascii="Arial" w:eastAsia="ＭＳ 明朝" ... />
    rPr = getattr(run._element, "rPr", None)
    if rPr is not None:
        rFonts = getattr(rPr, "rFonts", None)
        if rFonts is not None:
            # CJK を含む場合は eastAsia を最優先
            east = getattr(rFonts, "eastAsia", None)
            ascii_ = getattr(rFonts, "ascii", None)
            hAnsi = getattr(rFonts, "hAnsi", None)

            if has_cjk and east:
                return east  # 日本語run → eastAsiaフォントを「実際に使っているフォント」とみなす

            # 英数字中心の場合は ascii / hAnsi を優先
            if ascii_:
                return ascii_
            if hAnsi:
                return hAnsi

            # eastAsia だけ設定されているケースも拾う
            if east:
                return east

    # ---- 2) うまく取れなかった場合は，従来どおり python-docx のプロパティにフォールバック ----

    # 2-1) run 自身のフォント
    if run.font is not None and run.font.name:
        return run.font.name

    # 2-2) run のスタイルに設定されたフォント
    if run.style is not None and getattr(run.style, "font", None) is not None:
        if run.style.font.name:
            return run.style.font.name

    # 2-3) 段落スタイルに設定されたフォント
    if paragraph.style is not None and getattr(paragraph.style, "font", None) is not None:
        if paragraph.style.font.name:
            return paragraph.style.font.name

    # 2-4) Normal スタイルに設定されたフォント
    try:
        normal = doc.styles["Normal"]
        if normal.font.name:
            return normal.font.name
    except Exception:
        pass

    # それでも分からなければ None
    return None



def get_effective_font_size(run, paragraph, doc):
    """
    run → run.style → paragraph.style → Normal の順にフォントサイズを探索。
    見つからなければ None を返す。
    （戻り値は docx.shared.Pt オブジェクトか None）
    """
    if run.font is not None and run.font.size:
        return run.font.size

    if run.style is not None and getattr(run.style, "font", None) is not None:
        if run.style.font.size:
            return run.style.font.size

    if paragraph.style is not None and getattr(paragraph.style, "font", None) is not None:
        if paragraph.style.font.size:
            return paragraph.style.font.size

    try:
        normal = doc.styles["Normal"]
        if normal.font.size:
            return normal.font.size
    except Exception:
        pass

    return None


def _is_mincho_font(name: str) -> bool:
    """フォント名から明朝体かどうかをざっくり判定"""
    if not name:
        return False
    return ("明朝" in name) or ("Mincho" in name)


def _is_gothic_font(name: str) -> bool:
    """フォント名からゴシック体かどうかをざっくり判定"""
    if not name:
        return False
    return ("ゴシック" in name) or ("Gothic" in name)


# =========================================================
# ヘルパー：文書解析
# =========================================================
def analyze_docx(file_bytes: bytes) -> Dict[str, Any]:
    """
    アップロードされた .docx バイナリからスタイル情報を解析して、
    結果を dict で返す。
    """
    if Document is None:
        raise RuntimeError("python-docx がインポートできません。環境にインストールしてください。")

    bio = BytesIO(file_bytes)
    doc = Document(bio)

    paragraphs_info: List[Dict[str, Any]] = []
    font_counter: Counter[str] = Counter()
    font_paragraph_counter: Counter[str] = Counter()
    style_counter: Counter[str] = Counter()

    all_run_fonts: Counter[str] = Counter()
    all_run_sizes: Counter[Any] = Counter()
    body_run_fonts: Counter[str] = Counter()  # Normal スタイルなど「本文」候補用

    # -------- 段落ごとに走査 --------
    for idx, p in enumerate(doc.paragraphs, start=1):
        text = p.text or ""
        style_name = p.style.name if p.style is not None else "(No style)"
        style_counter[style_name] += 1

        para_fonts: List[str] = []
        para_sizes: List[Any] = []
        run_infos: List[Dict[str, Any]] = []
        char_count = len(text)

        for r in p.runs:
            if not r.text:
                continue
            eff_font = get_effective_font_name(r, p, doc)
            eff_size = get_effective_font_size(r, p, doc)
            if eff_font:
                all_run_fonts[eff_font] += len(r.text)
                para_fonts.append(eff_font)
            if eff_size:
                all_run_sizes[eff_size] += len(r.text)
                para_sizes.append(eff_size)

            run_infos.append(
                {
                    "text": r.text,
                    "font": eff_font,
                    "size": eff_size,
                }
            )

        # 本文候補（Normal, 本文, Body Text）のフォント集計
        if style_name in ["Normal", "本文", "Body Text"]:
            for f in para_fonts:
                body_run_fonts[f] += 1

        para_font_set = sorted(set(para_fonts))
        para_size_set = sorted(
            set(para_sizes),
            key=lambda x: float(x.pt) if hasattr(x, "pt") and x else 0.0,
        )

        paragraphs_info.append(
            dict(
                index=idx,
                text=text,
                style=style_name,
                fonts=para_font_set,
                sizes=para_size_set,
                char_count=char_count,
                runs=run_infos,  # ★ run 情報を持たせる
            )
        )

    # -------- ベースフォント（全体）の推定 --------
    if body_run_fonts:
        base_font_name = body_run_fonts.most_common(1)[0][0]
    elif all_run_fonts:
        base_font_name = all_run_fonts.most_common(1)[0][0]
    else:
        base_font_name = None

    if all_run_sizes:
        base_size_obj = all_run_sizes.most_common(1)[0][0]
    else:
        base_size_obj = None

    base_font_size_pt = float(base_size_obj.pt) if base_size_obj is not None and hasattr(base_size_obj, "pt") else None

    # -------- ベース明朝／ゴシックの推定 --------
    base_mincho_font: Optional[str] = None
    base_gothic_font: Optional[str] = None

    if all_run_fonts:
        # 出現頻度順に走査して最初に見つかったものを採用
        for fname, _ in all_run_fonts.most_common():
            if _is_mincho_font(fname):
                base_mincho_font = fname
                break
        for fname, _ in all_run_fonts.most_common():
            if _is_gothic_font(fname):
                base_gothic_font = fname
                break

    # -------- 段落ごとの異常系・注意点を検出 --------
    anomalies: List[Dict[str, Any]] = []

    for info in paragraphs_info:
        idx = info["index"]
        text = info["text"]
        style_name = info["style"]
        fonts = info["fonts"]
        sizes = info["sizes"]

        # 1) 段落内で複数フォントが混在
        if len(fonts) >= 2:
            sizes_pt = [float(s.pt) if hasattr(s, "pt") and s else None for s in sizes]
            anomalies.append(
                dict(
                    kind="mixed_fonts_in_paragraph",
                    kind_label="段落内で複数フォントが混在",
                    reason=(
                        f"この段落では {len(fonts)} 種類のフォント（{', '.join(fonts)}）が使われています。"
                        "本文や見出しは通常、1種類のフォントに統一することを推奨します。"
                    ),
                    index=idx,
                    style=style_name,
                    fonts=fonts,
                    sizes=sizes_pt,
                    text_preview=text[:50],
                )
            )

        # 2) ベースフォント以外が使われている本文段落
        if base_font_name and style_name in ["Normal", "本文", "Body Text"]:
            other_fonts = [f for f in fonts if f != base_font_name]
            if other_fonts:
                anomalies.append(
                    dict(
                        kind="body_paragraph_uses_other_fonts",
                        kind_label="本文段落でベースフォント以外を使用",
                        reason=(
                            f"本文用スタイル（{style_name}）の段落ですが、"
                            f"ベースフォント「{base_font_name}」以外のフォント（{', '.join(other_fonts)}）が使われています。"
                            "本文は同じフォントに揃えることを推奨します。"
                        ),
                        index=idx,
                        style=style_name,
                        fonts=fonts,
                        sizes=[float(s.pt) if hasattr(s, "pt") and s else None for s in sizes],
                        text_preview=text[:50],
                    )
                )

        # 3) 見出しらしいのに Normal スタイル（簡易判定）
        if style_name in ["Normal", "本文", "Body Text"]:
            stripped = text.strip()
            if 0 < len(stripped) <= 30:
                if (
                    stripped[0].isdigit()
                    or stripped.startswith("第")
                    or stripped.startswith("（")
                    or stripped.startswith("(")
                ):
                    anomalies.append(
                        dict(
                            kind="heading_like_but_normal_style",
                            kind_label="見出しらしいが本文スタイルのまま",
                            reason=(
                                "行頭が番号や「第」で始まっており見出しの可能性がありますが、"
                                "スタイルが Normal（本文）になっています。"
                                "見出し用スタイル（見出し1〜3など）を適用することで目次やナビゲーションが使いやすくなります。"
                            ),
                            index=idx,
                            style=style_name,
                            fonts=fonts,
                            sizes=[float(s.pt) if hasattr(s, "pt") and s else None for s in sizes],
                            text_preview=text[:50],
                        )
                    )

    # -------- フォント集計（段落代表フォント × 文字数） --------
    for info in paragraphs_info:
        text = info["text"]
        fonts = info["fonts"]
        char_count = info["char_count"]
        if not fonts or char_count == 0:
            continue
        main_font = fonts[0]
        font_counter[main_font] += char_count
        font_paragraph_counter[main_font] += 1

    font_summary: List[Dict[str, Any]] = []
    for font_name, chars in font_counter.items():
        font_summary.append(
            dict(
                font=font_name,
                char_count=chars,
                paragraphs=font_paragraph_counter[font_name],
            )
        )
    font_summary = sorted(font_summary, key=lambda x: x["char_count"], reverse=True)

    total_chars = sum(info["char_count"] for info in paragraphs_info)

    # -------- ベース明朝／ゴシック以外のフォント使用段落の抽出 --------
    paragraphs_other_fonts: List[Dict[str, Any]] = []
    allowed_fonts = set()
    if base_mincho_font:
        allowed_fonts.add(base_mincho_font)
    if base_gothic_font:
        allowed_fonts.add(base_gothic_font)

    if allowed_fonts:
        for info in paragraphs_info:
            run_infos = info["runs"]
            idx = info["index"]
            style_name = info["style"]

            # 段落内に「許容フォント以外」が含まれるかどうか
            has_other = False
            for r in run_infos:
                f = r["font"]
                if f is None:
                    continue
                if f not in allowed_fonts:
                    has_other = True
                    break

            if not has_other:
                continue

            # 異フォント部分に <> を付けたテキストを構築
            parts: List[str] = []
            for r in run_infos:
                t = r["text"] or ""
                f = r["font"]
                if f is None or f in allowed_fonts:
                    parts.append(t)
                else:
                    parts.append(f"<{t}>")
            marked_text = "".join(parts)

            # 段落内で使われているフォント一覧
            para_fonts = sorted({r["font"] for r in run_infos if r["font"]})

            paragraphs_other_fonts.append(
                dict(
                    index=idx,
                    style=style_name,
                    fonts=para_fonts,
                    marked_text=marked_text,
                )
            )

    result: Dict[str, Any] = dict(
        base_font_name=base_font_name,
        base_font_size_pt=base_font_size_pt,
        base_mincho_font=base_mincho_font,
        base_gothic_font=base_gothic_font,
        paragraphs=paragraphs_info,
        paragraphs_other_fonts=paragraphs_other_fonts,
        anomalies=anomalies,
        font_summary=font_summary,
        style_counter=style_counter,
        total_chars=total_chars,
    )
    return result


# =========================================================
# レポート生成（Word）
# =========================================================
def build_docx_report(analysis: Dict[str, Any]) -> BytesIO:
    if Document is None:
        raise RuntimeError("python-docx がインポートできません。")

    doc = Document()
    doc.add_heading("スタイルチェックレポート", level=1)

    base_font = analysis.get("base_font_name")
    base_size = analysis.get("base_font_size_pt")
    base_mincho = analysis.get("base_mincho_font")
    base_gothic = analysis.get("base_gothic_font")

    p = doc.add_paragraph()
    if base_font or base_size:
        p.add_run("推定ベースフォント（全体）: ")
        bf = p.add_run(f"{base_font or '不明'}")
        bf.bold = True
        if base_size:
            p.add_run(f" / {base_size:.1f} pt")
    else:
        p.add_run("ベースフォント（全体）を推定できませんでした。")

    p2 = doc.add_paragraph()
    p2.add_run("推定ベース明朝フォント: ")
    p2.add_run(base_mincho or "（明朝系フォントが検出されませんでした）").bold = True

    p3 = doc.add_paragraph()
    p3.add_run("推定ベースゴシックフォント: ")
    p3.add_run(base_gothic or "（ゴシック系フォントが検出されませんでした）").bold = True

    doc.add_paragraph(f"総文字数: {analysis.get('total_chars', 0)}")

    # フォント集計表
    doc.add_heading("フォント別 集計", level=2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "フォント名"
    hdr[1].text = "文字数"
    hdr[2].text = "段落数"

    for item in analysis.get("font_summary", []):
        row_cells = table.add_row().cells
        row_cells[0].text = str(item["font"])
        row_cells[1].text = str(item["char_count"])
        row_cells[2].text = str(item["paragraphs"])

    # ベース明朝／ゴシック以外のフォント使用段落
    doc.add_heading("ベース明朝／ゴシック以外のフォント使用箇所", level=2)
    other_paras: List[Dict[str, Any]] = analysis.get("paragraphs_other_fonts", [])
    if not other_paras:
        doc.add_paragraph("ベース明朝／ゴシック以外のフォントを使用している段落は検出されませんでした。")
    else:
        for item in other_paras:
            idx = item["index"]
            style = item["style"]
            fonts = ", ".join(item["fonts"])
            marked_text = item["marked_text"]

            para = doc.add_paragraph(style="List Number")
            head = para.add_run(f"段落 {idx}（スタイル: {style} / フォント: {fonts}）")
            head.bold = True
            para.add_run("\n")
            para.add_run(
                "※ `<>` で囲まれている部分が、ベース明朝／ゴシック以外のフォントで書かれている箇所です。\n"
            )
            para.add_run(marked_text)

    # スタイル使用状況
    doc.add_heading("スタイル使用状況", level=2)
    style_counter: Counter[str] = analysis.get("style_counter", Counter())
    for style_name, cnt in style_counter.most_common():
        doc.add_paragraph(f"{style_name}: {cnt}", style="List Bullet")

    # 異常・注意箇所
    doc.add_heading("異常・注意箇所一覧", level=2)
    anomalies: List[Dict[str, Any]] = analysis.get("anomalies", [])
    if not anomalies:
        doc.add_paragraph("異常・注意箇所は検出されませんでした。")
    else:
        for an in anomalies:
            kind_label = an.get("kind_label", an.get("kind", ""))
            idx = an.get("index")
            style = an.get("style", "")
            fonts = ", ".join(an.get("fonts", []))
            reason = an.get("reason", "")
            text_preview = an.get("text_preview", "").replace("\n", " ")

            para = doc.add_paragraph(style="List Number")
            run_title = para.add_run(f"[{kind_label}] 段落 {idx} ")
            run_title.bold = True
            para.add_run(f"(スタイル: {style}, フォント: {fonts})\n")

            r_reason_label = para.add_run("理由：")
            r_reason_label.bold = True
            para.add_run(reason + "\n")

            r_content_label = para.add_run("内容：")
            r_content_label.bold = True
            para.add_run(text_preview)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# =========================================================
# レポート生成（PDF：簡易版）
# =========================================================
def build_pdf_report(analysis: Dict[str, Any]) -> BytesIO:
    """
    reportlab を使用したテキストベースの簡易 PDF レポート。
    日本語フォント HeiseiMin-W3 を利用して文字化けを防ぐ。
    """
    if not HAS_REPORTLAB:
        raise RuntimeError("reportlab がインポートできません。")

    buffer = BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4

    base_font = analysis.get("base_font_name") or "不明"
    base_size = analysis.get("base_font_size_pt")
    total_chars = analysis.get("total_chars", 0)
    base_mincho = analysis.get("base_mincho_font") or "（明朝系フォントなし）"
    base_gothic = analysis.get("base_gothic_font") or "（ゴシック系フォントなし）"

    y = height - 40
    c.setFont("HeiseiMin-W3", 14)
    c.drawString(40, y, "スタイルチェックレポート")
    y -= 30

    c.setFont("HeiseiMin-W3", 10)
    c.drawString(40, y, f"ベースフォント（全体）: {base_font} / {base_size or '-'} pt")
    y -= 15
    c.drawString(40, y, f"ベース明朝フォント: {base_mincho}")
    y -= 15
    c.drawString(40, y, f"ベースゴシックフォント: {base_gothic}")
    y -= 15
    c.drawString(40, y, f"総文字数: {total_chars}")
    y -= 25

    c.setFont("HeiseiMin-W3", 11)
    c.drawString(40, y, "フォント別 集計")
    y -= 15
    c.setFont("HeiseiMin-W3", 9)
    for item in analysis.get("font_summary", []):
        line = f"- {item['font']}: {item['char_count']}文字 / {item['paragraphs']}段落"
        c.drawString(50, y, line[:80])
        y -= 12
        if y < 60:
            c.showPage()
            y = height - 40
            c.setFont("HeiseiMin-W3", 9)

    # ベース明朝／ゴシック以外のフォント使用段落（先頭数件）
    y -= 10
    c.setFont("HeiseiMin-W3", 11)
    c.drawString(40, y, "ベース明朝／ゴシック以外のフォント使用箇所（抜粋）")
    y -= 15
    c.setFont("HeiseiMin-W3", 9)
    other_paras: List[Dict[str, Any]] = analysis.get("paragraphs_other_fonts", [])[:20]
    if not other_paras:
        c.drawString(50, y, "ベース明朝／ゴシック以外のフォントを使用している段落は検出されませんでした。")
        y -= 15
    else:
        for item in other_paras:
            idx = item["index"]
            fonts = ", ".join(item["fonts"])
            marked_text = item["marked_text"].replace("\n", " ")
            line1 = f"段落 {idx} | フォント: {fonts}"
            c.drawString(50, y, line1[:100])
            y -= 12
            if y < 60:
                c.showPage()
                y = height - 40
                c.setFont("HeiseiMin-W3", 9)
            line2 = f"内容: {marked_text}"
            c.drawString(60, y, line2[:110])
            y -= 16
            if y < 60:
                c.showPage()
                y = height - 40
                c.setFont("HeiseiMin-W3", 9)

    # 異常・注意箇所（簡易）
    y -= 10
    c.setFont("HeiseiMin-W3", 11)
    c.drawString(40, y, "異常・注意箇所（先頭数件）")
    y -= 15
    c.setFont("HeiseiMin-W3", 9)

    anomalies: List[Dict[str, Any]] = analysis.get("anomalies", [])[:50]
    if not anomalies:
        c.drawString(50, y, "異常・注意箇所は検出されませんでした。")
    else:
        for an in anomalies:
            kind_label = an.get("kind_label", an.get("kind", ""))
            idx = an.get("index")
            fonts = ", ".join(an.get("fonts", []))
            reason = an.get("reason", "")
            text_preview = an.get("text_preview", "").replace("\n", " ")

            line1 = f"[{kind_label}] 段落 {idx} | フォント: {fonts}"
            c.drawString(50, y, line1[:100])
            y -= 12
            if y < 60:
                c.showPage()
                y = height - 40
                c.setFont("HeiseiMin-W3", 9)

            line2 = f"理由: {reason}"
            c.drawString(60, y, line2[:110])
            y -= 12
            if y < 60:
                c.showPage()
                y = height - 40
                c.setFont("HeiseiMin-W3", 9)

            line3 = f"内容: {text_preview}"
            c.drawString(60, y, line3[:110])
            y -= 16
            if y < 60:
                c.showPage()
                y = height - 40
                c.setFont("HeiseiMin-W3", 9)

    c.showPage()
    c.save()
    buffer.seek(0)
    return buffer


# =========================================================
# Streamlit UI
# =========================================================
st.set_page_config(
    page_title="スタイルチェック（Word）",
    page_icon="🔎",
    layout="wide",
)

st.title("🔎 Word スタイルチェック")
st.caption("Word(.docx) ファイルのフォント・スタイルのゆらぎを検出します。")

if Document is None:
    st.error("python-docx がインポートできません。環境に `python-docx` をインストールしてください。")
    st.stop()

uploaded_file = st.file_uploader(
    "Word ファイル（.docx）をアップロードしてください",
    type=["docx"],
    accept_multiple_files=False,
)

run_check = st.button("スタイルチェックを実行", type="primary")

if uploaded_file and run_check:
    with st.spinner("スタイル情報を解析中です..."):
        file_bytes = uploaded_file.read()
        analysis = analyze_docx(file_bytes)

    st.session_state["style_check_analysis"] = analysis

    base_font = analysis.get("base_font_name")
    base_size = analysis.get("base_font_size_pt")
    base_mincho = analysis.get("base_mincho_font")
    base_gothic = analysis.get("base_gothic_font")
    total_chars = analysis.get("total_chars", 0)

    st.subheader("1) ベースフォントの推定")
    if base_font or base_size:
        st.write(
            f"- 推定ベースフォント（全体）: **{base_font or '不明'}**"
            + (f" / **{base_size:.1f} pt**" if base_size else "")
        )
    else:
        st.info("ベースフォント（全体）を推定できませんでした。")

    st.write(f"- 推定ベース明朝フォント: **{base_mincho or '（明朝系フォントが検出されませんでした）'}**")
    st.write(f"- 推定ベースゴシックフォント: **{base_gothic or '（ゴシック系フォントが検出されませんでした）'}**")
    st.write(f"- 総文字数: **{total_chars}**")

    # 2) フォント集計
    st.subheader("2) フォント別 集計")
    font_summary = analysis.get("font_summary", [])
    if font_summary:
        st.dataframe(font_summary, use_container_width=True)
    else:
        st.write("フォント集計がありません。")

    # 2.5) ベース明朝／ゴシック以外のフォント使用段落
    st.subheader("2.5) ベース明朝／ゴシック以外のフォント使用箇所（抜書き）")
    other_paras: List[Dict[str, Any]] = analysis.get("paragraphs_other_fonts", [])
    if not (base_mincho or base_gothic):
        st.info("明朝系／ゴシック系フォントが検出できなかったため、このチェックはスキップされています。")
    elif not other_paras:
        st.success("ベース明朝／ゴシック以外のフォントを使用している段落は検出されませんでした。")
    else:
        st.caption("＜＞で囲まれた部分が、ベース明朝／ゴシック以外のフォントで書かれている箇所です。")
        for item in other_paras:
            idx = item["index"]
            style_name = item["style"]
            fonts = ", ".join(item["fonts"])
            marked_text = item["marked_text"]
            with st.expander(f"段落 {idx}（スタイル: {style_name} / フォント: {fonts}）", expanded=False):
                st.code(marked_text, language="text")

    # 3) スタイル使用状況
    st.subheader("3) スタイル使用状況")
    style_counter = analysis.get("style_counter", Counter())
    if style_counter:
        for style_name, cnt in style_counter.most_common():
            st.write(f"- **{style_name}**: {cnt}")
    else:
        st.write("スタイル情報が取得できませんでした。")

    # 4) 異常・注意箇所
    st.subheader("4) 異常・注意箇所")
    anomalies: List[Dict[str, Any]] = analysis.get("anomalies", [])
    if not anomalies:
        st.success("異常・注意箇所は検出されませんでした。")
    else:
        st.warning(f"{len(anomalies)} 件の注意箇所が見つかりました。")

        rows = []
        for an in anomalies:
            rows.append(
                dict(
                    種類=an.get("kind_label", an.get("kind", "")),
                    段落番号=an.get("index"),
                    スタイル=an.get("style", ""),
                    フォント一覧=", ".join(an.get("fonts", [])),
                    理由=an.get("reason", ""),
                    内容プレビュー=an.get("text_preview", "").replace("\n", " "),
                )
            )

        try:
            import pandas as pd

            df_an = pd.DataFrame(rows)
            st.dataframe(df_an, use_container_width=True, height=400)
        except Exception:
            for r in rows:
                st.write(
                    f"- [{r['種類']}] 段落 {r['段落番号']} "
                    f"(スタイル: {r['スタイル']} / フォント: {r['フォント一覧']})"
                )
                st.write(f"　理由: {r['理由']}")
                st.write(f"　内容: {r['内容プレビュー']}")

    # 5) レポートダウンロード
    st.subheader("5) レポートのダウンロード")

    col1, col2 = st.columns(2)
    with col1:
        try:
            docx_buffer = build_docx_report(analysis)
            st.download_button(
                label="📄 Word レポート (.docx) をダウンロード",
                data=docx_buffer,
                file_name="style_check_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        except Exception as e:
            st.error(f"Word レポート生成でエラーが発生しました: {e}")

    with col2:
        if HAS_REPORTLAB:
            try:
                pdf_buffer = build_pdf_report(analysis)
                st.download_button(
                    label="📑 PDF レポート (.pdf) をダウンロード",
                    data=pdf_buffer,
                    file_name="style_check_report.pdf",
                    mime="application/pdf",
                )
            except Exception as e:
                st.error(f"PDF レポート生成でエラーが発生しました: {e}")
        else:
            st.info("PDF レポート生成には reportlab が必要です。（現在は無効）")

else:
    st.info("Word ファイルをアップロードしてから「スタイルチェックを実行」ボタンを押してください。")
