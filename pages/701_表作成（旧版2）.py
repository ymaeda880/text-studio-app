# -*- coding: utf-8 -*-
# pages/701_表作成（旧版2）.py
#
# Excel貼り付け → スタイル選択 → Word .docx 生成
# ・画面プレビュー表示（rowspan/colspan 対応）
# ・表下の注テキスト
# ・サイドバーに基本フォントサイズ
# ・列幅モード：均等 / 自動（文字数） / 手動（cm指定）
# ・結合方式：セルに「<同上＞」「<同左＞」が書かれている場合のみ結合
#   - <同上＞ なら上の非マーカーセルと縦結合（連続すれば上端アンカーに伸長）
#   - <同左＞ なら左の非マーカーセルと横結合（連続すれば左端アンカーに伸長）
#   - 同値自動結合は廃止
# ・Word ではセルを物理結合しない（中身を空にして罫線と塗りで見かけ上の結合）
# ・HTML プレビューは rowspan/colspan で実際に結合

from __future__ import annotations
import io, csv, re, unicodedata
from typing import List, Tuple, Optional

import streamlit as st
import pandas as pd

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# セル内改行マーカー（<改行> / ＜改行＞ の両方に対応）
_BREAK_RE = re.compile(r"[<＜]\s*改行\s*[>＞]")

# ============== ユーティリティ ==============
def _detect_delimiter(text: str) -> str:
    lines = [ln for ln in text.splitlines() if ln.strip()]
    if not lines:
        return "\t"
    head = lines[0]
    if "\t" in head:
        return "\t"
    if head.count(",") >= head.count(";") and head.count(",") > 0:
        return ","
    if head.count(";") > 0:
        return ";"
    if re.search(r"\s{2,}", head):
        return r"\s+"
    return "\t"

def _parse_table(text: str) -> List[List[str]]:
    text = text.strip("\n\r ")
    if not text:
        return []
    delim = _detect_delimiter(text)
    if delim == r"\s+":
        rows = [re.split(r"\s{2,}", ln.strip()) for ln in text.splitlines() if ln.strip()]
    else:
        reader = csv.reader(io.StringIO(text), delimiter=("\t" if delim == "\t" else delim))
        rows = [list(r) for r in reader]
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return []
    maxc = max(len(r) for r in rows)
    rows = [r + [""] * (maxc - len(r)) for r in rows]
    return rows

def _hex_to_rgb(h: str) -> Tuple[int, int, int]:
    s = h.strip().lstrip("#")
    if len(s) == 3:
        s = "".join(c * 2 for c in s)
    return int(s[0:2], 16), int(s[2:4], 16), int(s[4:6], 16)

def _set_cell_shading(cell, hex_color: Optional[str]):
    if not hex_color:
        return
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#").upper())
    tcPr.append(shd)

def _set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("left", "right", "top", "bottom", "insideH", "insideV"):
        if edge in kwargs:
            val, sz, color = kwargs[edge]
            tag = OxmlElement(f"w:{edge}")
            tag.set(qn("w:val"), val)
            tag.set(qn("w:sz"), str(sz))
            tag.set(qn("w:color"), color)
            tcBorders.append(tag)

def _apply_table_outer_border(table, val="single", sz=12, color="000000"):
    rows = table.rows
    cols = table.columns
    for c in range(len(cols)):
        _set_cell_border(rows[0].cells[c], top=(val, sz, color))
        _set_cell_border(rows[-1].cells[c], bottom=(val, sz, color))
    for r in range(len(rows)):
        _set_cell_border(rows[r].cells[0], left=(val, sz, color))
        _set_cell_border(rows[r].cells[-1], right=(val, sz, color))

def _apply_table_inner_borders(table, vertical=True, horizontal=True, val="single", sz=6, color="000000"):
    rows = table.rows
    cols = table.columns
    if horizontal:
        for r in range(len(rows) - 1):
            for c in range(len(cols)):
                _set_cell_border(rows[r].cells[c], bottom=(val, sz, color))
    if vertical:
        for r in range(len(rows)):
            for c in range(len(cols) - 1):
                _set_cell_border(rows[r].cells[c], right=(val, sz, color))

def _apply_table_borders_robust(
    table,
    spans,
    *,
    inner_h=True,
    inner_v=True,
    outer=True,
    sz_inner=6,
    sz_outer=12,
    color="000000",
):
    """
    物理結合なし版の罫線設定：
      - spans からアンカー行列 anchor[r][c] を作成
      - 同じ anchor 同士の境目には内側罫線を引かない（見かけ上の結合）
      - anchor が変わる境目にだけ線を引く
      - 外周はアンカー境界に沿って太線を引く
    """
    R = len(spans)
    if R == 0:
        return
    C = len(spans[0])

    # --- アンカー行列を作る（各セルがどの結合ブロックに属するか） ---
    anchor: List[List[Optional[Tuple[int, int]]]] = [[None for _ in range(C)] for __ in range(R)]
    for r in range(R):
        for c in range(C):
            info = spans[r][c]
            if info["skip"]:
                continue
            rs = info["rowspan"]
            cs = info["colspan"]
            for rr in range(r, r + rs):
                for cc in range(c, c + cs):
                    anchor[rr][cc] = (r, c)

    # -----------------------------
    # 1) 内側 横罫線：anchor が変わるところだけ境界線
    # -----------------------------
    if inner_h:
        for r in range(R - 1):
            for c in range(C):
                if anchor[r][c] != anchor[r + 1][c]:
                    up = table.cell(r, c)
                    down = table.cell(r + 1, c)
                    _set_cell_border(up, bottom=("single", sz_inner, color))
                    _set_cell_border(down, top=("single", sz_inner, color))

    # -----------------------------
    # 2) 内側 縦罫線：anchor が変わるところだけ境界線
    # -----------------------------
    if inner_v:
        for r in range(R):
            for c in range(C - 1):
                if anchor[r][c] != anchor[r][c + 1]:
                    left = table.cell(r, c)
                    right = table.cell(r, c + 1)
                    _set_cell_border(left, right=("single", sz_inner, color))
                    _set_cell_border(right, left=("single", sz_inner, color))

    # -----------------------------
    # 3) 外周：各結合ブロックの外側を太線で囲む
    # -----------------------------
    if outer and anchor:
        # 上端：各列ごとに「上に別ブロック／外側」が来ているセルに top を引く
        for c in range(C):
            for r in range(R):
                if anchor[r][c] is None:
                    continue
                if r == 0 or anchor[r - 1][c] != anchor[r][c]:
                    cell = table.cell(r, c)
                    _set_cell_border(cell, top=("single", sz_outer, color))
                    break

        # 下端
        for c in range(C):
            for r in range(R - 1, -1, -1):
                if anchor[r][c] is None:
                    continue
                if r == R - 1 or anchor[r + 1][c] != anchor[r][c]:
                    cell = table.cell(r, c)
                    _set_cell_border(cell, bottom=("single", sz_outer, color))
                    break

        # 左端
        for r in range(R):
            for c in range(C):
                if anchor[r][c] is None:
                    continue
                if c == 0 or anchor[r][c - 1] != anchor[r][c]:
                    cell = table.cell(r, c)
                    _set_cell_border(cell, left=("single", sz_outer, color))
                    break

        # 右端
        for r in range(R):
            for c in range(C - 1, -1, -1):
                if anchor[r][c] is None:
                    continue
                if c == C - 1 or anchor[r][c + 1] != anchor[r][c]:
                    cell = table.cell(r, c)
                    _set_cell_border(cell, right=("single", sz_outer, color))
                    break

def _apply_font_run(run, font_name: str, size_pt: int, color_hex: str, bold: bool = False):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(size_pt)
    r, g, b = _hex_to_rgb(color_hex)
    run.font.color.rgb = RGBColor(r, g, b)
    run.bold = bold

# ---- 列幅用 ----
def _visual_len(s: str) -> int:
    if s is None:
        return 0
    t = 0
    for ch in str(s):
        t += 2 if unicodedata.east_asian_width(ch) in ("W", "F") else 1
    return t

def _compute_col_widths_cm(rows: List[List[str]], total_cm=16.0, min_cm=2.2, max_cm=8.0) -> List[float]:
    if not rows:
        return []
    n = len(rows[0])
    scores = []
    for c in range(n):
        mx = 0
        for r in range(len(rows)):
            mx = max(mx, _visual_len(rows[r][c]))
        scores.append(int(mx * 1.1) or 1)
    ssum = sum(scores) or 1
    raw = [total_cm * (sc / ssum) for sc in scores]
    clamped = [max(min_cm, min(max_cm, x)) for x in raw]
    # 合計補正
    diff = total_cm - sum(clamped)
    if abs(diff) > 1e-6:
        room = [(max_cm - w) if diff > 0 else (w - min_cm) for w in clamped]
        room_sum = sum(x for x in room if x > 0) or 1
        adj = [(diff * (r / room_sum) if r > 0 else 0) for r in room]
        clamped = [w + a for w, a in zip(clamped, adj)]
    return clamped

def _widths_to_pct(widths_cm: List[float]) -> List[float]:
    tot = sum(widths_cm) or 1
    return [round(w * 100.0 / tot, 2) for w in widths_cm]

def _apply_docx_col_widths(table, widths_cm: List[float]):
    # 固定レイアウト + 表幅 + グリッド + セル幅
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    # 表幅
    tot_cm = sum(widths_cm) or 1.0
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    # 1cm ≒ 567 dxa
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(tot_cm * 567)))
    # グリッド作り直し
    for child in list(table._tbl.iterchildren()):
        if child.tag == qn("w:tblGrid"):
            table._tbl.remove(child)
    grid = OxmlElement("w:tblGrid")
    for wcm in widths_cm:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(int(wcm * 567)))
        grid.append(gc)
    table._tbl.insert(1, grid)
    # 各セル幅
    for c, wcm in enumerate(widths_cm):
        for r in range(len(table.rows)):
            table.cell(r, c).width = Cm(wcm)

# ========= マーカー方式の結合 =========
# マーカー判定（全角/半角<>と空白の揺れに対応）
_MARK_UP_RE   = re.compile(r'^[<＜]\s*同上\s*[>＞]$')
_MARK_LEFT_RE = re.compile(r'^[<＜]\s*同左\s*[>＞]$')

def _is_mark_up(s: Optional[str]) -> bool:
    if s is None: return False
    s = str(s).strip()
    return bool(_MARK_UP_RE.match(s))

def _is_mark_left(s: Optional[str]) -> bool:
    if s is None: return False
    s = str(s).strip()
    return bool(_MARK_LEFT_RE.match(s))

def _compute_spans_markers(
    rows: List[List[str]],
    *,
    use_up: bool = True,      # 「<同上＞」を有効にする
    use_left: bool = True,    # 「<同左＞」を有効にする
    ignore_header: bool = True,
    strict_rect: bool = True  # 横結合は矩形制約（rowspan一致）を要求
):
    """
    マーカー方式の結合規則：
      ・セルが「<同上＞」なら “上の非マーカーセル” に縦結合（連続なら上端まで伸長）
      ・セルが「<同左＞」なら “左の非マーカーセル” に横結合（連続なら左端まで伸長）
    ※ 同値自動結合は行わない
    """
    if not rows:
        return []

    R, C = len(rows), len(rows[0])
    spans = [[{"rowspan": 1, "colspan": 1, "skip": False} for _ in range(C)] for __ in range(R)]
    start_r = 1 if (ignore_header and R >= 1) else 0

    # --- 1) 縦方向（<同上＞）---
    if use_up:
        # 上から下へ走査し、アンカー(非マーカー)のrowspanを伸ばす
        for r in range(start_r, R):
            for c in range(C):
                if not _is_mark_up(rows[r][c]):
                    continue
                # 上方の最も近い「非<同上>」セルをアンカーに（連続マーカー対応）
                a = r - 1
                while a >= start_r and _is_mark_up(rows[a][c]):
                    a -= 1
                if a >= 0 and not _is_mark_up(rows[a][c]):
                    spans[a][c]["rowspan"] += 1
                    spans[r][c]["skip"] = True
                # 先頭行など、アンカーが見つからなければ無視

    # --- 2) 横方向（<同左＞）---
    if use_left:
        # 左から右へ走査し、アンカー(非マーカー)のcolspanを伸ばす
        for r in range(start_r, R):
            for c in range(C):
                if not _is_mark_left(rows[r][c]):
                    continue
                # 左方の最も近い「非<同左>」セルをアンカーに（連続マーカー対応）
                b = c - 1
                while b >= 0 and _is_mark_left(rows[r][b]):
                    b -= 1
                # アンカー存在かつアンカー自体がskipでない（既に別結合の一部でない）
                if b >= 0 and (not _is_mark_left(rows[r][b])) and (not spans[r][b]["skip"]):
                    # 矩形制約：アンカー～現在までの各セルのrowspanが同じならOK
                    if strict_rect:
                        base_rs = spans[r][b]["rowspan"]
                        ok = True
                        for cc in range(b, c):
                            if spans[r][cc]["rowspan"] != base_rs:
                                ok = False; break
                        if not ok:
                            continue
                    spans[r][b]["colspan"] += 1
                    spans[r][c]["skip"] = True
                # 左端やアンカー不在は無視

    return spans


def _merge_docx_by_spans(
    table,
    rows: List[List[str]],
    spans,
    font_name: str,
    base_size: int,
    header_size: int,
    header_fg: str,
    body_fg: str,
    header_bg: Optional[str],
    body_bg: Optional[str],
    zebra: bool,
    header_same: bool,
    body_bg_on: bool,
):
    R, C = len(rows), len(rows[0])

    hb = header_bg or "#EEEEEE"
    zebra_alt = "#F7F9FC"
    body_fill_default = (body_bg if body_bg_on else None)

    # --- 1) 全セル空に ---
    for r in range(R):
        for c in range(C):
            table.cell(r, c).text = ""

    # --- 2) 物理結合は絶対に行わない ---
    # （すべて skip とアンカーで解決する）
    # ★重要：これが罫線問題を完全修復する
    pass

    # --- 3) テキストを anchor セルだけへ（<改行> → Word 段落、空白行なし） ---
    for r in range(R):
        for c in range(C):
            if spans[r][c]["skip"]:
                continue

            cell = table.cell(r, c)
            raw = str(rows[r][c])

            # 改行統一
            raw = raw.replace("\r\n", "\n").replace("\r", "\n")

            # 改行マーカーで分割
            parts_raw = _BREAK_RE.split(raw)

            parts = []
            for p in parts_raw:
                p = p.strip("\n")
                if p.strip() == "":
                    continue
                parts.append(p)

            if not parts:
                parts = [""]

            # 最初の段落
            p0 = cell.paragraphs[0]
            p0.text = parts[0]

            # 段落追加
            for extra in parts[1:]:
                cell.add_paragraph(extra)

    # --- 4) 背景 + フォント ---
    for r in range(R):
        for c in range(C):
            if spans[r][c]["skip"]:
                continue
            cell = table.cell(r, c)

            if r == 0:
                _set_cell_shading(cell, hb)
            else:
                fill = zebra_alt if (zebra and (r % 2 == 1)) else body_fill_default
                if fill:
                    _set_cell_shading(cell, fill)

            for p in cell.paragraphs:
                for run in p.runs:
                    if r == 0:
                        _apply_font_run(run, font_name, header_size, header_fg, bold=True)
                    else:
                        _apply_font_run(run, font_name, base_size, body_fg, bold=False)



def _apply_table_inner_borders_after_merge(table, spans, vertical=True, horizontal=True, val="single", sz=6, color="000000"):
    """（旧版）結合後、見えるセルだけに内側罫線を当てる関数 — 現在は未使用"""
    R, C = len(spans), len(spans[0])
    if horizontal:
        for r in range(R - 1):
            for c in range(C):
                if not spans[r][c]["skip"]:
                    _set_cell_border(table.cell(r, c), bottom=(val, sz, color))
    if vertical:
        for r in range(R):
            for c in range(C - 1):
                if not spans[r][c]["skip"]:
                    _set_cell_border(table.cell(r, c), right=(val, sz, color))

# def _build_html_table_with_spans(
#     rows: List[List[str]],
#     spans,
#     *,
#     header_bg: Optional[str],
#     header_fg: str,
#     header_size: int,
#     header_bold: bool,
#     body_bg: Optional[str],
#     body_fg: str,
#     body_size: int,
#     font_name: str,
#     zebra: bool,
#     inner_h: bool,
#     inner_v: bool,
#     outer: bool,
#     note_text: Optional[str] = None,
#     col_width_pct: Optional[List[float]] = None,
# ) -> str:
#     """rowspan/colspan を反映した HTML テーブルを生成"""
#     if not rows:
#         return ""
#     R, C = len(rows), len(rows[0])

#     border_css = "border-collapse:collapse;"
#     outer_css = "border:1px solid #000;" if outer else "border:0;"

#     if inner_h and inner_v:
#         td_border_css = "border:1px solid #000;"
#     elif inner_h and not inner_v:
#         td_border_css = "border-style:solid; border-color:#000; border-width:1px 0;"
#     elif inner_v and not inner_h:
#         td_border_css = "border-style:solid; border-color:#000; border-width:0 1px;"
#     else:
#         td_border_css = "border:0;"

#     td_base = "padding:6px 10px; vertical-align:top;"
#     th_weight = "font-weight:700;" if header_bold else "font-weight:400;"
#     zebra_alt = "#F7F9FC"

#     html = [f'<table style="{border_css}{outer_css} width:100%;">']

#     if col_width_pct:
#         html.append("<colgroup>")
#         for p in col_width_pct:
#             html.append(f'<col style="width:{p}%">')
#         html.append("</colgroup>")

#     # thead（行0のみ）
#     html.append("<thead><tr>")
#     hb = header_bg or "#EEEEEE"
#     for c in range(C):
#         info = spans[0][c]
#         if info["skip"]:
#             continue
#         attrs = []
#         if info["rowspan"] > 1:
#             attrs.append(f'rowspan="{info["rowspan"]}"')
#         if info["colspan"] > 1:
#             attrs.append(f'colspan="{info["colspan"]}"')
#         html.append(
#             f'<th {" ".join(attrs)} style="{td_base}{td_border_css}'
#             f'background:{hb}; color:{header_fg}; '
#             f'font-family:{font_name}; font-size:{header_size}pt; {th_weight} text-align:left;">'
#             f'{rows[0][c]}'
#             f"</th>"
#         )
#     html.append("</tr></thead>")

#     # tbody（1行目以降）
#     html.append("<tbody>")
#     for r in range(1, R):
#         row_bg = (zebra_alt if (zebra and r % 2 == 1) else (body_bg or "transparent"))
#         html.append(f'<tr style="background:{row_bg};">')
#         for c in range(C):
#             info = spans[r][c]
#             if info["skip"]:
#                 continue
#             attrs = []
#             if info["rowspan"] > 1:
#                 attrs.append(f'rowspan="{info["rowspan"]}"')
#             if info["colspan"] > 1:
#                 attrs.append(f'colspan="{info["colspan"]}"')
#             html.append(
#                 f'<td {" ".join(attrs)} style="{td_base}{td_border_css}'
#                 f'font-family:{font_name}; font-size:{body_size}pt; color:{body_fg}; text-align:left;">'
#                 f'{rows[r][c]}'
#                 f"</td>"
#             )
#         html.append("</tr>")
#     html.append("</tbody></table>")

#     if note_text and note_text.strip():
#         esc = (
#             note_text.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
#         ).replace("\n", "<br/>")
#         html.append(
#             f"<div style='margin-top:8px; font-family:{font_name}; font-size:{body_size}pt; color:#444;'>{esc}</div>"
#         )
#     return "\n".join(html)
def _build_html_table_with_spans(
    rows: List[List[str]],
    spans,
    *,
    header_bg: Optional[str],
    header_fg: str,
    header_size: int,
    header_bold: bool,
    body_bg: Optional[str],
    body_fg: str,
    body_size: int,
    font_name: str,
    zebra: bool,
    inner_h: bool,
    inner_v: bool,
    outer: bool,
    note_text: Optional[str] = None,
    col_width_pct: Optional[List[float]] = None,
) -> str:
    """rowspan/colspan を反映した HTML テーブルを生成（<改行> は <br/> に変換）"""
    if not rows:
        return ""
    R, C = len(rows), len(rows[0])

    border_css = "border-collapse:collapse;"
    outer_css = "border:1px solid #000;" if outer else "border:0;"

    if inner_h and inner_v:
        td_border_css = "border:1px solid #000;"
    elif inner_h and not inner_v:
        td_border_css = "border-style:solid; border-color:#000; border-width:1px 0;"
    elif inner_v and not inner_h:
        td_border_css = "border-style:solid; border-color:#000; border-width:0 1px;"
    else:
        td_border_css = "border:0;"

    td_base = "padding:6px 10px; vertical-align:top;"
    th_weight = "font-weight:700;" if header_bold else "font-weight:400;"
    zebra_alt = "#F7F9FC"

    html = [f'<table style="{border_css}{outer_css} width:100%;">']

    if col_width_pct:
        html.append("<colgroup>")
        for p in col_width_pct:
            html.append(f'<col style="width:{p}%">')
        html.append("</colgroup>")

    # thead（行0のみ）
    html.append("<thead><tr>")
    hb = header_bg or "#EEEEEE"
    for c in range(C):
        info = spans[0][c]
        if info["skip"]:
            continue
        attrs = []
        if info["rowspan"] > 1:
            attrs.append(f'rowspan="{info["rowspan"]}"')
        if info["colspan"] > 1:
            attrs.append(f'colspan="{info["colspan"]}"')

        # --- ヘッダー文字列：<改行> / ＜改行＞ → <br/>
        raw_header = str(rows[0][c])
        header_text = _BREAK_RE.sub("<br/>", raw_header)

        html.append(
            f'<th {" ".join(attrs)} style="{td_base}{td_border_css}'
            f'background:{hb}; color:{header_fg}; '
            f'font-family:{font_name}; font-size:{header_size}pt; {th_weight} text-align:left;">'
            f'{header_text}'
            f"</th>"
        )
    html.append("</tr></thead>")

    # tbody（1行目以降）
    html.append("<tbody>")
    for r in range(1, R):
        row_bg = (zebra_alt if (zebra and r % 2 == 1) else (body_bg or "transparent"))
        html.append(f'<tr style="background:{row_bg};">')
        for c in range(C):
            info = spans[r][c]
            if info["skip"]:
                continue
            attrs = []
            if info["rowspan"] > 1:
                attrs.append(f'rowspan="{info["rowspan"]}"')
            if info["colspan"] > 1:
                attrs.append(f'colspan="{info["colspan"]}"')

            # --- 本文セル：<改行> / ＜改行＞ → <br/>
            raw_body = str(rows[r][c])
            body_text = _BREAK_RE.sub("<br/>", raw_body)

            html.append(
                f'<td {" ".join(attrs)} style="{td_base}{td_border_css}'
                f'font-family:{font_name}; font-size:{body_size}pt; color:{body_fg}; text-align:left;">'
                f'{body_text}'
                f"</td>"
            )
        html.append("</tr>")
    html.append("</tbody></table>")

    # 注テキスト（こちらは従来通り：HTMLエスケープ＋\n→<br/>）
    if note_text and note_text.strip():
        esc = (
            note_text.strip().replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        ).replace("\n", "<br/>")
        html.append(
            f"<div style='margin-top:8px; font-family:{font_name}; font-size:{body_size}pt; color:#444;'>{esc}</div>"
        )
    return "\n".join(html)

# ============== UI ==============
st.set_page_config(page_title="🧾 Word 表作成", page_icon="🧾", layout="wide")
st.title("🧾 Word 表作成（Excel貼り付け → スタイル指定 → .docx ＋ プレビュー）")

with st.sidebar:
    st.markdown("### 1) 基本設定")
    example = st.checkbox("サンプルデータを入れる", value=False)
    base_size = st.slider("基本フォントサイズ（pt）", 8, 16, 10, 1)
    header_same = st.checkbox("ヘッダーも同じサイズ（OFFなら +1pt）", value=False)

    st.markdown("### 2) マーカーでセル結合")
    use_up = st.checkbox("「<同上＞」で縦結合を有効化", value=True)
    use_left = st.checkbox("「<同左＞」で横結合を有効化", value=True)
    ignore_header = st.checkbox("ヘッダー行（1行目）は結合対象に含めない", value=True)

st.markdown("**① Excel からコピーした表（TSV/CSV）を貼り付け**")
default_text = ""
if example:
    default_text = (
        "品目\t分類\t数量\t単価(円)\t金額(円)\n"
        "りんご\t青果\t12\t120\t1440\n"
        "<同上＞\t<同左＞\t8\t120\t960\n"
        "みかん\t青果\t8\t80\t640\n"
        "ぶどう\t青果\t2\t450\t900\n"
        "<同上＞\t<同左＞\t2\t450\t900"
    )
text = st.text_area("ここに貼り付け", value=default_text, height=200)

st.markdown("**② 注（表の下に表示・任意）**")
note_text = st.text_area("注をここに貼り付け（任意）", value="", height=100)

st.markdown("**③ スタイル（プリセット → 詳細調整可）**")
PRESETS = {
    "シンプル（横罫のみ・薄グレー見出し）": dict(
        header_bg="#EEEEEE",
        header_fg="#000000",
        header_bold=True,
        body_bg=None,
        body_fg="#000000",
        font_name="Meiryo",
        inner_h=True,
        inner_v=False,
        outer=True,
        zebra=False,
    ),
    "ゼブラ（横＋縦罫・見出し濃色）": dict(
        header_bg="#2F5597",
        header_fg="#FFFFFF",
        header_bold=True,
        body_bg=None,
        body_fg="#222222",
        font_name="Meiryo",
        inner_h=True,
        inner_v=True,
        outer=True,
        zebra=True,
    ),
    "横罫のみ（モノトーン）": dict(
        header_bg=None,
        header_fg="#000000",
        header_bold=True,
        body_bg=None,
        body_fg="#000000",
        font_name="Meiryo",
        inner_h=True,
        inner_v=False,
        outer=False,
        zebra=False,
    ),
    "ボックス＋見出し色": dict(
        header_bg="#DDEBF7",
        header_fg="#000000",
        header_bold=True,
        body_bg="#FFFFFF",
        body_fg="#000000",
        font_name="Meiryo",
        inner_h=True,
        inner_v=True,
        outer=True,
        zebra=False,
    ),
    "濃色ヘッダー": dict(
        header_bg="#1F4E79",
        header_fg="#FFFFFF",
        header_bold=True,
        body_bg="#FFFFFF",
        body_fg="#222222",
        font_name="Meiryo",
        inner_h=True,
        inner_v=True,
        outer=True,
        zebra=False,
    ),
}

c1, c2 = st.columns([1.2, 1])
with c1:
    preset = st.radio("プリセット", list(PRESETS.keys()), index=3)
with c2:
    st.caption("詳細設定で上書き可能")

# プリセット値を初期値として読み出し
_p = PRESETS[preset]
_default_header_bg = _p["header_bg"] or "#EEEEEE"

with st.expander("🔧 詳細設定（必要な場合のみ）", expanded=False):
    colA, colB, colC = st.columns(3)
    with colA:
        header_bg = st.color_picker("ヘッダー背景色", _default_header_bg)
        header_fg = st.color_picker("ヘッダー文字色", _p["header_fg"])
        header_bold = st.checkbox("ヘッダー太字", value=_p["header_bold"])
    with colB:
        body_bg_on = st.checkbox("本文 背景を塗る", value=_p["body_bg"] is not None)
        body_bg_val = st.color_picker("本文背景色", _p["body_bg"] or "#FFFFFF")
        body_fg = st.color_picker("本文文字色", _p["body_fg"])
    with colC:
        font_name = st.selectbox(
            "フォント",
            ["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"],
            index=["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"].index(_p["font_name"])
            if _p["font_name"] in ["Meiryo", "Yu Gothic", "MS PGothic", "MS Gothic", "Calibri", "Arial"]
            else 0,
        )
        inner_h = st.checkbox("横罫線（内側）", value=_p["inner_h"])
        inner_v = st.checkbox("縦罫線（内側）", value=_p["inner_v"])
        outer = st.checkbox("外枠（ボックス）", value=_p["outer"])
        zebra = st.checkbox("ゼブラ行", value=_p["zebra"])
        rounded_like = st.checkbox("外枠を太め＋余白（角丸風）", value=False)

# 詳細設定を開かない場合のデフォルト（未定義ガード）
if "header_bg" not in locals():
    header_bg = _default_header_bg
    header_fg = _p["header_fg"]
    header_bold = _p["header_bold"]
    body_bg_on = _p["body_bg"] is not None
    body_bg_val = _p["body_bg"] or "#FFFFFF"
    body_fg = _p["body_fg"]
    font_name = _p["font_name"]
    inner_h = _p["inner_h"]
    inner_v = _p["inner_v"]
    outer = _p["outer"]
    zebra = _p["zebra"]
    rounded_like = False

# ============== 列幅モード（手動を追加） ==============
col_width_mode = st.radio("列幅モード", ["均等", "自動（文字数で可変）", "手動（cm指定）"], index=0, horizontal=True)
TOTAL_CM = 16.0
body_size = int(base_size)
header_size = int(base_size if header_same else base_size + 1)

# 入力テキストから先に列数を推定して手動UIを出す
rows_preview = _parse_table(text) if text.strip() else []
n_cols_preview = len(rows_preview[0]) if rows_preview else 0

# 手動幅の状態
if "manual_widths" not in st.session_state:
    st.session_state.manual_widths = []
if rows_preview and (not st.session_state.manual_widths or len(st.session_state.manual_widths) != n_cols_preview):
    st.session_state.manual_widths = [round(TOTAL_CM / max(n_cols_preview, 1), 2)] * n_cols_preview

if col_width_mode == "手動（cm指定）" and rows_preview:
    with st.expander("✋ 手動で列幅を設定（合計は約16.0cm推奨）", expanded=True):
        ccols = st.columns(min(n_cols_preview, 6))  # 6列ずつ折り返し
        for i in range(n_cols_preview):
            st.session_state.manual_widths[i] = ccols[i % 6].number_input(
                f"列{i+1} 幅(cm)", min_value=0.5, max_value=20.0, value=float(st.session_state.manual_widths[i]), step=0.1, key=f"mw_{i}"
            )
        st.write(f"合計：**{sum(st.session_state.manual_widths):.2f} cm** / 推奨 {TOTAL_CM:.1f} cm")
        b1, b2, b3 = st.columns(3)
        with b1:
            if st.button("均等にする", use_container_width=True):
                st.session_state.manual_widths = [round(TOTAL_CM / n_cols_preview, 2)] * n_cols_preview
        with b2:
            if st.button("自動案で埋める（文字数）", use_container_width=True) and rows_preview:
                st.session_state.manual_widths = [round(x, 2) for x in _compute_col_widths_cm(rows_preview, total_cm=TOTAL_CM)]
        with b3:
            if st.button("合計を16cmにスケール", use_container_width=True):
                s = sum(st.session_state.manual_widths) or 1.0
                st.session_state.manual_widths = [round(x * TOTAL_CM / s, 2) for x in st.session_state.manual_widths]

st.markdown("---")
make_btn = st.button("▶ 表作成（Word .docx を生成＆画面表示）", type="primary", use_container_width=True)

# ============== 作成処理 ==============
if make_btn:
    rows = _parse_table(text)
    if not rows:
        st.error("表データが読めませんでした。テキストを確認してください。")
        st.stop()

    # 列幅決定
    if col_width_mode == "均等":
        widths_cm = [TOTAL_CM / len(rows[0])] * len(rows[0])
    elif col_width_mode == "自動（文字数で可変）":
        widths_cm = _compute_col_widths_cm(rows, total_cm=TOTAL_CM)
    else:  # 手動（cm指定）
        if not st.session_state.manual_widths or len(st.session_state.manual_widths) != len(rows[0]):
            st.session_state.manual_widths = [TOTAL_CM / len(rows[0])] * len(rows[0])
        widths_cm = st.session_state.manual_widths[:]
    widths_pct = _widths_to_pct(widths_cm)

    # プレビュー用DF（raw表示。結合はHTML/Wordで反映）
    df = pd.DataFrame(rows[1:], columns=rows[0])
    st.success(f"表を読み込みました：{df.shape[0]}行 × {df.shape[1]}列（列幅モード: {col_width_mode}）")
    st.dataframe(df, use_container_width=True)

    # ===== Word 生成 =====
    doc = Document()
    if rounded_like:
        for sec in doc.sections:
            sec.left_margin = Pt(36)
            sec.right_margin = Pt(36)
            sec.top_margin = Pt(36)
            sec.bottom_margin = Pt(36)

    n_rows = len(rows)
    n_cols = len(rows[0])
    table = doc.add_table(rows=n_rows, cols=n_cols)
    _apply_docx_col_widths(table, widths_cm)  # 列幅は先に設定

    # spans 計算（マーカー方式）
    spans = _compute_spans_markers(
        rows,
        use_up=use_up,
        use_left=use_left,
        ignore_header=ignore_header,
        strict_rect=True,
    )

    # 実セルは結合しない：アンカー以外は空欄＋罫線で見かけだけ結合
    _merge_docx_by_spans(
        table,
        rows,
        spans,
        font_name=font_name,
        base_size=body_size,
        header_size=header_size,
        header_fg=header_fg,
        body_fg=body_fg,
        header_bg=(header_bg or None),
        body_bg=(body_bg_val if body_bg_on else None),
        zebra=zebra,
        header_same=header_same,
        body_bg_on=body_bg_on,
    )

    # 罫線（spans/anchor ベースで、結合ブロック内部には線を引かない）
    _apply_table_borders_robust(
        table, spans,
        inner_h=inner_h,
        inner_v=inner_v,
        outer=outer,
        sz_inner=6,
        sz_outer=(16 if rounded_like else 12),
        color="000000",
    )

    # 注（任意）
    if note_text.strip():
        p = doc.add_paragraph()
        run = p.add_run(note_text.strip())
        _apply_font_run(run, font_name, body_size, "#444444")

    # 保存 & ダウンロード
    buf = io.BytesIO()
    doc.save(buf)
    st.download_button(
        "📥 Word（.docx）をダウンロード",
        data=buf.getvalue(),
        file_name="table_generated.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

    # ===== HTMLプレビュー（rowspan/colspan 反映） =====
    st.subheader("🔍 作成結果（画面プレビュー）")
    html = _build_html_table_with_spans(
        rows,
        spans,
        header_bg=(header_bg or "#EEEEEE"),
        header_fg=header_fg,
        header_size=header_size,
        header_bold=True,
        body_bg=(body_bg_val if body_bg_on else None),
        body_fg=body_fg,
        body_size=body_size,
        font_name=font_name,
        zebra=zebra,
        inner_h=inner_h,
        inner_v=inner_v,
        outer=outer,
        note_text=note_text,
        col_width_pct=widths_pct,
    )
    st.markdown(html, unsafe_allow_html=True)
