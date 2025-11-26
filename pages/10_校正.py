# pages/10_æ ¡æ­£.py â€” è§£æï¼ˆæ ¡æ­£æ–¹é‡ï¼šãƒšãƒ¼ã‚¸/è¡Œ/ç†ç”±ï¼‰ã‚ªãƒ³ãƒªãƒ¼æ¥µç°¡ç‰ˆ
# ãƒ»åŸæ–‡ã¯ã€Œ1è¡Œ=1è¡Œã€ã‚’å³å¯†ä¿æŒã—ã¦ãƒ†ãƒ¼ãƒ–ãƒ«åŒ–ï¼ˆCJKæŠ˜è¿”ã—/é•·èªZWSPï¼‰
# ãƒ»æ ¡æ­£æ–¹é‡ã¯Markdownè¡¨ã‚’ãƒ‘ãƒ¼ã‚¹ã—ã¦PDF/Wordä¸Šã®è¡¨ã«æ•´å½¢
# ãƒ»ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰å½¢å¼ã¯ PDF ã¾ãŸã¯ Word ã®ã©ã¡ã‚‰ã‹ï¼ˆãƒ‡ãƒ•ã‚©ãƒ«ãƒˆ PDFï¼‰

from __future__ import annotations
from typing import List, Tuple, Dict
from pathlib import Path
import sys
import io
import base64
import datetime as _dt
import streamlit as st
from openai import OpenAI

# ===== å…±æœ‰ãƒ©ã‚¤ãƒ–ãƒ©ãƒªï¼ˆcommon_libï¼‰ã‚’ãƒ‘ã‚¹ã«è¿½åŠ  =====
PROJECTS_ROOT = Path(__file__).resolve().parents[3]  # or 3 for pages
if str(PROJECTS_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECTS_ROOT))

# === lib ã®åˆ©ç”¨ï¼ˆèª­è¾¼ / è²¼ã‚Šä»˜ã‘æ•´å½¢ï¼‰ â€»OCRãªã— ===
from lib.text_loaders import (
    load_text_generic, extract_pdf_text, load_text_from_paste
)

from lib.proofreading.prompts import (
    MODE_DEFS,
    COMMON_PROMPT,
    get_analyze_instruction,
    build_system_prompt,
)

from lib.proofreading.explanation import render_proof_policy_logic_expander


# ------------------------------------------------------------
# UIå®šæ•°
# ------------------------------------------------------------
st.set_page_config(page_title="Text Studio / è§£æï¼ˆæ ¡æ­£æ–¹é‡ï¼‰", page_icon="ğŸ“", layout="wide")

MODEL_OPTIONS = ["gpt-5-mini", "gpt-5-nano"]
DEFAULT_MODEL = "gpt-5-mini"
DEFAULT_MODE = "é€šå¸¸æ ¡æ­£"
LINES_PER_PAGE = 40  # 1ãƒšãƒ¼ã‚¸ã‚ãŸã‚Šã®è¡¨ç¤ºè¡Œæ•°

# ã‚»ãƒƒã‚·ãƒ§ãƒ³åˆæœŸåŒ–
if "chat_model" not in st.session_state:
    st.session_state["chat_model"] = DEFAULT_MODEL
# if "proof_mode" not in st.session_state:
#     st.session_state["proof_mode"] = DEFAULT_MODE

# ------------------------------------------------------------
# è¡¨ç¤ºãƒ¦ãƒ¼ãƒ†ã‚£ãƒªãƒ†ã‚£
# ------------------------------------------------------------
def display_pdf_bytes(data: bytes, height: int = 600):
    """Streamlit PDFè¡¨ç¤ºï¼ˆstreamlit[pdf] ãŒã‚ã‚Œã° st.pdfï¼‰ã€‚ãªã‘ã‚Œã° iframe åŸ‹ã‚è¾¼ã¿ã€‚"""
    try:
        st.pdf(data, height=height)  # Streamlit 1.31+ / streamlit[pdf]
    except Exception:
        b64 = base64.b64encode(data).decode()
        st.markdown(
            f'<iframe src="data:application/pdf;base64,{b64}" width="100%" height="{height}px"></iframe>',
            unsafe_allow_html=True
        )

def to_numbered_lines(raw: str) -> List[str]:
    return raw.replace("\r\n", "\n").replace("\r", "\n").split("\n")

def render_preview_with_numbers(lines: List[str], lines_per_page: int) -> str:
    return "\n".join(f"[{(i//lines_per_page)+1}:{(i%lines_per_page)+1:02d}] {t}" for i, t in enumerate(lines))

# ------------------------------------------------------------
# OpenAIã‚¯ãƒ©ã‚¤ã‚¢ãƒ³ãƒˆ
# ------------------------------------------------------------
def openai_client() -> OpenAI:
    return OpenAI(api_key=st.secrets["OPENAI_API_KEY"])




def render_policy_preview(*, mode: str) -> str:
    analyze_base = get_analyze_instruction(mode)

    st.subheader("ğŸ§­ è§£æãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆSystemï¼‰")
    #st.code(analyze_base, language="markdown")
    # --- è§£æãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆSystemï¼‰ã‚’æŠ˜ã‚ŠãŸãŸã¿ ---
    with st.expander("ğŸ§­ è§£æãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆè§£æãƒ¢ãƒ¼ãƒ‰ï¼‰", expanded=False):
        st.code(analyze_base, language="markdown")

    # --- å…±é€šæ–¹é‡ã‚’æŠ˜ã‚ŠãŸãŸã¿ ---
    with st.expander("ğŸ“‹ å…±é€šæ–¹é‡ï¼ˆæ¯å›ä»˜ä¸ï¼‰", expanded=False):
        st.code(COMMON_PROMPT.strip(), language="markdown")

    # --- è¿½åŠ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’æŠ˜ã‚ŠãŸãŸã¿ ---
    with st.expander("âœï¸ è¿½åŠ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼ˆä»»æ„ï¼‰", expanded=False):
        return st.text_area(
            "è¿½åŠ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’å…¥åŠ›",
            key="extra_user_prompt",
            placeholder="ä¾‹ï¼‰å¤–æ¥èªã¯ã‚«ã‚¿ã‚«ãƒŠå„ªå…ˆã€‚è£½å“åã‚„å›ºæœ‰åè©ã¯åŸæ–‡ã©ãŠã‚Šã«ä¿æŒã€‚",
            height=100,
        ) or ""


# ------------------------------------------------------------
# Markdownè¡¨ â†’ é…åˆ—ï¼ˆé /è¡Œ/é‡è¦åº¦/åŸæ–‡/ä¿®æ­£æ¡ˆ/ç†ç”±ï¼‰
# ------------------------------------------------------------
def _parse_plan_md_tables(md: str) -> List[Dict[str, str]]:
    items: List[Dict[str, str]] = []
    if not md:
        return items

    to_jp = {
        "page": "é ", "line": "è¡Œ", "issue": "é‡è¦åº¦",
        "original": "åŸæ–‡", "suggestion": "ä¿®æ­£æ¡ˆ", "reason": "ç†ç”±",
        "é ": "é ", "ãƒšãƒ¼ã‚¸": "é ", "è¡Œ": "è¡Œ",
        "é‡è¦åº¦": "é‡è¦åº¦", "å•é¡Œç‚¹": "é‡è¦åº¦", "åŸæ–‡": "åŸæ–‡",
        "ä¿®æ­£æ¡ˆ": "ä¿®æ­£æ¡ˆ", "ææ¡ˆ": "ä¿®æ­£æ¡ˆ", "ç†ç”±": "ç†ç”±", "æ ¹æ‹ ": "ç†ç”±",
    }
    expected = ["é ", "è¡Œ", "é‡è¦åº¦", "åŸæ–‡", "ä¿®æ­£æ¡ˆ", "ç†ç”±"]

    def _norm_head_cell(s: str) -> str:
        key = s.strip().lower()
        return to_jp.get(key, to_jp.get(s.strip(), s.strip()))

    def _row_to_dict(ln: str, cols_jp: List[str]) -> Dict[str, str] | None:
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if len(cells) < len(cols_jp):
            return None
        row = dict(zip(cols_jp, cells[:len(cols_jp)]))
        for k in expected:
            row.setdefault(k, "")
        return row

    lines = [ln for ln in (l.strip() for l in md.splitlines()) if ln]
    header_seen = False
    cols_jp: List[str] = []

    for ln in lines:
        if not (ln.startswith("|") and "|" in ln):
            continue

        raw_cells = [c.strip() for c in ln.strip("|").split("|")]

        # åŒºåˆ‡ã‚Šè¡Œï¼ˆ--- / :--- / ---: / :---: ç­‰ï¼‰ã¯ã‚¹ã‚­ãƒƒãƒ—
        def _is_md_separator(cell: str) -> bool:
            s = cell.strip().replace("-", "").replace(":", "")
            return s == ""

        if all(_is_md_separator(c) for c in raw_cells):
            continue

        if not header_seen:
            norm = [_norm_head_cell(c) for c in raw_cells]
            if len(set(norm) & set(expected)) >= 3:
                cols_jp = [c if c in expected else "" for c in norm]
                if len(cols_jp) < 6:
                    cols_jp += [""] * (6 - len(cols_jp))
                cols_jp = [cols_jp[i] or expected[i] for i in range(6)]
                header_seen = True
                continue

        if header_seen and cols_jp:
            row = _row_to_dict(ln, cols_jp)
            if row:
                items.append(row)

    return items

# ------------------------------------------------------------
# è§£æçµæœã® Word / PDF ç”Ÿæˆ
# ------------------------------------------------------------
def build_policy_docx_bytes(
    *, original_numbered_preview: str, plan_md: str,
    model: str, mode: str, extra_prompt: str, src_name: str
) -> Tuple[bytes, str]:
    """Wordï¼ˆ.docxï¼‰ã§ãƒ¬ãƒãƒ¼ãƒˆå‡ºåŠ›ï¼šåŸæ–‡=ç­‰å¹…é¢¨/è‡ªå‹•æŠ˜è¿”ã—ã€æ ¡æ­£æ–¹é‡=Markdownæ–‡å­—åˆ—ãã®ã¾ã¾"""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        data = "\n".join([
            "=== æ ¡æ­£æ–¹é‡ãƒ¬ãƒãƒ¼ãƒˆï¼ˆTXTãƒ•ã‚©ãƒ¼ãƒ«ãƒãƒƒã‚¯ï¼‰ ===",
            f"ç”Ÿæˆæ—¥æ™‚: {_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
            f"ãƒ¢ãƒ‡ãƒ«: {model}",
            f"ãƒ¢ãƒ¼ãƒ‰: {mode}",
            f"è¿½åŠ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆ: {extra_prompt.strip()}" if extra_prompt.strip() else "",
            f"å…¥åŠ›ãƒ•ã‚¡ã‚¤ãƒ«: {src_name or '-'}",
            "\n--- åŸæ–‡ï¼ˆè¡Œç•ªå·ã¤ãï¼‰ ---\n",
            original_numbered_preview or "(ç©º)",
            "\n--- æ ¡æ­£æ–¹é‡ï¼ˆMarkdownï¼‰ ---\n",
            plan_md or "(ãªã—)",
        ]).encode("utf-8")
        return data, ".txt"

    doc = Document()
    h = doc.add_heading("æ ¡æ­£æ–¹é‡ãƒ¬ãƒãƒ¼ãƒˆ", 0)
    try:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        pass

    p = doc.add_paragraph()
    p.add_run("ç”Ÿæˆæ—¥æ™‚: ").bold = True; p.add_run(_dt.datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    p.add_run("\nãƒ¢ãƒ‡ãƒ«: ").bold = True; p.add_run(model)
    p.add_run("\nãƒ¢ãƒ¼ãƒ‰: ").bold = True; p.add_run(mode)
    if extra_prompt.strip():
        p.add_run("\nè¿½åŠ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆ: ").bold = True; p.add_run(extra_prompt.strip())
    p.add_run("\nå…¥åŠ›ãƒ•ã‚¡ã‚¤ãƒ«: ").bold = True; p.add_run(src_name or "-")

    doc.add_heading("åŸæ–‡ï¼ˆè¡Œç•ªå·ã¤ãï¼‰", level=1)
    para = doc.add_paragraph()
    r = para.add_run(original_numbered_preview if original_numbered_preview else "(ç©º)")
    try:
        r.font.name = "Courier New"; r.font.size = Pt(10)
    except Exception:
        pass

    doc.add_heading("æ ¡æ­£æ–¹é‡ï¼ˆMarkdownï¼‰", level=1)
    para2 = doc.add_paragraph()
    r2 = para2.add_run(plan_md if plan_md else "(ãªã—)")
    try:
        r2.font.name = "Courier New"; r2.font.size = Pt(10)
    except Exception:
        pass

    bio = io.BytesIO(); doc.save(bio)
    return bio.getvalue(), ".docx"

def build_policy_pdf_bytes(
    *, original_numbered_preview: str, plan_md: str,
    model: str, mode: str, extra_prompt: str, src_name: str
) -> bytes:
    """æ•´å½¢PDFï¼šåŸæ–‡ã¯1è¡Œ=1è¡Œã®è¡Œãƒ†ãƒ¼ãƒ–ãƒ«ã€æ ¡æ­£æ–¹é‡ã¯Markdownè¡¨ã‚’ãƒ‘ãƒ¼ã‚¹ã—ã¦è¡¨æç”»"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_LEFT, TA_CENTER
        from reportlab.lib import colors
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
        from reportlab.lib.units import mm
        import re
    except Exception:
        return b""

    # æ—¥æœ¬èªCIDãƒ•ã‚©ãƒ³ãƒˆ
    font_name = None
    for fname in ("HeiseiMin-W3", "HeiseiKakuGo-W5"):
        try:
            pdfmetrics.registerFont(UnicodeCIDFont(fname)); font_name = fname; break
        except Exception:
            continue
    if not font_name:
        return b""

    buf = io.BytesIO()
    pagesize = A4
    margin = 18 * mm
    doc = SimpleDocTemplate(
        buf, pagesize=pagesize,
        leftMargin=margin, rightMargin=margin,
        topMargin=20 * mm, bottomMargin=18 * mm,
        title="æ ¡æ­£æ–¹é‡ãƒ¬ãƒãƒ¼ãƒˆ",
    )

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleJP", fontName=font_name, fontSize=18, leading=22, alignment=TA_CENTER, spaceAfter=8))
    styles.add(ParagraphStyle(name="Meta", fontName=font_name, fontSize=10, leading=13, alignment=TA_LEFT, spaceAfter=6))
    styles.add(ParagraphStyle(name="H1", fontName=font_name, fontSize=13, leading=16, spaceBefore=6, spaceAfter=6))
    styles.add(ParagraphStyle(name="Body", fontName=font_name, fontSize=10, leading=14))
    styles.add(ParagraphStyle(name="MonoCJK", fontName=font_name, fontSize=9.5, leading=13, wordWrap="CJK"))

    text_width = pagesize[0] - doc.leftMargin - doc.rightMargin

    def _page_number(canvas, doc_):
        canvas.setFont(font_name, 9)
        canvas.drawRightString(pagesize[0] - doc_.rightMargin, 12 * mm, f"{doc_.page}")

    # åŸæ–‡ï¼ˆ1è¡Œ=1è¡Œã®ãƒ†ãƒ¼ãƒ–ãƒ«è¡¨ç¤ºã€CJKæŠ˜è¿”ã—ï¼‹è‹±æ•°å­—é•·èªã«ZWSPï¼‰
    ZWSP = "&#8203;"
    LONG_TOKEN = re.compile(r"([A-Za-z0-9_/#%~\-\.\?\=&]{30,})")
    def _soften(s: str) -> str:
        return LONG_TOKEN.sub(lambda m: ZWSP.join(m.group(1)[i:i+30] for i in range(0, len(m.group(1)), 30)), s)

    def _make_original_table(text: str) -> Table:
        rows = []
        for raw in (text or "").splitlines() or ["(ç©º)"]:
            s = raw.replace("\t", "    ")
            s = _soften(s).replace("  ", "&nbsp;&nbsp;")
            p = Paragraph(s if s else "&nbsp;", styles["MonoCJK"])
            rows.append([p])
        col_w = max(10, text_width - 2)  # æ é‡ãªã‚Šé˜²æ­¢ã«2ptãƒãƒ¼ã‚¸ãƒ³
        t = Table(rows, colWidths=[col_w], splitByRow=1)
        t.setStyle(TableStyle([
            ("BOX", (0, 0), (-1, -1), 0.4, colors.grey),
            ("BACKGROUND", (0, 0), (-1, -1), colors.whitesmoke),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ]))
        return t

    # æ ¡æ­£æ–¹é‡ãƒ†ãƒ¼ãƒ–ãƒ«
    items = _parse_plan_md_tables(plan_md)
    headers = ["é ", "è¡Œ", "é‡è¦åº¦", "åŸæ–‡", "ä¿®æ­£æ¡ˆ", "ç†ç”±"]
    table_data = [headers] + [[
        it.get("é ",""), it.get("è¡Œ",""), it.get("é‡è¦åº¦",""),
        it.get("åŸæ–‡",""), it.get("ä¿®æ­£æ¡ˆ",""), it.get("ç†ç”±","")
    ] for it in items]

    # åˆ—å¹…
    col_w = [14*mm, 14*mm, 18*mm]
    remain = text_width - sum(col_w)
    w_long = remain / 3
    col_w += [w_long, w_long, w_long]
    # col_w += [remain*0.36, remain*0.32, remain*0.32]

    def _p(s: str) -> Paragraph:
        return Paragraph((s or "").replace("\n", "<br/>"), styles["Body"])

    table_para = [table_data[0]] + [[_p(x) for x in r] for r in table_data[1:]]
    tbl = Table(table_para, colWidths=col_w, repeatRows=1)
    zebra = colors.Color(0.95, 0.95, 0.98)
    header_bg = colors.Color(0.88, 0.90, 0.95)
    grid = colors.Color(0.75, 0.78, 0.85)
    st_cmds = [
        ("BACKGROUND", (0, 0), (-1, 0), header_bg),
        ("ALIGN", (0, 0), (2, -1), "CENTER"),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTNAME", (0, 0), (-1, -1), font_name),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("GRID", (0, 0), (-1, -1), 0.3, grid),
    ]
    for i in range(1, len(table_para)):
        if i % 2 == 1:
            st_cmds.append(("BACKGROUND", (0, i), (-1, i), zebra))
    tbl.setStyle(TableStyle(st_cmds))

    # çµ„ã¿ç«‹ã¦
    story = []
    story.append(Paragraph("æ ¡æ­£æ–¹é‡ãƒ¬ãƒãƒ¼ãƒˆ", styles["TitleJP"]))
    meta_lines = [
        f"ç”Ÿæˆæ—¥æ™‚ï¼š{_dt.datetime.now():%Y-%m-%d %H:%M:%S}",
        f"ãƒ¢ãƒ‡ãƒ«ï¼š{model}",
        f"ãƒ¢ãƒ¼ãƒ‰ï¼š{mode}",
        f"å…¥åŠ›ãƒ•ã‚¡ã‚¤ãƒ«ï¼š{src_name or '-'}",
    ]
    if (extra_prompt or "").strip():
        meta_lines.insert(3, f"è¿½åŠ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆï¼š{extra_prompt.strip()}")
    for ln in meta_lines:
        story.append(Paragraph(ln, styles["Meta"]))

    story.append(Paragraph("åŸæ–‡ï¼ˆè¡Œç•ªå·ã¤ãï¼‰", styles["H1"]))
    story.append(_make_original_table(original_numbered_preview or "(ç©º)"))
    story.append(Spacer(1, 6*mm))

    story.append(Paragraph("æ ¡æ­£æ–¹é‡ï¼ˆãƒ†ãƒ¼ãƒ–ãƒ«ï¼‰", styles["H1"]))
    story.append(tbl)

    doc.build(story, onFirstPage=_page_number, onLaterPages=_page_number)
    pdf_bytes = buf.getvalue(); buf.close()
    return pdf_bytes

# ------------------------------------------------------------
# è§£æï¼ˆæ ¡æ­£æ–¹é‡ã®ã¿ï¼‰
# ------------------------------------------------------------
# def analyze_issues(model: str, lines: List[str], lines_per_page: int, mode: str, extra: str) -> str:
#     client = openai_client()
#     md_tables: List[str] = []
#     total_pages = (len(lines) + lines_per_page - 1) // lines_per_page

#     # â˜… ã“ã“ã ã‘å¤‰æ›´ï¼šãƒ¢ãƒ¼ãƒ‰ï¼‹è¿½åŠ ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‹ã‚‰ System ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆã‚’ç”Ÿæˆ
#     sys_inst_template = build_system_prompt(mode=mode, extra=extra)

#     for pg in range(total_pages):
#         start = pg * lines_per_page
#         end = min((pg + 1) * lines_per_page, len(lines))
#         page_chunk = [f"[{(i//lines_per_page)+1}:{(i%lines_per_page)+1:02d}] {lines[i]}" for i in range(start, end)]
#         page_text = "\n".join(page_chunk)

#         resp = client.chat.completions.create(
#             model=model,
#             messages=[
#                 {"role": "system", "content": sys_inst_template},
#                 {"role": "user", "content": f"æ¬¡ã®ãƒ†ã‚­ã‚¹ãƒˆï¼ˆã“ã®ãƒšãƒ¼ã‚¸ã®ã¿ï¼‰ã‚’è§£æã—ã¦ãã ã•ã„ï¼š\n---\n{page_text}"},
#             ],
#         )
#         md_tables.append(resp.choices[0].message.content.strip())

#     out = []
#     for i, tbl in enumerate(md_tables, 1):
#         out.append(f"#### é  {i}\n\n{tbl}\n")
#     return "\n".join(out)

def analyze_issues(model: str, lines: List[str], lines_per_page: int, mode: str, extra: str) -> str:
    """
    ãƒšãƒ¼ã‚¸å˜ä½ã§ GPT ã‹ã‚‰ Markdown è¡¨ã‚’å–å¾—ã—ã€
    ã¾ã¨ã‚ã¦ 1 å€‹ã®å¤§ããªè¡¨ï¼ˆãƒ˜ãƒƒãƒ€ãƒ¼ 1 å›ã ã‘ï¼‰ã«çµ„ã¿ç›´ã—ã¦è¿”ã™ã€‚
    """
    client = openai_client()
    md_tables: List[str] = []
    total_pages = (len(lines) + lines_per_page - 1) // lines_per_page

    # System ãƒ—ãƒ­ãƒ³ãƒ—ãƒˆ
    sys_inst_template = build_system_prompt(mode=mode, extra=extra)

    # --- 1) GPT ã«ãƒšãƒ¼ã‚¸ã”ã¨æŠ•ã’ã¦ Markdown è¡¨ã‚’é›†ã‚ã‚‹ ---
    for pg in range(total_pages):
        start = pg * lines_per_page
        end = min((pg + 1) * lines_per_page, len(lines))
        page_chunk = [
            f"[{(i // lines_per_page) + 1}:{(i % lines_per_page) + 1:02d}] {lines[i]}"
            for i in range(start, end)
        ]
        page_text = "\n".join(page_chunk)

        resp = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": sys_inst_template},
                {"role": "user", "content": f"æ¬¡ã®ãƒ†ã‚­ã‚¹ãƒˆï¼ˆã“ã®ãƒšãƒ¼ã‚¸ã®ã¿ï¼‰ã‚’è§£æã—ã¦ãã ã•ã„ï¼š\n---\n{page_text}"},
            ],
        )
        md_tables.append(resp.choices[0].message.content.strip())

    # --- 2) ã¾ã¨ã‚ã¦ãƒ‘ãƒ¼ã‚¹ã—ã¦ items ã«ã™ã‚‹ ---
    raw_md = "\n\n".join(md_tables)
    items = _parse_plan_md_tables(raw_md)

    # âœ… ãƒ˜ãƒƒãƒ€ãƒ¼ãã£ãã‚Šãªè¡Œï¼ˆé /è¡Œ/é‡è¦åº¦/åŸæ–‡/ä¿®æ­£æ¡ˆ/ç†ç”±ï¼‰ã‚’é™¤å¤–ã™ã‚‹
    def _is_header_like(row: Dict[str, str]) -> bool:
        def norm(s: str) -> str:
            return (s or "").strip().lower()

        return (
            norm(row.get("é ", "")) in {"é ", "page", "ãƒšãƒ¼ã‚¸"} and
            norm(row.get("è¡Œ", "")) in {"è¡Œ", "line"} and
            norm(row.get("é‡è¦åº¦", "")) in {"é‡è¦åº¦", "issue", "å•é¡Œç‚¹"} and
            norm(row.get("åŸæ–‡", "")) in {"åŸæ–‡", "original"} and
            norm(row.get("ä¿®æ­£æ¡ˆ", "")) in {"ä¿®æ­£æ¡ˆ", "suggestion", "ææ¡ˆ"} and
            norm(row.get("ç†ç”±", "")) in {"ç†ç”±", "reason", "æ ¹æ‹ "}
        )

    items = [row for row in items if not _is_header_like(row)]

    # --- 3) items ã‹ã‚‰ã€Œãƒ˜ãƒƒãƒ€ãƒ¼1å›ã ã‘ã® Markdown è¡¨ã€ã‚’å†æ§‹ç¯‰ ---
    def esc_cell(s: str) -> str:
        s = (s or "").replace("|", r"\|")
        s = s.replace("\n", "<br>")
        return s


    header = "| é  | è¡Œ | é‡è¦åº¦ | åŸæ–‡ | ä¿®æ­£æ¡ˆ | ç†ç”± |"
    sep    = "| --- | --- | --- | --- | --- | --- |"

    rows = [header, sep]
    for it in items:
        rows.append(
            "| {page} | {line} | {imp} | {orig} | {sugg} | {reason} |".format(
                page=esc_cell(it.get("é ", "")),
                line=esc_cell(it.get("è¡Œ", "")),
                imp=esc_cell(it.get("é‡è¦åº¦", "")),
                orig=esc_cell(it.get("åŸæ–‡", "")),
                sugg=esc_cell(it.get("ä¿®æ­£æ¡ˆ", "")),
                reason=esc_cell(it.get("ç†ç”±", "")),
            )
        )

    # ã“ã®æ–‡å­—åˆ—ãŒ HTML ãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ç”¨ plan_md ã‹ã¤ PDF è§£æç”¨ã®ã‚½ãƒ¼ã‚¹ã«ãªã‚‹
    return "\n".join(rows)


# ------------------------------------------------------------
# ç”»é¢ï¼ˆè§£æã®ã¿ï¼‰
# ------------------------------------------------------------
st.title("ğŸ“ æ–‡ç« ã®æ ¡æ­£")
st.caption(
    "æ•°æ®µè½ç¨‹åº¦ã®æ–‡ç« ã‚’è²¼ã‚Šä»˜ã‘ã¦ï¼Œè§£æã®ãƒœã‚¿ãƒ³ã‚’æŠ¼ã—ã¦ãã ã•ã„ï¼è§£æã«ã¯å°‘ã—æ™‚é–“ãŒã‹ã‹ã‚Šã¾ã™ï¼"
    "æ ¡æ­£ã«ã¯AIã‚’ä½¿ç”¨ã—ã¦ã„ã¾ã™ã®ã§ï¼Œå€‹äººæƒ…å ±ã‚„æ©Ÿå¯†æ–‡æ›¸ã®å–ã‚Šæ‰±ã„ã«ã¯æ³¨æ„ã—ã¦ãã ã•ã„ã€‚"
)
st.write("æœ¬æ–‡ã‚’å…¥åŠ›ã—ã¦ **â‘  è§£æ** ã‚’å®Ÿè¡Œã™ã‚‹ã¨ã€ãƒšãƒ¼ã‚¸/è¡Œ/ç†ç”±ã¤ãã®æ ¡æ­£æ–¹é‡ï¼ˆMarkdownè¡¨ï¼‰ã‚’ç”Ÿæˆã—ã¾ã™ã€‚")

render_proof_policy_logic_expander()

# 1) åˆæœŸåŒ–ï¼ˆæœ€åˆã®1å›ã ã‘ï¼‰
if "chat_model" not in st.session_state:
    st.session_state["chat_model"] = DEFAULT_MODEL


with st.sidebar:
    st.header("è¨­å®š")
    st.radio(
        "ğŸ§  ä½¿ç”¨ãƒ¢ãƒ‡ãƒ«",
        MODEL_OPTIONS,
        key="chat_model",
    )
    st.selectbox(
        "ğŸ›  è§£æãƒ¢ãƒ¼ãƒ‰",
        list(MODE_DEFS.keys()),
        index=list(MODE_DEFS.keys()).index(st.session_state.get("proof_mode", DEFAULT_MODE)),
        key="proof_mode",
        help="\n\n".join([f"ãƒ»{k}: {v['desc']}" for k, v in MODE_DEFS.items()]),
    )
    lpp = st.number_input("ãƒšãƒ¼ã‚¸è¡Œæ•°ï¼ˆè¡¨ç¤ºç”¨ï¼‰", min_value=20, max_value=100, value=LINES_PER_PAGE, step=5)

    # ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰å½¢å¼ï¼ˆPDF/Wordã®ã¿ã€‚ãƒ‡ãƒ•ã‚©ãƒ«ãƒˆPDFï¼‰
    _DL_LABELS = {"pdf": "PDF (.pdf)", "word": "Word (.docx)"}
    dl_choice_key = st.radio(
        "ğŸ“¦ ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰å½¢å¼ï¼ˆè§£æãƒ¬ãƒãƒ¼ãƒˆï¼‰",
        options=list(_DL_LABELS.keys()),
        format_func=lambda k: _DL_LABELS[k],
        index=0,  # ãƒ‡ãƒ•ã‚©ãƒ«ãƒˆ PDF
        key="dl_format_radio",
    )

extra_prompt = render_policy_preview(mode=st.session_state["proof_mode"])
st.markdown("---")

# ===== å…¥åŠ›ï¼ˆãƒ•ã‚¡ã‚¤ãƒ« / è²¼ã‚Šä»˜ã‘ï¼‰ =====
tab_paste, tab_file = st.tabs(["ğŸ“ è²¼ã‚Šä»˜ã‘ãƒ†ã‚­ã‚¹ãƒˆ","ğŸ“ ãƒ•ã‚¡ã‚¤ãƒ«ã‹ã‚‰"])

src_text = ""
used_file_name = None

# å…ˆã«åˆæœŸå€¤ï¼ˆå¿…è¦ãªã‚‰ä¸€åº¦ã ã‘ï¼‰
if "pasted_text" not in st.session_state:
    st.session_state["pasted_text"] = ""

# def _on_paste():
#     st.session_state["has_paste"] = bool(st.session_state["pasted_text"].strip())


with tab_file:
    col_u, col_btn1 = st.columns([3, 1])
    with col_u:
        up = st.file_uploader(".docx / .txt / .pdf ã‚’ã‚¢ãƒƒãƒ—ãƒ­ãƒ¼ãƒ‰", type=["docx", "txt", "pdf"])
    with col_btn1:
        do_analyze_file = st.button("â‘  è§£æï¼ˆãƒ•ã‚¡ã‚¤ãƒ«ï¼‰", type="primary", use_container_width=True, disabled=not up, key="btn_analyze_file")

    if up:
        used_file_name = up.name
        name = up.name.lower()
        if name.endswith(".pdf"):
            data = up.read()
            try:
                stats = extract_pdf_text(data)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

            st.subheader("ğŸ“„ PDFãƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼")
            display_pdf_bytes(data, height=600)

            if int(stats.get("visible", 0)) < 20:
                st.warning("ã“ã®PDFã¯ç”»åƒPDFï¼ˆãƒ†ã‚­ã‚¹ãƒˆå±¤ãªã—ï¼‰ã¨åˆ¤å®šã—ã¾ã—ãŸã€‚OCRãƒ„ãƒ¼ãƒ«ã§ãƒ†ã‚­ã‚¹ãƒˆåŒ–ã—ã¦ã‹ã‚‰å†åº¦ãŠè©¦ã—ãã ã•ã„ã€‚")
                st.stop()
            else:
                src_text = (stats.get("text") or "").strip()
        else:
            try:
                src_text = load_text_generic(up)
            except RuntimeError as e:
                st.error(str(e)); st.stop()

with tab_paste:
   
    pasted = st.text_area(
        "ã“ã“ã«æœ¬æ–‡ã‚’è²¼ã‚Šä»˜ã‘",
        height=260,
        key="pasted_text",
        placeholder="ã“ã“ã«æœ¬æ–‡ã‚’è²¼ã‚Šä»˜ã‘ã¦ãã ã•ã„ï¼ˆæ”¹è¡Œã¯ä¿æŒã•ã‚Œã¾ã™ï¼‰ã€‚",
    )

    do_analyze_paste = st.button("â‘  è§£æï¼ˆè²¼ã‚Šä»˜ã‘ï¼‰", type="primary", use_container_width=True)

    if do_analyze_paste:
        if not pasted.strip():
            st.warning("ãƒ†ã‚­ã‚¹ãƒˆã‚’è²¼ã‚Šä»˜ã‘ã¦ãã ã•ã„ã€‚")
            st.stop()

        src_text = load_text_from_paste(
            pasted,
            normalize=True,
            collapse_blanks=False,
            keep_blank_lines=1,
            trim_trailing=True,
        )
        used_file_name = "pasted_text.txt"


# ===== è§£æã®å®Ÿè¡Œ =====
if src_text:
    lines = to_numbered_lines(src_text)
    st.subheader("ğŸ‘€ è¡Œç•ªå·ä»˜ããƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ï¼ˆãƒ†ã‚­ã‚¹ãƒˆè¡¨ç¤ºï¼‰")
    st.caption(f"è¡¨ç¤ºä¸Šã®ãƒšãƒ¼ã‚¸è¡Œæ•°: {LINES_PER_PAGE} è¡Œ/ãƒšãƒ¼ã‚¸ï¼ˆæ“¬ä¼¼å‰²ã‚Šä»˜ã‘ï¼‰")
    st.text_area("åŸæ–‡ï¼ˆç•ªå·ä»˜ããƒ—ãƒ¬ãƒ“ãƒ¥ãƒ¼ï¼‰", value=render_preview_with_numbers(lines, LINES_PER_PAGE), height=260)

    want_analyze = (locals().get("do_analyze_file") or locals().get("do_analyze_paste"))
    if want_analyze:
        with st.spinner("è§£æä¸­ï¼ˆæ ¡æ­£æ–¹é‡ã‚’æŠ½å‡ºï¼‰â€¦"):
            plan_md = analyze_issues(
                st.session_state["chat_model"], lines, LINES_PER_PAGE,
                mode=st.session_state["proof_mode"], extra=extra_prompt
            )
        st.success("è§£æãŒå®Œäº†ã—ã¾ã—ãŸã€‚ãƒšãƒ¼ã‚¸/è¡Œ/ç†ç”±ã¤ãã§æ–¹é‡ã‚’è¡¨ç¤ºã—ã¾ã™ã€‚")
        st.subheader("ğŸ“‹ æ ¡æ­£æ–¹é‡ï¼ˆã¾ãšä½•ã‚’ã©ã†ç›´ã™ã‹ï¼‰")
        st.markdown(plan_md, unsafe_allow_html=False)

        # â–¼â–¼ è§£æãƒ¬ãƒãƒ¼ãƒˆã®ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰ï¼ˆPDF or Wordï¼‰ â–¼â–¼
        st.markdown("### â¤µï¸ è§£æãƒ¬ãƒãƒ¼ãƒˆã‚’ãƒ€ã‚¦ãƒ³ãƒ­ãƒ¼ãƒ‰")
        numbered_preview = render_preview_with_numbers(lines, LINES_PER_PAGE)
        file_base = (used_file_name or "pasted_text").rsplit(".", 1)[0]
        file_stub = f"æ ¡æ­£çµæœ_{file_base}"

        if dl_choice_key == "pdf":
            pdf_bytes = build_policy_pdf_bytes(
                original_numbered_preview=numbered_preview,
                plan_md=plan_md,
                model=st.session_state["chat_model"],
                mode=st.session_state["proof_mode"],
                extra_prompt=extra_prompt,
                src_name=used_file_name or "pasted_text.txt",
            )
            if pdf_bytes:
                st.download_button(
                    "PDFï¼ˆ.pdfï¼‰ã¨ã—ã¦ä¿å­˜",
                    data=pdf_bytes,
                    file_name=f"{file_stub}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                    key=f"dl_pdf_{file_stub}",
                )
            else:
                st.warning("PDF ã‚’ç”Ÿæˆã§ãã¾ã›ã‚“ã§ã—ãŸã€‚`pip install reportlab` ã‚’å®Ÿè¡Œã—ã€CIDãƒ•ã‚©ãƒ³ãƒˆï¼ˆHeiseiMin/HeiseiKakuGoï¼‰ãŒä½¿ãˆã‚‹ã‹ç¢ºèªã—ã¦ãã ã•ã„ã€‚")

        else:  # word
            data_docx, ext = build_policy_docx_bytes(
                original_numbered_preview=numbered_preview,
                plan_md=plan_md,
                model=st.session_state["chat_model"],
                mode=st.session_state["proof_mode"],
                extra_prompt=extra_prompt,
                src_name=used_file_name or "pasted_text.txt",
            )
            st.download_button(
                "Wordï¼ˆ.docxï¼‰ã¨ã—ã¦ä¿å­˜" if ext == ".docx" else "ãƒ†ã‚­ã‚¹ãƒˆï¼ˆ.txtï¼‰ã¨ã—ã¦ä¿å­˜",
                data=data_docx,
                file_name=f"{file_stub}{ext}",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == ".docx" else "text/plain",
                use_container_width=True,
                key=f"dl_word_{file_stub}",
            )

else:
    st.info("å…¥åŠ›ã‚¿ãƒ–ï¼ˆğŸ“/ğŸ“ï¼‰ã‹ã‚‰æœ¬æ–‡ã‚’æŒ‡å®šã—ã¦ã€â‘  è§£æã€ã‚’å®Ÿè¡Œã—ã¦ãã ã•ã„ã€‚")
